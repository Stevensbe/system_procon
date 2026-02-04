"""
Gera manual do usuário por módulo (DOCX) para o sistema PROCON-AM.
Inclui telas, campos específicos, fluxos e espaço para prints.
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from pathlib import Path


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(26)


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)


def add_meta(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(10)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(f"• {item}")


def add_section(doc, title, paragraphs=None, bullets=None):
    doc.add_heading(title, level=1)
    if paragraphs:
        for text in paragraphs:
            doc.add_paragraph(text)
    if bullets:
        add_bullets(doc, bullets)


def add_subsection(doc, title, paragraphs=None, bullets=None):
    doc.add_heading(title, level=2)
    if paragraphs:
        for text in paragraphs:
            doc.add_paragraph(text)
    if bullets:
        add_bullets(doc, bullets)


def add_print_placeholder(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True


def try_add_image(doc, path, width=Inches(6.5)):
    if path and Path(path).exists():
        doc.add_paragraph()
        doc.add_picture(str(path), width=width)
        return True
    return False


def main():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # Capa
    doc.add_paragraph()
    add_title(doc, "MANUAL DO USUÁRIO")
    add_subtitle(doc, "Sistema PROCON-AM")
    doc.add_paragraph()
    add_meta(
        doc,
        "Módulos: Fiscalização, Processos/Tramitação, Jurídico 1 e 2, Cobrança/GRM, PPA, "
        "Notificações, Caixa de Entrada, Portal Cidadão/Empresa, Atendimento",
    )
    add_meta(doc, f"Data: {datetime.now().strftime('%d/%m/%Y')}")
    doc.add_page_break()

    # Sumário
    doc.add_heading("SUMÁRIO", level=1)
    add_bullets(
        doc,
        [
            "1. Fiscalização",
            "2. Processos e Tramitação",
            "3. Jurídico 1 (Análise e Decisão)",
            "4. Jurídico 2 (Recursos)",
            "5. Cobrança / GRM",
            "6. PPA",
            "7. Notificações",
            "8. Caixa de Entrada",
            "9. Portal Cidadão / Empresa",
            "10. Atendimento",
        ],
    )
    doc.add_page_break()

    # ======================= 1. Fiscalização =======================
    add_section(
        doc,
        "1. Fiscalização",
        paragraphs=[
            "Módulo responsável pela lavratura de autos, coleta de dados em campo e emissão de notificações.",
        ],
    )
    add_subsection(
        doc,
        "1.1 Telas principais",
        bullets=[
            "Listagem de Autos (Banco / Supermercado / Posto / Diversos)",
            "Detalhe do Auto",
            "Auto de Infração",
            "Auto de Apreensão/Inutilização",
            "Notificações de Fiscalização",
        ],
    )
    add_subsection(
        doc,
        "1.2 Campos principais (por auto)",
        bullets=[
            "Dados do estabelecimento: Razão Social, Nome Fantasia, CNPJ, Endereço, Município/UF, CEP, Telefone.",
            "Dados do fiscal: Nome/ matrícula, assinatura.",
            "Dados do Auto: número, data/hora, local, tipo de irregularidade, cominação legal.",
            "Texto livre (relatório/observações) quando aplicável.",
        ],
    )
    add_subsection(
        doc,
        "1.3 Fluxo básico",
        bullets=[
            "Novo auto → Preencher campos → Salvar.",
            "Gerar documento (DOCX/PDF) quando necessário.",
            "Se houver infração: gerar Auto de Infração (AI).",
            "Se houver apreensão: criar Auto de Apreensão/Inutilização.",
            "Emitir Notificação de Fiscalização e registrar no processo.",
        ],
    )
    add_subsection(
        doc,
        "1.4 Validações e regras",
        bullets=[
            "CNPJ pode preencher automaticamente dados da empresa (quando disponível).",
            "Assinaturas ficam vinculadas ao auto e devem aparecer no detalhe.",
            "AI pode derivar do AC ou de parecer técnico.",
        ],
    )
    add_subsection(
        doc,
        "1.5 Prints",
        paragraphs=["(Adicionar prints reais das telas: Listagem, Detalhe, Formulário, Notificação)"],
    )
    add_print_placeholder(doc, "PRINT SUGERIDO: Listagem de Autos (Fiscalização)")
    add_print_placeholder(doc, "PRINT SUGERIDO: Formulário de Auto")
    add_print_placeholder(doc, "PRINT SUGERIDO: Detalhe do Auto")

    # ======================= 2. Processos e Tramitação =======================
    doc.add_page_break()
    add_section(
        doc,
        "2. Processos e Tramitação",
        paragraphs=[
            "Centraliza o dossiê do processo administrativo e permite tramitação entre setores e caixas pessoais.",
        ],
    )
    add_subsection(
        doc,
        "2.1 Telas principais",
        bullets=[
            "Lista de Processos",
            "Detalhe/Dossiê do Processo",
            "Tramitação (por setor ou caixa pessoal)",
        ],
    )
    add_subsection(
        doc,
        "2.2 Campos e informações exibidas",
        bullets=[
            "Número do processo (padrão SEI).",
            "Interessado, assunto, status atual.",
            "Histórico de tramitações.",
            "Documentos anexados (AC, AI, notificações, pareceres, decisões).",
        ],
    )
    add_subsection(
        doc,
        "2.3 Regras",
        bullets=[
            "Ao tramitar, o processo sai da fila do setor atual e aparece na fila do destino.",
            "Tramitação pode ser para setor ou caixa pessoal.",
        ],
    )
    add_subsection(
        doc,
        "2.4 Prints",
        paragraphs=["(Adicionar prints reais da lista e do dossiê)"],
    )
    add_print_placeholder(doc, "PRINT SUGERIDO: Lista de Processos")
    add_print_placeholder(doc, "PRINT SUGERIDO: Detalhe/Dossiê do Processo")

    # ======================= 3. Jurídico 1 =======================
    doc.add_page_break()
    add_section(
        doc,
        "3. Jurídico 1 (Análise e Decisão)",
        paragraphs=[
            "Responsável pela análise da defesa e emissão de parecer/decisão em 1ª instância.",
        ],
    )
    add_subsection(
        doc,
        "3.1 Telas principais",
        bullets=[
            "Petições Jurídico 1 (entrada de defesa)",
            "Detalhe da Petição",
            "Upload de Parecer/Decisão",
        ],
    )
    add_subsection(
        doc,
        "3.2 Campos principais",
        bullets=[
            "Número do processo",
            "Dados do autuado (CPF/CNPJ)",
            "Tipo de petição",
            "Arquivos anexos (DOCX/PDF)",
        ],
    )
    add_subsection(
        doc,
        "3.3 Regras",
        bullets=[
            "Após anexar decisão, é possível notificar a disponibilidade para download no portal.",
            "Status da petição muda para “Respondida” após notificar.",
            "Improcedente → arquivamento do processo.",
        ],
    )
    add_subsection(
        doc,
        "3.4 Prints",
        paragraphs=["(Adicionar prints reais da lista e do detalhe da petição)"],
    )
    add_print_placeholder(doc, "PRINT SUGERIDO: Petições Jurídico 1")
    add_print_placeholder(doc, "PRINT SUGERIDO: Detalhe/Upload de Decisão")

    # ======================= 4. Jurídico 2 =======================
    doc.add_page_break()
    add_section(
        doc,
        "4. Jurídico 2 (Recursos)",
        paragraphs=[
            "Responsável pela análise de recursos (2ª instância).",
        ],
    )
    add_subsection(
        doc,
        "4.1 Telas principais",
        bullets=[
            "Recursos Jurídico 2",
            "Detalhe do Recurso",
            "Upload de Parecer/Decisão (2ª instância)",
        ],
    )
    add_subsection(
        doc,
        "4.2 Regras",
        bullets=[
            "Recurso só aparece após decisão do Jurídico 1 e quando status estiver “aguardando recurso”.",
            "Decisão pode manter ou reformar a decisão anterior (redução/anulação).",
        ],
    )
    add_subsection(
        doc,
        "4.3 Prints",
        paragraphs=["(Adicionar prints reais da lista e do detalhe do recurso)"],
    )
    add_print_placeholder(doc, "PRINT SUGERIDO: Lista de Recursos Jurídico 2")
    add_print_placeholder(doc, "PRINT SUGERIDO: Detalhe/Decisão do Recurso")

    # ======================= 5. Cobrança / GRM =======================
    doc.add_page_break()
    add_section(
        doc,
        "5. Cobrança / GRM",
        paragraphs=[
            "Módulo responsável por gerar GRM, registrar pagamento e encaminhar dívida.",
        ],
    )
    add_subsection(
        doc,
        "5.1 Telas principais",
        bullets=[
            "Lista de GRM",
            "Detalhe da GRM",
            "Emissão de documento e notificação",
        ],
    )
    add_subsection(
        doc,
        "5.2 Regras",
        bullets=[
            "GRM só é gerada após o DAF receber o despacho.",
            "Status: aguardando, pago, parcelado, inadimplente.",
        ],
    )
    add_subsection(
        doc,
        "5.3 Prints",
        paragraphs=["(Adicionar prints reais da lista e detalhe da GRM)"],
    )
    add_print_placeholder(doc, "PRINT SUGERIDO: Lista de GRM")
    add_print_placeholder(doc, "PRINT SUGERIDO: Detalhe da GRM")

    # ======================= 6. PPA =======================
    doc.add_page_break()
    add_section(
        doc,
        "6. PPA",
        paragraphs=[
            "O PPA funciona como capa do processo e consolida informações importantes.",
        ],
    )
    add_subsection(
        doc,
        "6.1 Telas principais",
        bullets=[
            "Lista de PPA",
            "Detalhe do PPA",
            "Geração de DOCX/PDF",
        ],
    )
    add_subsection(
        doc,
        "6.2 Campos principais",
        bullets=[
            "Número do PPA (001/2026).",
            "Vínculo com processo/auto.",
            "Histórico e documentos anexados.",
        ],
    )
    add_subsection(
        doc,
        "6.3 Prints",
        paragraphs=["(Adicionar prints reais da lista e detalhe do PPA)"],
    )
    add_print_placeholder(doc, "PRINT SUGERIDO: Lista de PPA")
    add_print_placeholder(doc, "PRINT SUGERIDO: Detalhe do PPA")

    # ======================= 7. Notificações =======================
    doc.add_page_break()
    add_section(
        doc,
        "7. Notificações",
        paragraphs=[
            "Registro de notificações administrativas com envio manual e histórico.",
        ],
    )
    add_subsection(
        doc,
        "7.1 Telas principais",
        bullets=[
            "Listagem de notificações",
            "Criação de notificação",
            "Ações: visualizar, editar, excluir, enviar",
        ],
    )
    add_subsection(
        doc,
        "7.2 Prints",
        paragraphs=["(Adicionar prints reais da listagem e do formulário)"],
    )
    add_print_placeholder(doc, "PRINT SUGERIDO: Listagem de Notificações")
    add_print_placeholder(doc, "PRINT SUGERIDO: Formulário de Notificação")

    # ======================= 8. Caixa de Entrada =======================
    doc.add_page_break()
    add_section(
        doc,
        "8. Caixa de Entrada",
        paragraphs=[
            "Caixa unificada com itens do setor e da caixa pessoal.",
        ],
    )
    add_subsection(
        doc,
        "8.1 Operações comuns",
        bullets=[
            "Marcar como recebido",
            "Abrir processo",
            "Tramitar para setor ou caixa pessoal",
        ],
    )
    add_subsection(
        doc,
        "8.2 Prints",
        paragraphs=["(Adicionar prints reais da caixa de entrada)"],
    )
    add_print_placeholder(doc, "PRINT SUGERIDO: Caixa de Entrada")

    # ======================= 9. Portal Cidadão / Empresa =======================
    doc.add_page_break()
    add_section(
        doc,
        "9. Portal Cidadão / Empresa",
        paragraphs=[
            "Portal externo para acompanhamento de processos e envio de petições.",
        ],
    )
    add_subsection(
        doc,
        "9.1 Acompanhar Processo",
        bullets=[
            "Informar número do processo e CPF/CNPJ do autuado.",
            "Visualizar status e documentos liberados.",
        ],
    )
    add_subsection(
        doc,
        "9.2 Peticionamento",
        bullets=[
            "Tipos: Defesa, Juntada de Documento e Recurso (quando habilitado).",
            "Arquivos permitidos: DOCX e PDF.",
        ],
    )
    add_subsection(
        doc,
        "9.3 Prints",
        paragraphs=["(Adicionar prints reais do portal)"],
    )
    add_print_placeholder(doc, "PRINT SUGERIDO: Portal - Acompanhar Processo")
    add_print_placeholder(doc, "PRINT SUGERIDO: Portal - Peticionamento")

    # ======================= 10. Atendimento =======================
    doc.add_page_break()
    add_section(
        doc,
        "10. Atendimento",
        paragraphs=[
            "Registro de reclamações e gestão de filas.",
        ],
    )
    add_subsection(
        doc,
        "10.1 Telas principais",
        bullets=[
            "Fila de atendimento",
            "Novo atendimento/reclamação",
            "Painel LGPD",
        ],
    )
    add_subsection(
        doc,
        "10.2 Campos principais",
        bullets=[
            "Dados do consumidor",
            "Empresa reclamada",
            "Descrição e anexos",
            "Prazos e status",
        ],
    )
    add_subsection(
        doc,
        "10.3 Prints",
        paragraphs=["(Adicionar prints reais das telas de atendimento)"],
    )
    add_print_placeholder(doc, "PRINT SUGERIDO: Fila de Atendimento")
    add_print_placeholder(doc, "PRINT SUGERIDO: Formulário de Atendimento")

    # Fluxogramas oficiais (se existirem)
    doc.add_page_break()
    add_section(doc, "Apêndice - Fluxogramas (Referência)")
    if not try_add_image(doc, "docs/_tmp_doc0066_page1.png"):
        add_print_placeholder(doc, "Fluxograma oficial: adicionar imagem se disponível.")

    # Salvar
    output_dir = Path("docs")
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / "MANUAL_USUARIO_MODULOS_PROCON_DETALHADO.docx"
    doc.save(output_path)
    print(f"Manual gerado em: {output_path}")


if __name__ == "__main__":
    main()
