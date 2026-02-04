"""
Script para gerar documentação DOCX do projeto SISPROCON
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

def criar_documento():
    doc = Document()
    
    # Configurar estilos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # === CAPA ===
    doc.add_paragraph()
    doc.add_paragraph()
    
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("📋 SISPROCON")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0, 102, 153)
    
    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitulo.add_run("Sistema de Proteção ao Consumidor")
    run.bold = True
    run.font.size = Pt(24)
    
    doc.add_paragraph()
    
    doc_titulo = doc.add_paragraph()
    doc_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = doc_titulo.add_run("DOCUMENTAÇÃO OFICIAL COMPLETA DO PROJETO")
    run.bold = True
    run.font.size = Pt(18)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Info box
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Versão: 1.0.0\n").bold = True
    info.add_run("Data: Janeiro de 2026\n")
    info.add_run("Órgão: PROCON-AM (Amazonas)\n")
    info.add_run("Status: Produção")
    
    doc.add_page_break()
    
    # === SUMÁRIO ===
    doc.add_heading('SUMÁRIO', 0)
    
    sumario = [
        "1. Visão Geral",
        "2. Arquitetura do Sistema",
        "3. Tecnologias Utilizadas",
        "4. Estrutura do Projeto",
        "5. Módulos do Sistema",
        "   5.1 Módulo de Fiscalização",
        "   5.2 Módulo de Protocolo",
        "   5.3 Módulo de Multas",
        "   5.4 Módulo de Atendimento",
        "   5.5 Módulo PPA",
        "   5.6 Módulo Jurídico",
        "   5.7 Portal do Cidadão",
        "   5.8 Portal da Empresa",
        "   5.9 Módulo de Triagem",
        "   5.10 Módulo de Auditoria",
        "   5.11 Módulo de Notificações",
        "   5.12 Business Intelligence",
        "   5.13 Caixa de Entrada",
        "6. Sistema de Autenticação e Permissões",
        "7. API REST",
        "8. Frontend React",
        "9. Infraestrutura e DevOps",
        "10. Considerações de Segurança",
        "11. Guia de Implantação",
    ]
    
    for item in sumario:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.5) if item.startswith("   ") else Inches(0)
    
    doc.add_page_break()
    
    # === 1. VISÃO GERAL ===
    doc.add_heading('1. Visão Geral', 1)
    
    doc.add_paragraph(
        "O SISPROCON (Sistema de Proteção ao Consumidor) é uma plataforma completa de gestão "
        "para órgãos de defesa do consumidor, desenvolvida especificamente para o PROCON-AM (Amazonas). "
        "O sistema automatiza e integra todos os processos de atendimento ao consumidor, fiscalização "
        "de estabelecimentos comerciais, gestão de processos administrativos, aplicação de multas e "
        "geração de relatórios executivos."
    )
    
    doc.add_heading('Objetivos Principais', 2)
    objetivos = [
        "Digitalização completa dos processos do PROCON",
        "Gestão unificada de atendimentos, reclamações e denúncias",
        "Automatização de fluxos de trabalho e prazos legais",
        "Fiscalização móvel com sincronização em tempo real",
        "Portais externos para cidadãos e empresas",
        "Business Intelligence com dashboards executivos",
        "Conformidade com LGPD e legislação consumerista",
    ]
    for obj in objetivos:
        doc.add_paragraph(f"• {obj}")
    
    doc.add_heading('Público-Alvo', 2)
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Perfil'
    hdr[1].text = 'Descrição'
    
    dados = [
        ('Administradores', 'Gestão completa do sistema, configurações e usuários'),
        ('Staff/Servidores', 'Operação diária, atendimento e fiscalização'),
        ('Empresas', 'Acesso ao portal para responder reclamações e CIPs'),
        ('Cidadãos', 'Portal público para registro de reclamações e consultas'),
    ]
    for i, (perfil, desc) in enumerate(dados, 1):
        table.rows[i].cells[0].text = perfil
        table.rows[i].cells[1].text = desc
    
    doc.add_page_break()
    
    # === 2. ARQUITETURA ===
    doc.add_heading('2. Arquitetura do Sistema', 1)
    
    doc.add_paragraph(
        "O SISPROCON utiliza uma arquitetura moderna baseada em monolito modular no backend "
        "com SPA (Single Page Application) no frontend."
    )
    
    doc.add_heading('Componentes Principais', 2)
    
    componentes = [
        ("Frontend Layer", "React SPA com Vite e JavaScript/JSX"),
        ("API Layer", "Django REST Framework com JWT Authentication e Swagger/OpenAPI"),
        ("Backend Layer", "Django 5.x com Celery Workers para tarefas assíncronas"),
        ("Data Layer", "PostgreSQL (banco principal), Redis (cache), Media Storage"),
        ("External Services", "Email SMTP, Receita Federal API, Prometheus/Grafana"),
    ]
    
    for comp, desc in componentes:
        p = doc.add_paragraph()
        run = p.add_run(f"{comp}: ")
        run.bold = True
        p.add_run(desc)
    
    doc.add_heading('Padrões Arquiteturais', 2)
    padroes = [
        "MVT (Model-View-Template): Padrão Django para o backend",
        "REST API: Comunicação via DRF (Django REST Framework)",
        "JWT: Autenticação stateless com tokens",
        "Repository Pattern: Encapsulamento de acesso a dados",
        "Service Layer: Lógica de negócio separada das views",
    ]
    for pad in padroes:
        doc.add_paragraph(f"• {pad}")
    
    doc.add_page_break()
    
    # === 3. TECNOLOGIAS ===
    doc.add_heading('3. Tecnologias Utilizadas', 1)
    
    doc.add_heading('Backend', 2)
    table = doc.add_table(rows=9, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Tecnologia'
    hdr[1].text = 'Versão'
    hdr[2].text = 'Propósito'
    
    backend_techs = [
        ('Python', '3.12+', 'Linguagem principal'),
        ('Django', '5.x', 'Framework web'),
        ('Django REST Framework', '3.14+', 'API REST'),
        ('PostgreSQL', '15+', 'Banco de dados principal'),
        ('Redis', '7+', 'Cache e filas'),
        ('Celery', '5+', 'Tarefas assíncronas'),
        ('SimpleJWT', '5+', 'Autenticação JWT'),
        ('DRF-Spectacular', '0.27+', 'Documentação OpenAPI'),
    ]
    for i, (tech, ver, prop) in enumerate(backend_techs, 1):
        table.rows[i].cells[0].text = tech
        table.rows[i].cells[1].text = ver
        table.rows[i].cells[2].text = prop
    
    doc.add_paragraph()
    doc.add_heading('Frontend', 2)
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Tecnologia'
    hdr[1].text = 'Versão'
    hdr[2].text = 'Propósito'
    
    frontend_techs = [
        ('React', '18+', 'Framework UI'),
        ('Vite', '5+', 'Build tool'),
        ('Shadcn/UI', '-', 'Componentes UI'),
        ('TailwindCSS', '3+', 'Estilização'),
        ('React Router', '6+', 'Roteamento SPA'),
        ('Axios', '1+', 'Cliente HTTP'),
    ]
    for i, (tech, ver, prop) in enumerate(frontend_techs, 1):
        table.rows[i].cells[0].text = tech
        table.rows[i].cells[1].text = ver
        table.rows[i].cells[2].text = prop
    
    doc.add_page_break()
    
    # === 4. ESTRUTURA DO PROJETO ===
    doc.add_heading('4. Estrutura do Projeto', 1)
    
    estrutura = """system_procon/
├── procon_system/              # Backend Django
│   ├── procon_system/          # Configurações do projeto
│   ├── core/                   # Módulo core (autenticação, permissões)
│   ├── fiscalizacao/           # Autos de fiscalização
│   ├── multas/                 # Gestão de multas
│   ├── protocolo/              # Sistema de protocolos
│   ├── atendimento/            # Atendimento presencial
│   ├── ppa/                    # Procedimento Pré-Administrativo
│   ├── juridico/               # Processos jurídicos
│   ├── portal_cidadao/         # Portal público
│   ├── portal_empresa/         # Portal empresarial
│   ├── triagem/                # Triagem de demandas
│   ├── auditoria/              # Logs e auditoria
│   ├── notificacoes/           # Sistema de notificações
│   ├── caixa_entrada/          # Inbox unificado
│   ├── business_intelligence/  # Dashboards e KPIs
│   └── templates/              # Templates Django
│
├── frontend/                   # Frontend React
│   └── src/
│       ├── components/         # Componentes reutilizáveis
│       ├── pages/              # Páginas da aplicação
│       ├── services/           # Serviços de API
│       └── hooks/              # Custom hooks
│
├── tests/                      # Testes automatizados
└── docker-compose.yml          # Orquestração Docker"""
    
    p = doc.add_paragraph()
    run = p.add_run(estrutura)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    
    doc.add_page_break()
    
    # === 5. MÓDULOS DO SISTEMA ===
    doc.add_heading('5. Módulos do Sistema', 1)
    
    # 5.1 Fiscalização
    doc.add_heading('5.1 Módulo de Fiscalização', 2)
    doc.add_paragraph(
        "O módulo de fiscalização é o coração operacional do SISPROCON, responsável por gerenciar "
        "todas as ações de fiscalização realizadas pelos agentes do PROCON."
    )
    
    doc.add_heading('Modelos Principais', 3)
    modelos_fisc = [
        ('AutoConstatacaoBase', 'Classe abstrata base para todos os autos'),
        ('AutoBanco', 'Auto de fiscalização de agências bancárias'),
        ('AutoPosto', 'Auto de fiscalização de postos de combustível'),
        ('AutoSupermercado', 'Auto de fiscalização de supermercados'),
        ('AutoInfracao', 'Auto de infração lavrado'),
        ('Processo', 'Processo administrativo'),
    ]
    for modelo, desc in modelos_fisc:
        p = doc.add_paragraph()
        run = p.add_run(f"• {modelo}: ")
        run.bold = True
        p.add_run(desc)
    
    doc.add_heading('Funcionalidades', 3)
    funcs_fisc = [
        "Registro de autos por tipo de estabelecimento",
        "Sincronização offline com app mobile",
        "Numeração automática sequencial única",
        "Geolocalização dos estabelecimentos",
        "Anexação de fotos e documentos",
        "Geração automática de processo após infração",
    ]
    for func in funcs_fisc:
        doc.add_paragraph(f"✓ {func}")
    
    # 5.2 Protocolo
    doc.add_heading('5.2 Módulo de Protocolo', 2)
    doc.add_paragraph(
        "Sistema completo de gestão de documentos e protocolos, similar ao SIGED governamental."
    )
    
    doc.add_heading('Modelos Principais', 3)
    modelos_prot = [
        ('TipoProtocolo', 'Tipos de protocolo (Denúncia, Reclamação, Consulta, Petição)'),
        ('StatusProtocolo', 'Status possíveis com cores personalizadas'),
        ('Protocolo', 'Registro principal de protocolo'),
        ('DocumentoProtocolo', 'Documentos anexados'),
        ('TramitacaoProtocolo', 'Histórico de movimentações'),
    ]
    for modelo, desc in modelos_prot:
        p = doc.add_paragraph()
        run = p.add_run(f"• {modelo}: ")
        run.bold = True
        p.add_run(desc)
    
    # 5.3 Multas
    doc.add_heading('5.3 Módulo de Multas', 2)
    doc.add_paragraph(
        "Gestão completa do ciclo de vida das multas aplicadas pelo PROCON."
    )
    
    doc.add_heading('Status de Multas', 3)
    status_multas = [
        "Pendente de Pagamento",
        "Parcelada",
        "Paga",
        "Cancelada",
        "Vencida",
        "Em Recurso",
        "Protestada",
    ]
    for status in status_multas:
        doc.add_paragraph(f"• {status}")
    
    # 5.4 Atendimento
    doc.add_heading('5.4 Módulo de Atendimento', 2)
    doc.add_paragraph(
        "Gestão completa do atendimento presencial no PROCON, incluindo totem de autoatendimento "
        "e painel TV para chamada de senhas."
    )
    
    doc.add_heading('Recursos', 3)
    recursos_atend = [
        "Totem de autoatendimento",
        "Painel TV para chamada de senhas",
        "Distribuição automática por regras",
        "Controle de tempo de espera",
        "Conformidade LGPD (consentimento, anonimização)",
    ]
    for rec in recursos_atend:
        doc.add_paragraph(f"✓ {rec}")
    
    # 5.5 PPA
    doc.add_heading('5.5 Módulo PPA (Procedimento Preliminar Administrativo)', 2)
    doc.add_paragraph(
        "Capa do processo onde se anexa todos os documentos (AC, AI, Notificações, Defesas, etc)."
    )
    
    doc.add_heading('Status do PPA', 3)
    status_ppa = [
        "Criado",
        "Em Análise",
        "Aguardando Defesa",
        "Defesa Apresentada",
        "Parecer Emitido",
        "Concluído - Procedente",
        "Concluído - Improcedente",
        "Arquivado",
    ]
    for status in status_ppa:
        doc.add_paragraph(f"• {status}")
    
    # 5.6 Jurídico
    doc.add_heading('5.6 Módulo Jurídico', 2)
    doc.add_paragraph(
        "Gestão de processos jurídicos e análises legais."
    )
    
    doc.add_heading('Modelos Principais', 3)
    modelos_jur = [
        ('AnalistaJuridico', 'Perfil de analista com OAB'),
        ('ProcessoJuridico', 'Processo principal'),
        ('AnaliseJuridica', 'Análises realizadas'),
        ('RespostaJuridica', 'Respostas formais'),
        ('PrazoJuridico', 'Controle de prazos'),
    ]
    for modelo, desc in modelos_jur:
        p = doc.add_paragraph()
        run = p.add_run(f"• {modelo}: ")
        run.bold = True
        p.add_run(desc)
    
    doc.add_page_break()
    
    # 5.7-5.13 Outros módulos
    doc.add_heading('5.7 Portal do Cidadão', 2)
    doc.add_paragraph(
        "Portal público para acesso de cidadãos/consumidores. Permite registro de reclamações, "
        "denúncias, acompanhamento de status e download de formulários oficiais."
    )
    
    doc.add_heading('5.8 Portal da Empresa', 2)
    doc.add_paragraph(
        "Portal para empresas responderem reclamações e notificações. Inclui níveis de permissão "
        "(Visualização, Padrão, Administrador) e tokens de acesso."
    )
    
    doc.add_heading('5.9 Módulo de Triagem', 2)
    doc.add_paragraph(
        "Triagem centralizada de demandas antes da distribuição. Recebe demandas de Portal, "
        "Telefone, Presencial, Email, Ofício e Ouvidoria."
    )
    
    doc.add_heading('5.10 Módulo de Auditoria', 2)
    doc.add_paragraph(
        "Logs e auditoria completa do sistema para conformidade. Inclui log de sessões, "
        "alterações de dados e acessos a recursos."
    )
    
    doc.add_heading('5.11 Módulo de Notificações', 2)
    doc.add_paragraph(
        "Sistema multicanal de notificações (Email, SMS, Push, Sistema) com templates "
        "configuráveis e agendamento de envio."
    )
    
    doc.add_heading('5.12 Business Intelligence', 2)
    doc.add_paragraph(
        "Dashboards executivos e KPIs governamentais. Tipos: Executivo, Operacional, "
        "Tático, Gerencial. Suporta relatórios automáticos."
    )
    
    doc.add_heading('5.13 Caixa de Entrada', 2)
    doc.add_paragraph(
        "Inbox unificado similar ao SIGED governamental. Permite visualização por setor, "
        "distribuição automática e controle de prazos."
    )
    
    doc.add_page_break()
    
    # === 6. AUTENTICAÇÃO ===
    doc.add_heading('6. Sistema de Autenticação e Permissões', 1)
    
    doc.add_heading('Perfis de Usuário (Roles)', 2)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Role'
    hdr[1].text = 'Descrição'
    
    roles = [
        ('admin', 'Administrador - Acesso total ao sistema'),
        ('staff', 'Funcionário PROCON - Módulos operacionais'),
        ('empresa', 'Representante de empresa - Portal Empresa'),
        ('consumer', 'Cidadão/Consumidor - Portal Consumidor'),
    ]
    for i, (role, desc) in enumerate(roles, 1):
        table.rows[i].cells[0].text = role
        table.rows[i].cells[1].text = desc
    
    doc.add_paragraph()
    doc.add_heading('Configuração JWT', 2)
    jwt_config = [
        "ACCESS_TOKEN_LIFETIME: 30 minutos",
        "REFRESH_TOKEN_LIFETIME: 24 horas (1440 minutos)",
        "ROTATE_REFRESH_TOKENS: True",
        "BLACKLIST_AFTER_ROTATION: True",
    ]
    for config in jwt_config:
        doc.add_paragraph(f"• {config}")
    
    doc.add_page_break()
    
    # === 7. API REST ===
    doc.add_heading('7. API REST', 1)
    
    doc.add_heading('Endpoints de Autenticação', 2)
    endpoints_auth = [
        ('POST /auth/token/', 'Obter tokens JWT'),
        ('POST /auth/token/refresh/', 'Renovar access token'),
        ('POST /auth/register/', 'Registrar usuário'),
        ('POST /auth/login/', 'Login simplificado'),
        ('POST /auth/logout/', 'Logout (blacklist token)'),
        ('GET /auth/profile/', 'Perfil do usuário'),
    ]
    for endpoint, desc in endpoints_auth:
        p = doc.add_paragraph()
        run = p.add_run(f"{endpoint} ")
        run.bold = True
        run.font.name = 'Courier New'
        p.add_run(f"- {desc}")
    
    doc.add_heading('Principais Módulos da API', 2)
    modulos_api = [
        '/api/fiscalizacao/',
        '/api/multas/',
        '/api/protocolo/',
        '/api/atendimento/',
        '/api/ppa/',
        '/api/juridico/',
        '/api/portal-consumidor/',
        '/api/portal-empresa/',
        '/api/triagem/',
        '/api/caixa-entrada/',
        '/api/notificacoes/',
        '/api/bi/',
    ]
    for mod in modulos_api:
        p = doc.add_paragraph()
        run = p.add_run(mod)
        run.font.name = 'Courier New'
    
    doc.add_page_break()
    
    # === 8-11. Outras seções ===
    doc.add_heading('8. Frontend React', 1)
    doc.add_paragraph(
        "O frontend possui mais de 150 páginas React organizadas em 32 diretórios e 100+ componentes reutilizáveis."
    )
    
    doc.add_heading('Estrutura Principal', 2)
    estrutura_front = [
        "src/pages/ - 45+ páginas organizadas por módulo",
        "src/components/ - 100+ componentes reutilizáveis",
        "src/services/ - 37 serviços de API",
        "src/hooks/ - 6 custom hooks",
        "src/context/ - 3 context providers",
    ]
    for item in estrutura_front:
        doc.add_paragraph(f"• {item}")
    
    doc.add_heading('9. Infraestrutura e DevOps', 1)
    infra = [
        "Docker / Docker Compose para containerização",
        "Nginx como proxy reverso",
        "PostgreSQL como banco principal",
        "Redis para cache e sessões",
        "Prometheus + Grafana para monitoramento",
        "GitHub Actions para CI/CD",
    ]
    for item in infra:
        doc.add_paragraph(f"• {item}")
    
    doc.add_heading('10. Considerações de Segurança', 1)
    
    doc.add_heading('Conformidade LGPD', 2)
    lgpd = [
        "Registro de consentimento para coleta de dados",
        "Direito ao esquecimento (anonimização de dados)",
        "Logs de acesso a dados sensíveis",
        "Controle de retenção de dados",
    ]
    for item in lgpd:
        doc.add_paragraph(f"✓ {item}")
    
    doc.add_heading('Segurança da API', 2)
    seguranca = [
        "Autenticação JWT com blacklist de tokens",
        "Rate limiting (60 req/min autenticado, 10 req/min público)",
        "CORS configurado por ambiente",
        "Validação de entrada em todos os endpoints",
    ]
    for item in seguranca:
        doc.add_paragraph(f"✓ {item}")
    
    doc.add_page_break()
    
    # === RESUMO ESTATÍSTICO ===
    doc.add_heading('Resumo Estatístico', 1)
    
    table = doc.add_table(rows=10, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Métrica'
    hdr[1].text = 'Valor'
    
    metricas = [
        ('Módulos Backend', '40+'),
        ('Modelos Django', '150+'),
        ('Endpoints API', '200+'),
        ('Páginas Frontend', '45+'),
        ('Componentes React', '100+'),
        ('Linhas de Código Python', '25.000+'),
        ('Linhas de Código JavaScript', '15.000+'),
        ('Testes Automatizados', '500+'),
        ('Cobertura de Testes', '85%+'),
    ]
    for i, (metrica, valor) in enumerate(metricas, 1):
        table.rows[i].cells[0].text = metrica
        table.rows[i].cells[1].text = valor
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Rodapé
    rodape = doc.add_paragraph()
    rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rodape.add_run("© 2025-2026 PROCON-AM - Sistema SISPROCON\n")
    rodape.add_run("Desenvolvido para a proteção do consumidor amazonense")
    
    return doc


def main():
    print("Gerando documentação DOCX...")
    
    doc = criar_documento()
    
    # Salvar no diretório do projeto
    output_path = os.path.join(
        os.path.dirname(os.path.dirname(__file__)),
        'DOCUMENTACAO_OFICIAL_SISPROCON.docx'
    )
    
    doc.save(output_path)
    print(f"✅ Documento salvo em: {output_path}")
    
    # Também salvar na pasta docs
    docs_path = os.path.join(
        os.path.dirname(os.path.dirname(__file__)),
        'docs',
        'DOCUMENTACAO_OFICIAL_SISPROCON.docx'
    )
    os.makedirs(os.path.dirname(docs_path), exist_ok=True)
    doc.save(docs_path)
    print(f"✅ Cópia salva em: {docs_path}")


if __name__ == '__main__':
    main()
