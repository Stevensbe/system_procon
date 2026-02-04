from rest_framework import viewsets, status
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated, IsAdminUser
from rest_framework.parsers import MultiPartParser, FormParser
from django.conf import settings
from django.core.mail import EmailMessage
from django.core.files.base import ContentFile
from django.db.models import Q
from django.contrib.contenttypes.models import ContentType
from django.utils import timezone
from django.http import HttpResponse
from datetime import timedelta
from io import BytesIO
from pathlib import Path
from docx import Document
import shutil
import subprocess
import tempfile
import logging

from .models import (
    EvidenciaFotografica,
    AssinaturaDigital,
    NotificacaoEletronica,
    Processo,
    DocumentoProcesso,
    ControlePrazos,
    ConfiguracaoFiscalizacao,
    AutoBanco,
    AutoPosto,
    AutoSupermercado,
    AutoDiversos,
    AutoInfracao,
)
from ppa.models import ProcedimentoPreAdministrativo, AnexoPPA, MovimentacaoPPA
from .serializers import (
    EvidenciaFotograficaSerializer,
    AssinaturaDigitalSerializer,
    NotificacaoEletronicaSerializer,
    ControlePrazosSerializer,
    ConfiguracaoFiscalizacaoSerializer,
    DashboardFiscalizacaoAvancadoSerializer,
)

logger = logging.getLogger(__name__)


def _formatar_data_extenso(data):
    meses = [
        'janeiro', 'fevereiro', 'março', 'abril', 'maio', 'junho',
        'julho', 'agosto', 'setembro', 'outubro', 'novembro', 'dezembro'
    ]
    mes = meses[data.month - 1]
    return f"{data.day:02d} de {mes} de {data.year}"


def _formatar_numero_ppa_exibicao(numero):
    if not numero:
        return ''
    texto = str(numero).strip()
    texto = texto.replace('PPA:', '').replace('PPA-', '').replace('PPA', '').strip()
    texto = texto.replace(' ', '')
    if '/' in texto:
        seq, ano = texto.split('/', 1)
        if seq.isdigit() and ano.isdigit():
            return f"{int(seq):03d}/{ano}"
    return texto


def _obter_auto_notificacao(notificacao):
    return (
        notificacao.auto
        or notificacao.auto_posto
        or notificacao.auto_supermercado
        or notificacao.auto_diversos
        or notificacao.auto_infracao
    )


def _normalizar_numero_ppa(valor):
    if not valor:
        return ''
    texto = str(valor).strip().upper()
    texto = texto.replace('PPA:', '').replace('PPA', '').strip()
    return texto.replace(' ', '')


def _buscar_ppa_por_numero(numero):
    if not numero:
        return None
    candidatos = []
    bruto = str(numero).strip()
    normalizado = _normalizar_numero_ppa(bruto)
    for valor in [bruto, normalizado]:
        if valor and valor not in candidatos:
            candidatos.append(valor)

    if normalizado and '/' in normalizado:
        seq, ano = normalizado.split('/', 1)
        if seq.isdigit() and ano.isdigit():
            candidatos.append(f"PPA-{int(seq):05d}/{ano}")
            candidatos.append(f"{int(seq):03d}/{ano}")

    for candidato in candidatos:
        ppa = ProcedimentoPreAdministrativo.objects.filter(numero__iexact=candidato).first()
        if ppa:
            return ppa

    return None


def _set_paragraph_text(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ''
    else:
        paragraph.add_run(text)


def _gerar_documento_notificacao(notificacao):
    template_path = Path(settings.BASE_DIR) / 'fiscalizacao' / 'templates' / 'docs' / 'NotificacaoFiscalizacao.docx'
    doc = Document(template_path)

    auto = _obter_auto_notificacao(notificacao)
    data_envio = notificacao.data_envio or timezone.now()
    numero_notificacao = notificacao.numero or f"{notificacao.id:03d}/{data_envio.year}"
    numero_documento = f"{numero_notificacao}-DFISC/PROCON-AM"
    ppa = notificacao.ppa or _buscar_ppa_por_numero(notificacao.assunto)
    ppa_numero = _formatar_numero_ppa_exibicao(ppa.numero) if ppa else ''
    ppa_linha = f"PPA: {ppa_numero}" if ppa_numero else "PPA: -"

    cidade = (getattr(auto, 'municipio', '') or 'Manaus').title()
    data_extenso = _formatar_data_extenso(data_envio.date())
    destinatario_nome = (notificacao.destinatario_nome or getattr(auto, 'razao_social', '') or '').strip()
    cnpj = (notificacao.destinatario_cpf_cnpj or getattr(auto, 'cnpj', '') or '').strip()
    representante_legal = (notificacao.representante_legal or getattr(auto, 'responsavel_nome', '') or '').strip()

    endereco = (getattr(auto, 'endereco', '') or '').strip()
    cep = (getattr(auto, 'cep', '') or '').strip()
    estado = (getattr(auto, 'estado', '') or '').strip()
    linha_endereco = endereco.upper() if endereco else ''
    linha_cidade = " - ".join(filter(None, [(f"{cidade}-{estado}" if estado else cidade).upper(), f"CEP: {cep}" if cep else '']))

    mensagem = (notificacao.mensagem or '').strip()

    substituir_endereco = 0
    aguardando_destinatario = False
    for paragraph in doc.paragraphs:
        texto = paragraph.text.strip()
        if not texto:
            continue

        texto_upper = texto.upper()
        texto_lower = texto.lower()

        if texto_upper.startswith('NOTIFICA'):
            _set_paragraph_text(paragraph, f"NOTIFICAÇÃO N° {numero_documento}")
        elif texto_upper.startswith('PPA'):
            _set_paragraph_text(paragraph, ppa_linha)
        elif texto_lower.startswith('manaus,'):
            _set_paragraph_text(paragraph, f"{cidade}, {data_extenso}")
        elif texto == 'À' or texto == 'A':
            aguardando_destinatario = True
        elif aguardando_destinatario:
            if destinatario_nome:
                _set_paragraph_text(paragraph, destinatario_nome.upper())
            aguardando_destinatario = False
        elif texto_upper.startswith('CNPJ'):
            if cnpj:
                _set_paragraph_text(paragraph, f"CNPJ: {cnpj}")
        elif texto_upper.startswith('NA PESSOA'):
            if representante_legal:
                _set_paragraph_text(
                    paragraph,
                    f"NA PESSOA DE SEU REPRESENTANTE LEGAL: {representante_legal}"
                )
            substituir_endereco = 2
        elif substituir_endereco:
            if substituir_endereco == 2 and linha_endereco:
                _set_paragraph_text(paragraph, linha_endereco)
            elif substituir_endereco == 1 and linha_cidade:
                _set_paragraph_text(paragraph, linha_cidade)
            substituir_endereco -= 1
        elif ('denúncia' in texto_lower or 'denuncia' in texto_lower) and mensagem:
            _set_paragraph_text(paragraph, mensagem)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue(), numero_notificacao


def _converter_docx_para_pdf(docx_bytes, nome_base):
    """Converte um DOCX em PDF e retorna os bytes, ou None se falhar."""
    if not docx_bytes:
        return None
    safe_name = (nome_base or 'Notificacao').replace('/', '_')
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = Path(tmpdir) / f"{safe_name}.docx"
        pdf_path = Path(tmpdir) / f"{safe_name}.pdf"
        docx_path.write_bytes(docx_bytes)

        try:
            from docx2pdf import convert

            convert(str(docx_path), str(tmpdir))
            if pdf_path.exists():
                return pdf_path.read_bytes()
        except Exception:
            pass

        converter = shutil.which('soffice') or shutil.which('libreoffice')
        if converter:
            try:
                subprocess.run(
                    [converter, '--headless', '--convert-to', 'pdf', '--outdir', tmpdir, str(docx_path)],
                    check=True,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                )
                if pdf_path.exists():
                    return pdf_path.read_bytes()
            except Exception:
                return None
    return None


# --- VIEWS PARA EVIDÊNCIAS FOTOGRÁFICAS ---
class EvidenciaFotograficaViewSet(viewsets.ModelViewSet):
    """
    ViewSet para gerenciar evidências fotográficas de fiscalização
    """
    queryset = EvidenciaFotografica.objects.all()
    serializer_class = EvidenciaFotograficaSerializer
    permission_classes = [IsAuthenticated]
    parser_classes = [MultiPartParser, FormParser]
    
    def get_queryset(self):
        queryset = EvidenciaFotografica.objects.all()
        
        # Filtros
        auto_id = self.request.query_params.get('auto_id')
        tipo_evidencia = self.request.query_params.get('tipo_evidencia')
        fiscal_id = self.request.query_params.get('fiscal_id')
        
        if auto_id:
            queryset = queryset.filter(
                Q(auto_id=auto_id) | 
                Q(auto_posto_id=auto_id) | 
                Q(auto_supermercado_id=auto_id) | 
                Q(auto_diversos_id=auto_id) |
                Q(auto_infracao_id=auto_id)
            )
        
        if tipo_evidencia:
            queryset = queryset.filter(tipo_evidencia=tipo_evidencia)
        
        if fiscal_id:
            queryset = queryset.filter(fiscal_upload_id=fiscal_id)
        
        return queryset.select_related('fiscal_upload')
    
    def perform_create(self, serializer):
        serializer.save(fiscal_upload=self.request.user)
    
    @action(detail=False, methods=['get'])
    def por_auto(self, request):
        """Lista evidências por auto de infração"""
        auto_id = request.query_params.get('auto_id')
        if not auto_id:
            return Response({'error': 'auto_id é obrigatório'}, status=400)
        
        evidencias = self.get_queryset().filter(
            Q(auto_id=auto_id) | 
            Q(auto_posto_id=auto_id) | 
            Q(auto_supermercado_id=auto_id) | 
            Q(auto_diversos_id=auto_id)
        )
        
        serializer = self.get_serializer(evidencias, many=True)
        return Response(serializer.data)


# --- VIEWS PARA ASSINATURA DIGITAL ---
class AssinaturaDigitalViewSet(viewsets.ModelViewSet):
    """
    ViewSet para gerenciar assinaturas digitais
    """
    queryset = AssinaturaDigital.objects.all()
    serializer_class = AssinaturaDigitalSerializer
    permission_classes = [IsAuthenticated]
    parser_classes = [MultiPartParser, FormParser]
    
    def get_queryset(self):
        queryset = AssinaturaDigital.objects.all()
        
        # Filtros
        auto_id = self.request.query_params.get('auto_id')
        tipo_assinatura = self.request.query_params.get('tipo_assinatura')
        status = self.request.query_params.get('status')
        
        if auto_id:
            queryset = queryset.filter(
                Q(auto_id=auto_id) | 
                Q(auto_posto_id=auto_id) | 
                Q(auto_supermercado_id=auto_id) | 
                Q(auto_diversos_id=auto_id)
            )
        
        if tipo_assinatura:
            queryset = queryset.filter(tipo_assinatura=tipo_assinatura)
        
        if status:
            queryset = queryset.filter(status=status)
        
        return queryset
    
    @action(detail=True, methods=['post'])
    def assinar(self, request, pk=None):
        """Realiza a assinatura digital"""
        assinatura = self.get_object()
        
        # Verifica se não está expirada
        if assinatura.is_expired():
            return Response({'error': 'Assinatura expirada'}, status=400)
        
        # Verifica se já não foi assinada
        if assinatura.status == 'assinado':
            return Response({'error': 'Assinatura já realizada'}, status=400)
        
        # Processa a assinatura
        assinatura.status = 'assinado'
        assinatura.data_assinatura = timezone.now()
        assinatura.ip_assinatura = self.get_client_ip(request)
        assinatura.user_agent = request.META.get('HTTP_USER_AGENT', '')
        
        # Gera hash da assinatura (simulado)
        assinatura.hash_assinatura = f"hash_{assinatura.id}_{int(timezone.now().timestamp())}"
        
        assinatura.save()
        
        serializer = self.get_serializer(assinatura)
        return Response(serializer.data)
    
    @action(detail=False, methods=['get'])
    def pendentes(self, request):
        """Lista assinaturas pendentes"""
        assinaturas = self.get_queryset().filter(status='pendente')
        serializer = self.get_serializer(assinaturas, many=True)
        return Response(serializer.data)
    
    @action(detail=False, methods=['get'])
    def vencidas(self, request):
        """Lista assinaturas vencidas"""
        assinaturas = [a for a in self.get_queryset() if a.is_expired()]
        serializer = self.get_serializer(assinaturas, many=True)
        return Response(serializer.data)
    
    def get_client_ip(self, request):
        x_forwarded_for = request.META.get('HTTP_X_FORWARDED_FOR')
        if x_forwarded_for:
            ip = x_forwarded_for.split(',')[0]
        else:
            ip = request.META.get('REMOTE_ADDR')
        return ip


# --- VIEWS PARA NOTIFICAÇÃO ELETRÔNICA ---
class NotificacaoEletronicaViewSet(viewsets.ModelViewSet):
    """
    ViewSet para gerenciar notificações eletrônicas
    """
    queryset = NotificacaoEletronica.objects.all()
    serializer_class = NotificacaoEletronicaSerializer
    permission_classes = [IsAuthenticated]
    
    def get_queryset(self):
        queryset = NotificacaoEletronica.objects.all()
        
        # Filtros
        auto_id = self.request.query_params.get('auto_id')
        tipo_notificacao = self.request.query_params.get('tipo_notificacao')
        status = self.request.query_params.get('status')
        
        if auto_id:
            queryset = queryset.filter(
                Q(auto_id=auto_id) | 
                Q(auto_posto_id=auto_id) | 
                Q(auto_supermercado_id=auto_id) | 
                Q(auto_diversos_id=auto_id)
            )
        
        if tipo_notificacao:
            queryset = queryset.filter(tipo_notificacao=tipo_notificacao)
        
        if status:
            queryset = queryset.filter(status=status)
        
        return queryset

    def _obter_processo_para_auto(self, auto, notificacao):
        if not auto:
            return None

        if isinstance(auto, AutoInfracao):
            processo = Processo.objects.filter(auto_infracao=auto).first()
            if processo:
                return processo
            return Processo.objects.create(
                auto_infracao=auto,
                autuado=getattr(auto, 'razao_social', ''),
                cnpj=getattr(auto, 'cnpj', ''),
                status='aguardando_defesa',
                prioridade='normal',
                fiscal_responsavel=getattr(auto, 'fiscal_nome', '') or '',
                observacoes=f"Processo criado a partir da notificacao {notificacao.id} do auto de infracao {auto.numero}.",
            )

        content_type = ContentType.objects.get_for_model(auto.__class__)
        processo = Processo.objects.filter(
            auto_constatacao_content_type=content_type,
            auto_constatacao_id=auto.id,
        ).first()
        if processo:
            return processo

        fiscal_responsavel = ''
        if getattr(auto, 'fiscal_nome_1', None):
            fiscal_responsavel = auto.fiscal_nome_1
        elif getattr(auto, 'fiscal_nome', None):
            fiscal_responsavel = auto.fiscal_nome

        return Processo.objects.create(
            auto_constatacao_content_type=content_type,
            auto_constatacao_id=auto.id,
            autuado=getattr(auto, 'razao_social', ''),
            cnpj=getattr(auto, 'cnpj', ''),
            status='aguardando_auto_infracao',
            prioridade='normal',
            fiscal_responsavel=fiscal_responsavel,
            observacoes=f"Processo criado a partir da notificacao {notificacao.id} do auto {auto.numero}.",
        )

    def _mapear_sigla_ppa(self, auto):
        if isinstance(auto, AutoBanco):
            return 'BANCO'
        if isinstance(auto, AutoPosto):
            return 'POSTO'
        if isinstance(auto, AutoSupermercado):
            return 'SUPERMERCADO'
        if isinstance(auto, AutoDiversos):
            return 'DIVERSOS'
        return 'OUTROS'

    def _registrar_auto_ppa(self, ppa, auto, tipo_documento, tipo_movimentacao, descricao):
        if not ppa or not auto:
            return None
        content_type = ContentType.objects.get_for_model(auto)
        numero_documento = f"{tipo_documento} {getattr(auto, 'numero', '')}".strip()
        existente = AnexoPPA.objects.filter(
            ppa=ppa,
            tipo_documento=tipo_documento,
            content_type=content_type,
            object_id=auto.id,
        ).first()
        if existente:
            return existente
        anexo = AnexoPPA.objects.create(
            ppa=ppa,
            tipo_documento=tipo_documento,
            descricao=descricao,
            numero_documento=numero_documento,
            content_type=content_type,
            object_id=auto.id,
            anexado_por=self.request.user,
        )
        MovimentacaoPPA.objects.create(
            ppa=ppa,
            tipo_movimentacao=tipo_movimentacao,
            atendimento=descricao,
            usuario=self.request.user,
        )
        return anexo

    def _criar_ppa_para_notificacao(self, notificacao, auto):
        if not auto:
            return None
        interessado = (getattr(auto, 'razao_social', '') or '').strip()
        numero_auto = getattr(auto, 'numero', '') or ''
        assunto = f"Auto {numero_auto} - {interessado}".strip(' -') if numero_auto else interessado
        ppa = ProcedimentoPreAdministrativo.objects.create(
            sigla=self._mapear_sigla_ppa(auto),
            assunto=assunto or "Notificacao Fiscalizacao",
            interessado=interessado or "Nao informado",
            cnpj_interessado=getattr(auto, 'cnpj', '') or '',
            endereco_interessado=getattr(auto, 'endereco', '') or '',
            analista_responsavel=self.request.user,
            supervisor=None,
            status='notificado',
            criado_por=self.request.user,
        )
        MovimentacaoPPA.objects.create(
            ppa=ppa,
            tipo_movimentacao='criacao',
            atendimento=f"PPA criado a partir da notificacao {notificacao.id}.",
            usuario=self.request.user,
        )
        if isinstance(auto, AutoInfracao):
            self._registrar_auto_ppa(
                ppa,
                auto,
                'AI',
                'anexo_ai',
                f"Auto de Infracao {numero_auto} anexado",
            )
        else:
            self._registrar_auto_ppa(
                ppa,
                auto,
                'AC',
                'anexo_ac',
                f"Auto de Constatacao {numero_auto} anexado",
            )
        auto_infracao = notificacao.auto_infracao
        if auto_infracao and not isinstance(auto, AutoInfracao):
            self._registrar_auto_ppa(
                ppa,
                auto_infracao,
                'AI',
                'anexo_ai',
                f"Auto de Infracao {auto_infracao.numero} anexado",
            )
        return ppa

    def _obter_ppa_para_notificacao(self, notificacao, create_if_missing=False):
        if notificacao.ppa_id:
            return notificacao.ppa
        ppa_encontrado = _buscar_ppa_por_numero(notificacao.assunto)
        if ppa_encontrado:
            return ppa_encontrado
        if create_if_missing:
            auto = _obter_auto_notificacao(notificacao)
            return self._criar_ppa_para_notificacao(notificacao, auto)
        return None

    def perform_update(self, serializer):
        notificacao = serializer.save()
        try:
            self._gerar_e_registrar_documentos(
                notificacao,
                force_update=True,
                create_ppa=False,
            )
        except Exception as exc:
            logger.exception("Falha ao atualizar documentos da notificacao %s: %s", notificacao.id, exc)

    def _registrar_documento_processo(self, notificacao, processo, anexo_bytes, numero_notificacao, force_update=False):
        if not processo or not anexo_bytes:
            return None

        anexos = list(notificacao.anexos or [])
        existente = next(
            (
                item for item in anexos
                if isinstance(item, dict) and item.get('numero_notificacao') == numero_notificacao
            ),
            None
        )
        if existente and not force_update:
            return None

        documento = None
        if existente and existente.get('documento_id'):
            documento = DocumentoProcesso.objects.filter(id=existente['documento_id']).first()

        filename = f"Notificacao_{numero_notificacao.replace('/', '_')}.docx"
        if documento:
            documento.titulo = f"Notificacao Fiscalizacao {numero_notificacao}"
            documento.descricao = "Notificacao da fiscalizacao."
            documento.usuario_upload = getattr(self.request.user, 'username', '')
            documento.arquivo.save(filename, ContentFile(anexo_bytes), save=True)
        else:
            documento = DocumentoProcesso(
                processo=processo,
                tipo='outros',
                titulo=f"Notificacao Fiscalizacao {numero_notificacao}",
                descricao="Notificacao da fiscalizacao.",
                usuario_upload=getattr(self.request.user, 'username', ''),
            )
            documento.arquivo.save(filename, ContentFile(anexo_bytes), save=True)

        if existente:
            existente.update({
                'tipo': 'notificacao',
                'documento_id': documento.id,
                'arquivo': documento.arquivo.name,
            })
        else:
            anexos.append({
                'tipo': 'notificacao',
                'numero_notificacao': numero_notificacao,
                'documento_id': documento.id,
                'arquivo': documento.arquivo.name,
            })
        notificacao.anexos = anexos
        notificacao.save(update_fields=['anexos'])
        return documento

    def _registrar_documento_processo_pdf(self, notificacao, processo, anexo_bytes, numero_notificacao, force_update=False):
        if not processo or not anexo_bytes:
            return None

        anexos = list(notificacao.anexos or [])
        existente = next(
            (
                item for item in anexos
                if isinstance(item, dict)
                and item.get('tipo') == 'notificacao_pdf'
                and item.get('numero_notificacao') == numero_notificacao
            ),
            None
        )
        if existente and not force_update:
            return None

        documento = None
        if existente and existente.get('documento_id'):
            documento = DocumentoProcesso.objects.filter(id=existente['documento_id']).first()

        filename = f"Notificacao_{numero_notificacao.replace('/', '_')}.pdf"
        if documento:
            documento.titulo = f"Notificacao Fiscalizacao {numero_notificacao} (PDF)"
            documento.descricao = "Notificacao da fiscalizacao (PDF)."
            documento.usuario_upload = getattr(self.request.user, 'username', '')
            documento.arquivo.save(filename, ContentFile(anexo_bytes), save=True)
        else:
            documento = DocumentoProcesso(
                processo=processo,
                tipo='outros',
                titulo=f"Notificacao Fiscalizacao {numero_notificacao} (PDF)",
                descricao="Notificacao da fiscalizacao (PDF).",
                usuario_upload=getattr(self.request.user, 'username', ''),
            )
            documento.arquivo.save(filename, ContentFile(anexo_bytes), save=True)

        if existente:
            existente.update({
                'tipo': 'notificacao_pdf',
                'documento_id': documento.id,
                'arquivo': documento.arquivo.name,
            })
        else:
            anexos.append({
                'tipo': 'notificacao_pdf',
                'numero_notificacao': numero_notificacao,
                'documento_id': documento.id,
                'arquivo': documento.arquivo.name,
            })
        notificacao.anexos = anexos
        notificacao.save(update_fields=['anexos'])
        return documento

    def _gerar_docx_auto_constatacao(self, auto):
        try:
            from .views.utils_views import (
                gerar_documento_banco,
                gerar_documento_posto,
                gerar_documento_supermercado,
                gerar_documento_diversos,
            )

            request = getattr(self, 'request', None)
            if request is None:
                from django.test import RequestFactory
                request = RequestFactory().get('/')

            if isinstance(auto, AutoBanco):
                resultado = gerar_documento_banco(request, auto.id)
            elif isinstance(auto, AutoPosto):
                resultado = gerar_documento_posto(request, auto.id)
            elif isinstance(auto, AutoSupermercado):
                resultado = gerar_documento_supermercado(request, auto.id)
            elif isinstance(auto, AutoDiversos):
                resultado = gerar_documento_diversos(request, auto.id)
            else:
                return None

            if isinstance(resultado, Response) or getattr(resultado, 'status_code', 200) >= 400:
                return None
            if isinstance(resultado, HttpResponse):
                return resultado.content
        except Exception as exc:
            logger.exception("Falha ao gerar DOCX do auto %s: %s", getattr(auto, 'numero', ''), exc)
        return None

    def _gerar_docx_auto_infracao(self, auto):
        try:
            from .views.infracao_views import gerar_documento_infracao_docx

            request = getattr(self, 'request', None)
            if request is None:
                from django.test import RequestFactory
                request = RequestFactory().get('/')

            resultado = gerar_documento_infracao_docx(request, auto.id)
            if isinstance(resultado, Response) or getattr(resultado, 'status_code', 200) >= 400:
                return None
            if isinstance(resultado, HttpResponse):
                return resultado.content
        except Exception as exc:
            logger.exception("Falha ao gerar DOCX do auto de infracao %s: %s", getattr(auto, 'numero', ''), exc)
        return None

    def _registrar_documento_auto(self, processo, auto, force_update=False):
        if not processo or not auto:
            return None

        if isinstance(auto, AutoInfracao):
            titulo = f"Auto de Infracao {auto.numero}"
            docx_bytes = self._gerar_docx_auto_infracao(auto)
        else:
            titulo = f"Auto de Constatacao {auto.numero}"
            docx_bytes = self._gerar_docx_auto_constatacao(auto)

        if not docx_bytes:
            return None

        documento = DocumentoProcesso.objects.filter(
            processo=processo,
            titulo=titulo,
        ).first()

        if documento and not force_update:
            return documento

        safe_numero = (auto.numero or '').replace('/', '_')
        safe_titulo = titulo.replace(' ', '_').replace('/', '_')
        filename = f"{safe_titulo}_{safe_numero}.docx"
        usuario_upload = ''
        if self.request.user and getattr(self.request.user, 'is_authenticated', False):
            usuario_upload = self.request.user.get_full_name() or self.request.user.username

        if documento:
            documento.tipo = 'outros'
            documento.descricao = titulo
            documento.usuario_upload = usuario_upload
            documento.arquivo.save(filename, ContentFile(docx_bytes), save=True)
            return documento

        documento = DocumentoProcesso(
            processo=processo,
            tipo='outros',
            titulo=titulo,
            descricao=titulo,
            usuario_upload=usuario_upload,
        )
        documento.arquivo.save(filename, ContentFile(docx_bytes), save=True)
        return documento

    def _registrar_anexo_ppa(self, notificacao, ppa, anexo_bytes, numero_notificacao, force_update=False):
        if not ppa or not anexo_bytes:
            return None

        numero_documento = f"NOT {numero_notificacao}"
        existente = AnexoPPA.objects.filter(
            ppa=ppa,
            tipo_documento='NOT',
            numero_documento=numero_documento,
        ).first()
        if existente and not force_update:
            return None

        filename = f"NOT_{numero_notificacao.replace('/', '_')}.docx"
        if existente:
            anexo = existente
            anexo.descricao = 'Notificacao da fiscalizacao'
            anexo.anexado_por = self.request.user
            anexo.nome_arquivo_original = filename
            anexo.content_type = ContentType.objects.get_for_model(NotificacaoEletronica)
            anexo.object_id = notificacao.id
            anexo.arquivo.save(filename, ContentFile(anexo_bytes), save=True)
        else:
            anexo = AnexoPPA(
                ppa=ppa,
                tipo_documento='NOT',
                descricao='Notificacao da fiscalizacao',
                numero_documento=numero_documento,
                anexado_por=self.request.user,
                nome_arquivo_original=filename,
                content_type=ContentType.objects.get_for_model(NotificacaoEletronica),
                object_id=notificacao.id,
            )
            anexo.arquivo.save(filename, ContentFile(anexo_bytes), save=True)

        if not existente:
            MovimentacaoPPA.objects.create(
                ppa=ppa,
                tipo_movimentacao='anexo_notificacao',
                atendimento=f"Notificacao {numero_documento} anexada",
                usuario=self.request.user,
            )

        anexos = list(notificacao.anexos or [])
        if existente:
            atualizado = False
            for item in anexos:
                if isinstance(item, dict) and item.get('anexo_ppa_id') == anexo.id:
                    item.update({'arquivo': anexo.arquivo.name})
                    atualizado = True
                    break
            if not atualizado:
                anexos.append({
                    'tipo': 'ppa_notificacao',
                    'numero_notificacao': numero_notificacao,
                    'anexo_ppa_id': anexo.id,
                    'arquivo': anexo.arquivo.name,
                })
        else:
            anexos.append({
                'tipo': 'ppa_notificacao',
                'numero_notificacao': numero_notificacao,
                'anexo_ppa_id': anexo.id,
                'arquivo': anexo.arquivo.name,
            })
        notificacao.anexos = anexos
        notificacao.save(update_fields=['anexos'])
        return anexo

    def _gerar_e_registrar_documentos(self, notificacao, force_update=False, create_ppa=False):
        if not notificacao.numero:
            notificacao.save(update_fields=['numero'])

        anexo_bytes, numero_notificacao = _gerar_documento_notificacao(notificacao)
        pdf_bytes = _converter_docx_para_pdf(anexo_bytes, f"Notificacao_{numero_notificacao}")

        auto = _obter_auto_notificacao(notificacao)
        processo = notificacao.processo or self._obter_processo_para_auto(auto, notificacao)
        if processo and notificacao.processo_id != processo.id:
            notificacao.processo = processo
            notificacao.save(update_fields=['processo'])

        self._registrar_documento_auto(
            processo,
            auto,
            force_update=force_update,
        )

        auto_infracao = notificacao.auto_infracao
        if auto_infracao and not isinstance(auto, AutoInfracao):
            self._registrar_documento_auto(
                processo,
                auto_infracao,
                force_update=force_update,
            )

        self._registrar_documento_processo(
            notificacao,
            processo,
            anexo_bytes,
            numero_notificacao,
            force_update=force_update,
        )
        if pdf_bytes:
            self._registrar_documento_processo_pdf(
                notificacao,
                processo,
                pdf_bytes,
                numero_notificacao,
                force_update=force_update,
            )

        ppa = notificacao.ppa or self._obter_ppa_para_notificacao(
            notificacao,
            create_if_missing=create_ppa,
        )
        if ppa and notificacao.ppa_id != ppa.id:
            notificacao.ppa = ppa
            notificacao.save(update_fields=['ppa'])
        self._registrar_anexo_ppa(
            notificacao,
            ppa,
            anexo_bytes,
            numero_notificacao,
            force_update=force_update,
        )

        return anexo_bytes, numero_notificacao

    def perform_create(self, serializer):
        notificacao = serializer.save()
        auto = _obter_auto_notificacao(notificacao)
        update_fields = []
        if not notificacao.assunto:
            if auto and getattr(auto, 'numero', None):
                notificacao.assunto = f"Auto {auto.numero}"
            else:
                notificacao.assunto = "Notificacao Fiscalizacao"
            update_fields.append('assunto')
        if not notificacao.destinatario_nome and auto and getattr(auto, 'razao_social', None):
            notificacao.destinatario_nome = auto.razao_social
            update_fields.append('destinatario_nome')
        if not notificacao.destinatario_cpf_cnpj and auto and getattr(auto, 'cnpj', None):
            notificacao.destinatario_cpf_cnpj = auto.cnpj
            update_fields.append('destinatario_cpf_cnpj')
        if update_fields:
            notificacao.save(update_fields=update_fields)

        try:
            self._gerar_e_registrar_documentos(
                notificacao,
                create_ppa=True,
            )
            notificacao.logs_envio.append({
                'data': timezone.now().isoformat(),
                'acao': 'gerada',
                'status': 'rascunho'
            })
            notificacao.save(update_fields=['logs_envio'])
        except Exception as exc:
            logger.exception("Falha ao gerar documentos da notificacao %s: %s", notificacao.id, exc)
    
    @action(detail=True, methods=['post'])
    def enviar(self, request, pk=None):
        """Envia a notificação eletrônica"""
        notificacao = self.get_object()
        
        try:
            resultado = self._processar_envio(notificacao)
            if not resultado['ok']:
                return Response({'error': resultado['erro']}, status=400)
            serializer = self.get_serializer(notificacao)
            return Response(serializer.data)
        except Exception as e:
            self._registrar_erro_envio(notificacao, str(e))
            return Response({'error': f'Erro ao enviar: {str(e)}'}, status=500)
    
    def _processar_envio(self, notificacao):
        if notificacao.status in ['enviada', 'entregue', 'lida']:
            return {'ok': False, 'erro': 'Notificacao ja enviada'}

        if notificacao.tentativas_envio >= notificacao.max_tentativas:
            notificacao.status = 'erro'
            notificacao.save(update_fields=['status'])
            return {'ok': False, 'erro': 'Maximo de tentativas excedido'}

        if not notificacao.destinatario_email:
            self._registrar_erro_envio(notificacao, 'Email do destinatario nao informado')
            return {'ok': False, 'erro': 'Email do destinatario nao informado'}

        anexo_bytes, numero_notificacao = self._gerar_e_registrar_documentos(
            notificacao,
            force_update=True,
            create_ppa=True,
        )

        assunto = notificacao.assunto or f"Notificacao {numero_notificacao}"
        portal_url = getattr(settings, 'PORTAL_CIDADAO_URL', '')
        corpo = (notificacao.mensagem or '').strip() or "Segue em anexo a notificacao emitida pela Fiscalizacao."
        if portal_url:
            corpo += f"\n\nPara apresentar peticao, acesse: {portal_url}"
        else:
            corpo += "\n\nCaso deseje apresentar peticao, utilize o Portal do Cidadao."
        self._enviar_email(
            destinatario=notificacao.destinatario_email,
            assunto=assunto,
            corpo=corpo,
            anexo_bytes=anexo_bytes,
            numero_notificacao=numero_notificacao,
        )

        if notificacao.tipo_notificacao == 'auto_infracao' and notificacao.auto_infracao:
            auto_infracao = notificacao.auto_infracao
            data_notificacao = auto_infracao.data_notificacao or timezone.now().date()
            if auto_infracao.status != 'notificado':
                auto_infracao.status = 'notificado'
            if not auto_infracao.data_notificacao:
                auto_infracao.data_notificacao = data_notificacao
            auto_infracao.save(update_fields=['status', 'data_notificacao'])

            processo = notificacao.processo or Processo.objects.filter(auto_infracao=auto_infracao).first()
            if processo:
                processo.data_notificacao = data_notificacao
                processo.calcular_prazos()
                processo.atualizar_status(
                    'aguardando_defesa',
                    f"Notificação de multa enviada (AI {auto_infracao.numero})."
                )
                if notificacao.processo_id != processo.id:
                    notificacao.processo = processo
                    notificacao.save(update_fields=['processo'])

        notificacao.status = 'enviada'
        notificacao.data_envio = timezone.now()
        notificacao.tentativas_envio += 1
        notificacao.logs_envio.append({
            'data': timezone.now().isoformat(),
            'acao': 'enviada',
            'status': 'sucesso',
            'numero_notificacao': numero_notificacao,
        })
        notificacao.save(update_fields=['status', 'data_envio', 'tentativas_envio', 'logs_envio'])

        return {'ok': True}

    def _registrar_erro_envio(self, notificacao, erro):
        notificacao.tentativas_envio += 1
        notificacao.proxima_tentativa = timezone.now() + timedelta(hours=1)
        notificacao.logs_envio.append({
            'data': timezone.now().isoformat(),
            'acao': 'erro',
            'erro': erro
        })
        notificacao.save(update_fields=['tentativas_envio', 'proxima_tentativa', 'logs_envio'])

    def _enviar_email(self, destinatario, assunto, corpo, anexo_bytes, numero_notificacao):
        email = EmailMessage(
            subject=assunto,
            body=corpo,
            from_email=getattr(settings, 'DEFAULT_FROM_EMAIL', None),
            to=[destinatario],
        )
        email.attach(
            filename=f"Notificacao_{numero_notificacao.replace('/', '_')}.docx",
            content=anexo_bytes,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        email.send(fail_silently=False)
    
    @action(detail=False, methods=['get'])
    def pendentes(self, request):
        """Lista notificações pendentes"""
        notificacoes = self.get_queryset().filter(status='pendente')
        serializer = self.get_serializer(notificacoes, many=True)
        return Response(serializer.data)
    
    @action(detail=False, methods=['get'])
    def estatisticas(self, request):
        """Estatísticas de notificações"""
        queryset = self.get_queryset()
        
        stats = {
            'total': queryset.count(),
            'pendentes': queryset.filter(status='pendente').count(),
            'enviadas': queryset.filter(status='enviada').count(),
            'entregues': queryset.filter(status='entregue').count(),
            'lidas': queryset.filter(status='lida').count(),
            'erros': queryset.filter(status='erro').count(),
        }
        
        return Response(stats)


# --- VIEWS PARA CONTROLE DE PRAZOS ---
class ControlePrazosViewSet(viewsets.ModelViewSet):
    """
    ViewSet para gerenciar controle de prazos
    """
    queryset = ControlePrazos.objects.all()
    serializer_class = ControlePrazosSerializer
    permission_classes = [IsAuthenticated]
    
    def get_queryset(self):
        queryset = ControlePrazos.objects.all()
        
        # Filtros
        auto_id = self.request.query_params.get('auto_id')
        tipo_prazo = self.request.query_params.get('tipo_prazo')
        status = self.request.query_params.get('status')
        
        if auto_id:
            queryset = queryset.filter(
                Q(auto_id=auto_id) | 
                Q(auto_posto_id=auto_id) | 
                Q(auto_supermercado_id=auto_id) | 
                Q(auto_diversos_id=auto_id)
            )
        
        if tipo_prazo:
            queryset = queryset.filter(tipo_prazo=tipo_prazo)
        
        if status:
            queryset = queryset.filter(status=status)
        
        return queryset.select_related('responsavel')
    
    @action(detail=True, methods=['post'])
    def prorrogar(self, request, pk=None):
        """Prorroga um prazo"""
        prazo = self.get_object()
        nova_data = request.data.get('nova_data')
        
        if not nova_data:
            return Response({'error': 'nova_data é obrigatória'}, status=400)
        
        try:
            nova_data = timezone.datetime.fromisoformat(nova_data.replace('Z', '+00:00'))
        except ValueError:
            return Response({'error': 'Formato de data inválido'}, status=400)
        
        # Atualiza o prazo
        prazo.data_prorrogacao = nova_data
        prazo.status = 'prorrogado'
        prazo.historico_alteracoes.append({
            'data': timezone.now().isoformat(),
            'acao': 'prorrogado',
            'data_anterior': prazo.data_fim.isoformat(),
            'data_nova': nova_data.isoformat(),
            'usuario': request.user.username
        })
        prazo.save()
        
        serializer = self.get_serializer(prazo)
        return Response(serializer.data)
    
    @action(detail=False, methods=['get'])
    def vencendo(self, request):
        """Lista prazos vencendo (5 dias ou menos)"""
        prazos = [p for p in self.get_queryset() if p.is_vencendo()]
        serializer = self.get_serializer(prazos, many=True)
        return Response(serializer.data)
    
    @action(detail=False, methods=['get'])
    def vencidos(self, request):
        """Lista prazos vencidos"""
        prazos = [p for p in self.get_queryset() if p.is_vencido()]
        serializer = self.get_serializer(prazos, many=True)
        return Response(serializer.data)
    
    @action(detail=False, methods=['get'])
    def alertas(self, request):
        """Lista alertas de prazos"""
        alertas = []
        prazos = self.get_queryset()
        
        for prazo in prazos:
            dias_restantes = prazo.calcular_dias_restantes()
            if dias_restantes is not None:
                if dias_restantes <= 0:
                    alertas.append({
                        'tipo': 'vencido',
                        'prazo': prazo,
                        'dias': abs(dias_restantes)
                    })
                elif dias_restantes <= 5:
                    alertas.append({
                        'tipo': 'vencendo',
                        'prazo': prazo,
                        'dias': dias_restantes
                    })
        
        serializer = self.get_serializer([a['prazo'] for a in alertas], many=True)
        return Response({
            'alertas': alertas,
            'prazos': serializer.data
        })


# --- VIEWS PARA CONFIGURAÇÕES ---
class ConfiguracaoFiscalizacaoViewSet(viewsets.ModelViewSet):
    """
    ViewSet para gerenciar configurações de fiscalização
    """
    queryset = ConfiguracaoFiscalizacao.objects.all()
    serializer_class = ConfiguracaoFiscalizacaoSerializer
    permission_classes = [IsAuthenticated, IsAdminUser]
    
    def get_object(self):
        """Sempre retorna a configuração padrão"""
        return ConfiguracaoFiscalizacao.get_config()
    
    def list(self, request, *args, **kwargs):
        """Retorna a configuração atual"""
        config = self.get_object()
        serializer = self.get_serializer(config)
        return Response(serializer.data)


# --- VIEWS PARA DASHBOARD AVANÇADO ---
class DashboardFiscalizacaoAvancadoViewSet(viewsets.ViewSet):
    """
    ViewSet para dashboard avançado de fiscalização
    """
    permission_classes = [IsAuthenticated]
    
    @action(detail=False, methods=['get'])
    def estatisticas(self, request):
        """Estatísticas completas do dashboard"""
        # Estatísticas gerais
        total_autos = AutoBanco.objects.count() + AutoPosto.objects.count() + AutoSupermercado.objects.count() + AutoDiversos.objects.count()
        autos_este_mes = AutoBanco.objects.filter(data_fiscalizacao__month=timezone.now().month).count()
        autos_pendentes = AutoBanco.objects.filter(status='pendente').count()
        autos_vencidos = AutoBanco.objects.filter(status='vencido').count()
        
        # Evidências
        total_evidencias = EvidenciaFotografica.objects.count()
        evidencias_este_mes = EvidenciaFotografica.objects.filter(data_upload__month=timezone.now().month).count()
        
        # Assinaturas
        assinaturas_pendentes = AssinaturaDigital.objects.filter(status='pendente').count()
        assinaturas_vencidas = AssinaturaDigital.objects.filter(status='expirado').count()
        
        # Notificações
        notificacoes_pendentes = NotificacaoEletronica.objects.filter(status='pendente').count()
        notificacoes_enviadas = NotificacaoEletronica.objects.filter(status='enviada').count()
        notificacoes_entregues = NotificacaoEletronica.objects.filter(status='entregue').count()
        
        # Prazos
        prazos_vencendo = ControlePrazos.objects.filter(status='vencendo').count()
        prazos_vencidos = ControlePrazos.objects.filter(status='vencido').count()
        
        # Gráficos
        autos_por_tipo = {
            'Banco': AutoBanco.objects.count(),
            'Posto': AutoPosto.objects.count(),
            'Supermercado': AutoSupermercado.objects.count(),
            'Diversos': AutoDiversos.objects.count(),
        }
        
        # Alertas
        alertas_criticos = []
        alertas_importantes = []
        
        # Prazos vencidos
        prazos_vencidos_list = [p for p in ControlePrazos.objects.all() if p.is_vencido()]
        for prazo in prazos_vencidos_list:
            alertas_criticos.append({
                'tipo': 'prazo_vencido',
                'titulo': f'Prazo vencido: {prazo.descricao}',
                'descricao': f'Prazo vencido há {abs(prazo.calcular_dias_restantes())} dias',
                'data': prazo.data_fim
            })
        
        # Assinaturas vencidas
        assinaturas_vencidas_list = [a for a in AssinaturaDigital.objects.all() if a.is_expired()]
        for assinatura in assinaturas_vencidas_list:
            alertas_importantes.append({
                'tipo': 'assinatura_vencida',
                'titulo': f'Assinatura vencida: {assinatura.nome_assinante}',
                'descricao': f'Assinatura vencida em {assinatura.data_expiracao.strftime("%d/%m/%Y")}',
                'data': assinatura.data_expiracao
            })
        
        data = {
            'total_autos': total_autos,
            'autos_este_mes': autos_este_mes,
            'autos_pendentes': autos_pendentes,
            'autos_vencidos': autos_vencidos,
            'total_evidencias': total_evidencias,
            'evidencias_este_mes': evidencias_este_mes,
            'assinaturas_pendentes': assinaturas_pendentes,
            'assinaturas_vencidas': assinaturas_vencidas,
            'notificacoes_pendentes': notificacoes_pendentes,
            'notificacoes_enviadas': notificacoes_enviadas,
            'notificacoes_entregues': notificacoes_entregues,
            'prazos_vencendo': prazos_vencendo,
            'prazos_vencidos': prazos_vencidos,
            'autos_por_tipo': autos_por_tipo,
            'alertas_criticos': alertas_criticos,
            'alertas_importantes': alertas_importantes,
        }
        
        serializer = DashboardFiscalizacaoAvancadoSerializer(data)
        return Response(serializer.data)
