"""
Views para Autos de Infração - System Procon

Este módulo contém as views relacionadas aos Autos de Infração,
incluindo criação, listagem, atualização de status e relacionamentos com Autos de Constatação.
"""

from django.http import JsonResponse
from django.conf import settings
from pathlib import Path
from django.shortcuts import get_object_or_404
from django.db.models import Count, Q
from django.utils import timezone
from django.contrib.contenttypes.models import ContentType
from django.core.files.base import ContentFile
import logging
from rest_framework import generics, status
from rest_framework.decorators import api_view
from rest_framework.response import Response
from datetime import datetime, timedelta

from ..models import (
    AutoBanco,
    AutoPosto,
    AutoSupermercado,
    AutoDiversos,
    AutoInfracao,
    Processo,
    DocumentoProcesso,
    NotificacaoEletronica,
    HistoricoStatusInfracao,
    STATUS_INFRACAO_CHOICES,
)

logger = logging.getLogger(__name__)

from ..serializers import (
    AutoInfracaoSerializer,
    AutoInfracaoCreateSerializer,
    AutoInfracaoSimpleSerializer,
)


# ========================================
# VIEWS DE API - AUTO INFRAÇÃO
# ========================================

class AutoInfracaoListCreateAPIView(generics.ListCreateAPIView):
    queryset = AutoInfracao.objects.all()
    
    def get_serializer_class(self):
        if self.request.method == 'POST':
            return AutoInfracaoCreateSerializer
        return AutoInfracaoSerializer

    def perform_create(self, serializer):
        auto_tipo = serializer.validated_data.get('auto_tipo')
        auto_id = serializer.validated_data.get('auto_id')
        infracao = serializer.save()

        if auto_tipo and auto_id:
            auto_obj = _buscar_auto_constatacao(auto_tipo, auto_id)
            if auto_obj:
                _vincular_auto_infracao_ao_processo(infracao, auto_obj, self.request)
    
    def get_queryset(self):
        queryset = super().get_queryset()
        
        # Filtros específicos
        status_filter = self.request.query_params.get('status')
        numero = self.request.query_params.get('numero') or self.request.query_params.get('search')
        razao_social = self.request.query_params.get('razao_social')
        cnpj = self.request.query_params.get('cnpj')
        data_inicio = self.request.query_params.get('data_inicio')
        data_fim = self.request.query_params.get('data_fim')
        
        if status_filter:
            queryset = queryset.filter(status=status_filter)
        if numero:
            queryset = queryset.filter(numero__icontains=numero.strip())
        if razao_social:
            queryset = queryset.filter(razao_social__icontains=razao_social)
        if cnpj:
            queryset = queryset.filter(cnpj__icontains=cnpj)
        if data_inicio:
            queryset = queryset.filter(data_fiscalizacao__gte=data_inicio)
        if data_fim:
            queryset = queryset.filter(data_fiscalizacao__lte=data_fim)
        
        return queryset


class AutoInfracaoRetrieveUpdateDestroyAPIView(generics.RetrieveUpdateDestroyAPIView):
    queryset = AutoInfracao.objects.all()
    serializer_class = AutoInfracaoSerializer


@api_view(['GET'])
def gerar_documento_infracao_docx(request, pk):
    """
    Gera documento DOCX do Auto de Infração
    """
    try:
        auto = AutoInfracao.objects.get(pk=pk)
    except AutoInfracao.DoesNotExist:
        return Response({'error': 'Auto de Infração não encontrado'}, status=404)

    try:
        from django.http import HttpResponse
        docx_bytes = _gerar_docx_auto_infracao(auto)
        if not docx_bytes:
            return Response({'error': 'Erro ao gerar documento do Auto de Infração.'}, status=500)

        response = HttpResponse(
            docx_bytes,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f"attachment; filename=\"Auto_Infracao_{auto.numero.replace('/', '_')}\".docx"

        return response

    except Exception as e:
        return Response({'error': f'Erro ao gerar documento: {str(e)}'}, status=500)


def _gerar_docx_auto_infracao(auto):
    """Gera bytes do DOCX do Auto de Infração."""
    from docx import Document
    from io import BytesIO

    template_path = Path(settings.BASE_DIR) / 'fiscalizacao' / 'templates' / 'docs' / 'AutoInfracao.docx'
    if not template_path.exists():
        return None

    doc = Document(template_path)

    def set_cell_text(cell, text, bold=False):
        cell.text = ""
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(text)
        run.bold = bold

    if doc.paragraphs:
        title = doc.paragraphs[0]
        title.text = ""
        title_run = title.add_run(f"AUTO DE INFRAÇÃO Nº {auto.numero}")
        title_run.bold = True

    if not doc.tables:
        return None

    table = doc.tables[0]
    set_cell_text(table.cell(0, 0), f"RAZÃO SOCIAL: {auto.razao_social or ''}".strip(), bold=True)
    set_cell_text(table.cell(1, 0), f"NOME FANTASIA: {auto.nome_fantasia or ''}".strip(), bold=True)
    set_cell_text(table.cell(2, 0), f"ATIVIDADE: {auto.atividade or ''}".strip(), bold=True)
    set_cell_text(table.cell(3, 0), f"ENDEREÇO: {auto.endereco or ''}".strip(), bold=True)

    estado = auto.estado or "AM"
    estado_sigla = estado if len(estado) <= 2 else estado[:2]
    set_cell_text(table.cell(4, 0), f"CEP: {auto.cep or ''}".strip(), bold=True)
    set_cell_text(table.cell(4, 1), f"MUNICIPIO: {auto.municipio or ''}".strip(), bold=True)
    set_cell_text(table.cell(4, 3), f"ESTADO: {estado_sigla}".strip(), bold=True)

    set_cell_text(table.cell(5, 0), f"CNPJ: {auto.cnpj or ''}".strip(), bold=True)
    set_cell_text(table.cell(5, 2), f"TELEFONE: {auto.telefone or ''}".strip(), bold=True)

    intro_cell = table.cell(6, 0)
    intro_cell.text = ""
    intro_paragraph = intro_cell.paragraphs[0]
    intro_paragraph.add_run("     ")

    if auto.texto_origem:
        intro_paragraph.add_run(auto.texto_origem)
    else:
        origem_numero = auto.auto_constatacao_numero or "____/____"
        intro_paragraph.add_run("O presente Auto originou-se e a partir de ")
        if auto.notificacao_numero:
            origem_run = intro_paragraph.add_run(
                f"Auto de Constatação Nº {origem_numero},"
            )
            origem_run.bold = True
            intro_paragraph.add_run(
                f" referente a notificação nº {auto.notificacao_numero}."
            )
        else:
            origem_run = intro_paragraph.add_run(
                f"Auto de Constatação Nº {origem_numero}."
            )
            origem_run.bold = True

        intro_paragraph.add_run(
            " A Autuada fica pelo presente intimada para cumprir ou impugnar o Auto de Infração no prazo de 10 (dez) dias. "
            "No caso de impugnação, deverá ser encaminhada ao Chefe do Setor de Fiscalização do "
        )
        procon_run = intro_paragraph.add_run("PROCON/AM")
        procon_run.bold = True
        intro_paragraph.add_run(
            " no endereço de e-mail: fiscalizacaoprocon@procon.am.gov.br"
        )

    relatorio_cell = table.cell(7, 0)
    relatorio_cell.text = ""
    relatorio_title = relatorio_cell.paragraphs[0]
    relatorio_title_run = relatorio_title.add_run("Relatório")
    relatorio_title_run.bold = True

    relatorio_body = relatorio_cell.add_paragraph()
    relatorio_body.add_run(f"     {auto.relatorio or ''}".rstrip())

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _buscar_auto_constatacao(auto_tipo, auto_id):
    model_map = {
        'banco': AutoBanco,
        'posto': AutoPosto,
        'supermercado': AutoSupermercado,
        'autobanco': AutoBanco,
        'autoposto': AutoPosto,
        'autosupermercado': AutoSupermercado,
        'autodiversos': AutoDiversos,
        'diversos': AutoDiversos,
    }
    model_class = model_map.get(auto_tipo)
    if not model_class:
        return None
    try:
        return model_class.objects.get(pk=auto_id)
    except model_class.DoesNotExist:
        return None


def _vincular_auto_infracao_ao_processo(infracao, auto_obj, request=None):
    try:
        content_type = ContentType.objects.get_for_model(auto_obj)
        processo = Processo.objects.filter(
            auto_constatacao_content_type=content_type,
            auto_constatacao_id=auto_obj.id
        ).first()
        if not processo:
            return

        processo.auto_infracao = infracao
        processo.valor_multa = infracao.valor_multa
        if processo.status == 'aguardando_auto_infracao':
            processo.status = 'aguardando_defesa'
        processo.observacoes = (
            f"{processo.observacoes}\nAuto de Infração {infracao.numero} vinculado ao processo."
            if processo.observacoes else f"Auto de Infração {infracao.numero} vinculado ao processo."
        )
        processo.save(update_fields=['auto_infracao', 'valor_multa', 'status', 'observacoes'])

        docx_bytes = _gerar_docx_auto_infracao(infracao)
        if docx_bytes:
            nome_arquivo = f"Auto_Infracao_{infracao.numero.replace('/', '_')}.docx"
            DocumentoProcesso.objects.update_or_create(
                processo=processo,
                tipo='auto_infracao',
                titulo=f"Auto de Infração {infracao.numero}",
                defaults={
                    'arquivo': ContentFile(docx_bytes, name=nome_arquivo),
                    'descricao': 'Auto de Infração gerado automaticamente',
                    'usuario_upload': (
                        request.user.get_full_name() if request and request.user.is_authenticated else 'Sistema'
                    )
                }
            )
    except Exception as exc:
        logger.exception("Erro ao vincular auto de infração ao processo: %s", exc)


# ========================================
# === VALIDAÇÃO FORMAL ===
# ========================================

@api_view(['POST'])
def validar_formal_infracao(request, pk):
    """
    Registra validação formal do Auto de Infração.
    Espera:
    - status: 'valido' | 'erro' | 'pendente'
    - motivo: obrigatório quando status='erro'
    """
    auto = get_object_or_404(AutoInfracao, pk=pk)

    status_formal = (request.data.get('status') or request.data.get('validacao_formal_status') or '').strip().lower()
    motivo = (request.data.get('motivo') or request.data.get('validacao_formal_motivo') or '').strip()

    if status_formal not in {'valido', 'erro', 'pendente'}:
        return Response({'error': 'Status inválido para validação formal.'}, status=status.HTTP_400_BAD_REQUEST)

    if status_formal == 'erro' and not motivo:
        return Response({'motivo': 'Informe o motivo do erro formal.'}, status=status.HTTP_400_BAD_REQUEST)

    auto.validacao_formal_status = status_formal
    auto.validacao_formal_motivo = motivo if status_formal == 'erro' else ''
    auto.validado_em = timezone.now()
    if request.user and request.user.is_authenticated:
        auto.validado_por = request.user
    else:
        auto.validado_por = None

    if status_formal == 'erro' and auto.status != 'cancelado':
        auto.status = 'cancelado'

    auto.save()

    processo = Processo.objects.filter(auto_infracao=auto).first()
    if processo and status_formal == 'erro':
        motivo_txt = auto.validacao_formal_motivo or 'Sem motivo informado'
        linha = f"AI {auto.numero}: erro formal. Motivo: {motivo_txt}"
        processo.atualizar_status('arquivado', linha)
        processo.observacoes = (f"{processo.observacoes}\n{linha}".strip()
                                if processo.observacoes else linha)
        processo.save(update_fields=['observacoes'])

        try:
            referencia = NotificacaoEletronica.objects.filter(
                Q(processo=processo) | Q(auto_infracao=auto)
            ).order_by('-data_envio', '-id').first()
            destinatario_email = (referencia.destinatario_email or '').strip() if referencia else ''
            destinatario_nome = (referencia.destinatario_nome or '').strip() if referencia else (auto.razao_social or '')

            assunto = f"Processo {processo.numero_processo} arquivado"
            corpo = (
                f"Prezado(a) {destinatario_nome or 'interessado'},\n\n"
                f"Informamos que o processo {processo.numero_processo} "
                f"foi arquivado por erro formal no Auto de Infração {auto.numero}.\n\n"
                f"Motivo: {motivo_txt}\n\n"
                "Em caso de dúvidas, procure o PROCON/AM."
            )

            pendente = NotificacaoEletronica.objects.filter(
                tipo_notificacao='arquivamento',
                processo=processo,
                auto_infracao=auto,
                status__in=['pendente', 'erro'],
            ).order_by('-id').first()

            if pendente:
                pendente.destinatario_nome = destinatario_nome or pendente.destinatario_nome
                pendente.destinatario_email = destinatario_email or pendente.destinatario_email
                pendente.assunto = assunto
                pendente.mensagem = corpo
                pendente.status = 'pendente'
                pendente.save(update_fields=['destinatario_nome', 'destinatario_email', 'assunto', 'mensagem', 'status'])
            else:
                NotificacaoEletronica.objects.create(
                    processo=processo,
                    auto_infracao=auto,
                    tipo_notificacao='arquivamento',
                    destinatario_nome=destinatario_nome or 'Interessado',
                    destinatario_email=destinatario_email or None,
                    destinatario_cpf_cnpj=auto.cnpj or '',
                    representante_legal='',
                    assunto=assunto,
                    mensagem=corpo,
                    status='pendente',
                )
        except Exception as exc:
            logger.exception("Falha ao criar notificacao pendente de arquivamento %s: %s", processo.numero_processo, exc)

    return Response(AutoInfracaoSerializer(auto).data)


# ========================================
# VIEWS FUNCIONAIS PARA INFRAÇÕES
# ========================================

@api_view(['POST'])
def criar_infracao_de_auto(request):
    """
    Cria uma infração baseada em um auto de constatação existente.
    
    Espera:
    - auto_tipo: 'banco', 'posto', 'supermercado', 'diversos'
    - auto_id: ID do auto
    - dados da infração
    """
    try:
        auto_tipo = request.data.get('auto_tipo')
        auto_id = request.data.get('auto_id')
        
        # Mapear tipos para modelos
        model_map = {
            'banco': AutoBanco,
            'posto': AutoPosto,
            'supermercado': AutoSupermercado,
            'diversos': AutoDiversos,
            'autobanco': AutoBanco,
            'autoposto': AutoPosto,
            'autosupermercado': AutoSupermercado,
            'autodiversos': AutoDiversos,
        }
        
        if auto_tipo not in model_map:
            return Response(
                {'error': 'Tipo de auto inválido'}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Buscar o auto
        model_class = model_map[auto_tipo]
        auto_obj = get_object_or_404(model_class, id=auto_id)
        
        # Criar dados da infração
        infracao_data = request.data.copy()
        infracao_data.setdefault('auto_constatacao_numero', getattr(auto_obj, 'numero', ''))
        
        # Pré-preencher campos do auto de constatação
        infracao_data.setdefault('data_fiscalizacao', auto_obj.data_fiscalizacao)
        infracao_data.setdefault('hora_fiscalizacao', auto_obj.hora_fiscalizacao)
        infracao_data.setdefault('municipio', auto_obj.municipio)
        infracao_data.setdefault('estado', auto_obj.estado)
        infracao_data.setdefault('razao_social', auto_obj.razao_social)
        infracao_data.setdefault('nome_fantasia', getattr(auto_obj, 'nome_fantasia', ''))
        infracao_data.setdefault('atividade', getattr(auto_obj, 'atividade', ''))
        infracao_data.setdefault('endereco', auto_obj.endereco)
        infracao_data.setdefault('cep', getattr(auto_obj, 'cep', ''))
        infracao_data.setdefault('cnpj', auto_obj.cnpj)
        infracao_data.setdefault('telefone', getattr(auto_obj, 'telefone', ''))
        
        # Usar serializer para criar
        serializer = AutoInfracaoCreateSerializer(data=infracao_data)
        if serializer.is_valid():
            infracao = serializer.save()

            _vincular_auto_infracao_ao_processo(infracao, auto_obj, request)
            return Response(
                AutoInfracaoSerializer(infracao).data,
                status=status.HTTP_201_CREATED
            )
        else:
            return Response(
                serializer.errors,
                status=status.HTTP_400_BAD_REQUEST
            )
            
    except Exception as e:
        return Response(
            {'error': str(e)},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['GET'])
def infrações_por_auto(request, auto_tipo, auto_id):
    """
    Lista infrações relacionadas a um auto específico.
    """
    try:
        # Mapear tipos para modelos
        model_map = {
            'banco': AutoBanco,
            'posto': AutoPosto,
            'supermercado': AutoSupermercado,
            'diversos': AutoDiversos,
        }
        
        if auto_tipo not in model_map:
            return Response(
                {'error': 'Tipo de auto inválido'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        model_class = model_map[auto_tipo]
        auto_obj = get_object_or_404(model_class, id=auto_id)
        numero_auto = getattr(auto_obj, 'numero', None)

        if not numero_auto:
            return Response([])

        infracoes = AutoInfracao.objects.filter(
            auto_constatacao_numero=numero_auto
        ).order_by('-data_fiscalizacao')
        
        serializer = AutoInfracaoSimpleSerializer(infracoes, many=True)
        return Response(serializer.data)
        
    except Exception as e:
        return Response(
            {'error': str(e)},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['GET'])
def autos_com_potencial_infracao(request):
    """
    Lista autos que têm irregularidades mas ainda não têm infração criada.
    """
    try:
        resultado = []
        
        # AutoBanco com irregularidades
        bancos_irregulares = AutoBanco.objects.exclude(
            Q(nada_consta=True) | Q(sem_irregularidades=True)
        )
        
        for banco in bancos_irregulares:
            if banco.numero and AutoInfracao.objects.filter(auto_constatacao_numero=banco.numero).exists():
                continue
            resultado.append({
                'tipo': 'banco',
                'id': banco.id,
                'numero': banco.numero,
                'razao_social': banco.razao_social,
                'data_fiscalizacao': banco.data_fiscalizacao,
                'irregularidades': 'Diversas irregularidades bancárias'
            })
        
        # AutoPosto com irregularidades
        postos_irregulares = AutoPosto.objects.exclude(
            Q(nada_consta=True) | Q(sem_irregularidades=True)
        )
        
        for posto in postos_irregulares:
            if posto.numero and AutoInfracao.objects.filter(auto_constatacao_numero=posto.numero).exists():
                continue
            resultado.append({
                'tipo': 'posto',
                'id': posto.id,
                'numero': posto.numero,
                'razao_social': posto.razao_social,
                'data_fiscalizacao': posto.data_fiscalizacao,
                'irregularidades': 'Irregularidades em posto de combustível'
            })
        
        # AutoSupermercado com irregularidades
        supermercados_irregulares = AutoSupermercado.objects.exclude(
            nada_consta=True
        )
        
        for supermercado in supermercados_irregulares:
            if supermercado.numero and AutoInfracao.objects.filter(auto_constatacao_numero=supermercado.numero).exists():
                continue
            resultado.append({
                'tipo': 'supermercado',
                'id': supermercado.id,
                'numero': supermercado.numero,
                'razao_social': supermercado.razao_social,
                'data_fiscalizacao': supermercado.data_fiscalizacao,
                'irregularidades': 'Irregularidades em supermercado'
            })
        
        # Ordenar por data mais recente
        resultado.sort(key=lambda x: x['data_fiscalizacao'], reverse=True)
        
        return Response(resultado)
        
    except Exception as e:
        return Response(
            {'error': str(e)},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )


@api_view(['POST'])
def atualizar_status_infracao(request, pk):
    """
    Atualiza o status de uma infração e registra no histórico.
    """
    try:
        infracao = get_object_or_404(AutoInfracao, pk=pk)
        novo_status = request.data.get('status')
        observacoes = request.data.get('observacoes', '')
        usuario = request.data.get('usuario', 'Sistema')
        
        if novo_status not in [choice[0] for choice in STATUS_INFRACAO_CHOICES]:
            return Response(
                {'error': 'Status inválido'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        status_anterior = infracao.status
        infracao.status = novo_status
        if novo_status == 'notificado' and not infracao.data_notificacao:
            infracao.data_notificacao = timezone.now().date()
        infracao.save()
        
        # Registrar no histórico
        HistoricoStatusInfracao.objects.create(
            auto_infracao=infracao,
            status_anterior=status_anterior,
            status_novo=novo_status,
            observacoes=observacoes,
            usuario=usuario
        )
        
        return Response({
            'message': 'Status atualizado com sucesso',
            'status_anterior': status_anterior,
            'status_novo': novo_status
        })
        
    except Exception as e:
        return Response(
            {'error': str(e)},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )
