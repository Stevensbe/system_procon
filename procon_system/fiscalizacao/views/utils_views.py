"""
Views Utilitárias - System Procon

Este módulo contém views para funcionalidades auxiliares como:
- Validações (CNPJ, números)
- Uploads de arquivos
- Geração de documentos
- Busca e filtros
- Relatórios
"""

import io
import os
from django.conf import settings
from django.http import HttpResponse, JsonResponse
from django.shortcuts import get_object_or_404
from django.db.models import Q
from django.utils import timezone
from docx import Document
from rest_framework.decorators import api_view
from rest_framework.response import Response
from rest_framework import status
from django.db.models import Count
import re

from ..models import (
    AutoBanco,
    AutoPosto,
    AutoSupermercado,
    AutoDiversos,
    AutoInfracao,
    AnexoAuto,
    SequenciaAutos,
)

from ..serializers import (
    AutoSimpleSerializer,
)


# ========================================
# VIEWS DE VALIDAÇÃO
# ========================================

@api_view(['POST'])
def validar_cnpj(request):
    """
    Valida formato do CNPJ.
    """
    try:
        cnpj = request.data.get('cnpj', '').strip()
        
        if not cnpj:
            return Response({
                'valido': False,
                'erro': 'CNPJ não fornecido'
            })
        
        # Remove formatação
        cnpj_numeros = re.sub(r'[^\d]', '', cnpj)
        
        # Verifica se tem 14 dígitos
        if len(cnpj_numeros) != 14:
            return Response({
                'valido': False,
                'erro': 'CNPJ deve ter 14 dígitos'
            })
        
        # Verifica se não é sequência de números iguais
        if len(set(cnpj_numeros)) == 1:
            return Response({
                'valido': False,
                'erro': 'CNPJ não pode ter todos os dígitos iguais'
            })
        
        # Validação dos dígitos verificadores
        def calcular_digito(cnpj_base, pesos):
            soma = sum(int(cnpj_base[i]) * pesos[i] for i in range(len(pesos)))
            resto = soma % 11
            return 0 if resto < 2 else 11 - resto
        
        # Primeiro dígito verificador
        pesos1 = [5, 4, 3, 2, 9, 8, 7, 6, 5, 4, 3, 2]
        digito1 = calcular_digito(cnpj_numeros[:12], pesos1)
        
        # Segundo dígito verificador
        pesos2 = [6, 5, 4, 3, 2, 9, 8, 7, 6, 5, 4, 3, 2]
        digito2 = calcular_digito(cnpj_numeros[:13], pesos2)
        
        # Verifica se os dígitos conferem
        if int(cnpj_numeros[12]) == digito1 and int(cnpj_numeros[13]) == digito2:
            # Formatar CNPJ
            cnpj_formatado = f"{cnpj_numeros[:2]}.{cnpj_numeros[2:5]}.{cnpj_numeros[5:8]}/{cnpj_numeros[8:12]}-{cnpj_numeros[12:14]}"
            
            return Response({
                'valido': True,
                'cnpj_formatado': cnpj_formatado,
                'cnpj_numeros': cnpj_numeros
            })
        else:
            return Response({
                'valido': False,
                'erro': 'CNPJ inválido - dígitos verificadores não conferem'
            })
            
    except Exception as e:
        return Response({
            'valido': False,
            'erro': f'Erro na validação: {str(e)}'
        })


@api_view(['GET'])
def proximos_numeros(request):
    """
    Retorna os próximos números disponíveis para cada tipo de auto.
    """
    try:
        ano_atual = timezone.now().year
        
        # Buscar ou criar sequência para o ano atual
        sequencia, created = SequenciaAutos.objects.get_or_create(
            ano=ano_atual,
            defaults={'ultimo_numero': 0}
        )
        
        proximo_numero = sequencia.ultimo_numero + 1
        numero_formatado = f"{proximo_numero:03d}/{ano_atual}"
        
        return Response({
            'proximo_numero': proximo_numero,
            'numero_formatado': numero_formatado,
            'ano': ano_atual,
            'total_gerados_ano': sequencia.ultimo_numero
        })
        
    except Exception as e:
        return Response({
            'error': str(e)
        }, status=500)


# ========================================
# VIEWS DE BUSCA
# ========================================

@api_view(['GET'])
def buscar_autos(request):
    """
    Busca unificada em todos os tipos de auto.
    """
    try:
        query = request.GET.get('q', '').strip()
        tipo = request.GET.get('tipo', 'todos')  # 'banco', 'posto', 'supermercado', 'diversos', 'todos'
        limite = int(request.GET.get('limite', 20))
        
        if not query:
            return Response({
                'resultados': [],
                'total': 0,
                'message': 'Termo de busca não fornecido'
            })
        
        resultados = []
        
        # Função para buscar em um modelo específico
        def buscar_modelo(model_class, tipo_nome):
            if tipo != 'todos' and tipo != tipo_nome:
                return []
            
            autos = model_class.objects.filter(
                Q(numero__icontains=query) |
                Q(razao_social__icontains=query) |
                Q(cnpj__icontains=query) |
                Q(nome_fantasia__icontains=query)
            ).order_by('-data_fiscalizacao')[:limite]
            
            return [{
                'tipo': tipo_nome,
                'id': auto.id,
                'numero': auto.numero,
                'razao_social': auto.razao_social,
                'cnpj': auto.cnpj,
                'data_fiscalizacao': auto.data_fiscalizacao,
                'municipio': auto.municipio,
                'url': f'/fiscalizacao/{tipo_nome}s/{auto.id}/'
            } for auto in autos]
        
        # Buscar em todos os modelos
        resultados.extend(buscar_modelo(AutoBanco, 'banco'))
        resultados.extend(buscar_modelo(AutoPosto, 'posto'))
        resultados.extend(buscar_modelo(AutoSupermercado, 'supermercado'))
        resultados.extend(buscar_modelo(AutoDiversos, 'diversos'))
        
        # Ordenar por data mais recente
        resultados.sort(key=lambda x: x['data_fiscalizacao'], reverse=True)
        
        # Limitar resultados
        resultados = resultados[:limite]
        
        return Response({
            'resultados': resultados,
            'total': len(resultados),
            'termo_busca': query,
            'tipo_filtro': tipo
        })
        
    except Exception as e:
        return Response({
            'error': str(e),
            'resultados': [],
            'total': 0
        }, status=500)


@api_view(['GET'])
def autos_constatacao_disponiveis(request):
    """
    Lista autos de constatação para integrações (PPA, agenda, etc).
    """
    try:
        limite = int(request.GET.get('limite', 50))
        if limite <= 0:
            raise ValueError
        limite = min(limite, 200)
    except (TypeError, ValueError):
        return Response(
            {'error': 'limite invalido. Informe um inteiro positivo.'},
            status=400
        )

    tipo = (request.GET.get('tipo', 'todos') or 'todos').lower()
    somente_disponiveis = bool(request.GET.get('disponiveis'))

    modelos = [
        ('banco', AutoBanco),
        ('posto', AutoPosto),
        ('supermercado', AutoSupermercado),
        ('diversos', AutoDiversos),
    ]

    def serialize_auto(item, tipo_nome):
        return {
            'id': item.id,
            'tipo': tipo_nome,
            'numero': item.numero,
            'razao_social': getattr(item, 'razao_social', ''),
            'empresa_autuada': getattr(
                item,
                'empresa_autuada',
                getattr(item, 'razao_social', '')
            ),
            'cnpj': getattr(item, 'cnpj', ''),
            'municipio': getattr(item, 'municipio', ''),
            'data_fiscalizacao': (
                item.data_fiscalizacao.isoformat()
                if getattr(item, 'data_fiscalizacao', None)
                else None
            ),
            'disponivel': True,
        }

    resultados = []
    for nome, modelo in modelos:
        if tipo not in ('todos', nome):
            continue

        queryset = modelo.objects.all().order_by('-data_fiscalizacao')
        autos = queryset[:limite]
        resultados.extend(serialize_auto(auto, nome) for auto in autos)

    resultados.sort(
        key=lambda registro: registro['data_fiscalizacao'] or '',
        reverse=True
    )

    return Response({
        'results': resultados,
        'total': len(resultados),
        'filters': {
            'tipo': tipo,
            'somente_disponiveis': somente_disponiveis,
            'limite_por_tipo': limite,
        }
    })


# ========================================
# VIEWS DE UPLOAD
# ========================================

@api_view(['POST'])
def upload_anexo(request):
    """
    Upload de anexos para autos.
    """
    try:
        arquivo = request.FILES.get('arquivo')
        content_type_id = request.data.get('content_type_id')
        object_id = request.data.get('object_id')
        descricao = request.data.get('descricao', '')
        
        if not arquivo:
            return Response({
                'error': 'Nenhum arquivo fornecido'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        if not content_type_id or not object_id:
            return Response({
                'error': 'content_type_id e object_id são obrigatórios'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # Verificar tamanho do arquivo (max 10MB por padrão)
        max_size = getattr(settings, 'MAX_FILE_SIZE_MB', 10) * 1024 * 1024
        if arquivo.size > max_size:
            return Response({
                'error': f'Arquivo muito grande. Máximo permitido: {max_size // (1024*1024)}MB'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # Verificar extensão
        allowed_extensions = getattr(settings, 'ALLOWED_EXTENSIONS', 
                                   ['pdf', 'doc', 'docx', 'jpg', 'jpeg', 'png'])
        
        file_extension = arquivo.name.split('.')[-1].lower()
        if file_extension not in allowed_extensions:
            return Response({
                'error': f'Extensão não permitida. Permitidas: {", ".join(allowed_extensions)}'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # Criar anexo
        anexo = AnexoAuto.objects.create(
            content_type_id=content_type_id,
            object_id=object_id,
            arquivo=arquivo,
            descricao=descricao
        )
        
        return Response({
            'id': anexo.id,
            'arquivo_url': anexo.arquivo.url,
            'descricao': anexo.descricao,
            'enviado_em': anexo.enviado_em,
            'message': 'Arquivo enviado com sucesso'
        }, status=status.HTTP_201_CREATED)
        
    except Exception as e:
        return Response({
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


# ========================================
# VIEWS DE GERAÇÃO DE DOCUMENTOS
# ========================================





@api_view(['GET'])
def gerar_documento_banco(request, pk):
    """
    Gera documento Word para Auto de Banco.
    """
    try:
        auto = get_object_or_404(AutoBanco, pk=pk)
        
        # Criar documento
        doc = Document()
        
        # Título
        titulo = doc.add_heading('AUTO DE CONSTATAÇÃO - AGÊNCIA BANCÁRIA', 0)
        titulo.alignment = 1  # Centralizar
        
        # Informações básicas
        doc.add_paragraph(f"Número: {auto.numero}")
        doc.add_paragraph(f"Data: {auto.data_fiscalizacao}")
        doc.add_paragraph(f"Hora: {auto.hora_fiscalizacao}")
        
        # Dados da empresa
        doc.add_heading('DADOS DO ESTABELECIMENTO', level=1)
        doc.add_paragraph(f"Razão Social: {auto.razao_social}")
        doc.add_paragraph(f"Nome Fantasia: {auto.nome_fantasia}")
        doc.add_paragraph(f"CNPJ: {auto.cnpj}")
        doc.add_paragraph(f"Endereço: {auto.endereco}")
        doc.add_paragraph(f"Município: {auto.municipio}")
        
        # Irregularidades encontradas
        doc.add_heading('IRREGULARIDADES CONSTATADAS', level=1)
        
        if auto.nada_consta:
            doc.add_paragraph("✓ NADA CONSTA")
        elif auto.sem_irregularidades:
            doc.add_paragraph("✓ NÃO FORAM ENCONTRADAS IRREGULARIDADES")
        else:
            if auto.ausencia_cartaz_informativo:
                doc.add_paragraph("• Ausência de Cartaz Informativo")
            if auto.ausencia_profissional_libras:
                doc.add_paragraph("• Ausência de profissional de LIBRAS")
            # ... outras irregularidades
        
        if auto.observacoes:
            doc.add_heading('OBSERVAÇÕES', level=1)
            doc.add_paragraph(auto.observacoes)
        
        # Salvar em buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Preparar resposta
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="auto_banco_{auto.numero}.docx"'
        
        return response
        
    except Exception as e:
        return Response({
            'error': str(e)
        }, status=500)


@api_view(['GET'])
def gerar_documento_posto(request, pk):
    """
    Gera documento Word para Auto de Posto.
    """
    try:
        auto = get_object_or_404(AutoPosto, pk=pk)

        template_path = os.path.join(
            settings.BASE_DIR,
            'fiscalizacao',
            'templates',
            'docs',
            'Posto.docx',
        )
        if not os.path.exists(template_path):
            return Response({'error': 'Template do Auto de Posto nao encontrado.'}, status=500)

        doc = Document(template_path)

        def format_date(value):
            if not value:
                return '____/____/____'
            if isinstance(value, str):
                return value
            return value.strftime('%d/%m/%Y')

        def format_time(value):
            if not value:
                return '_____:_____'
            if isinstance(value, str):
                return value[:5]
            return value.strftime('%H:%M')

        def mark_checkbox(value):
            return '(X)' if value else '( )'

        origem = auto.origem or ''
        replacements = {
            '{{numero}}': auto.numero or '',
            '{{razao_social}}': auto.razao_social or '',
            '{{nome_fantasia}}': auto.nome_fantasia or '',
            '{{porte}}': auto.porte or '',
            '{{atuacao}}': auto.atuacao or '',
            '{{atividade}}': auto.atividade or '',
            '{{endereco}}': auto.endereco or '',
            '{{cep}}': auto.cep or '',
            '{{municipio}}': auto.municipio or '',
            '{{cnpj}}': auto.cnpj or '',
            '{{telefone}}': auto.telefone or '',
            '{{origem_outros}}': auto.origem_outros or '',
            '{{hora_fiscalizacao}}': format_time(auto.hora_fiscalizacao),
            '{{data_fiscalizacao}}': format_date(auto.data_fiscalizacao),
            '{{prazo_envio_documentos}}': str(auto.prazo_envio_documentos or ''),
            '{{info_adicionais}}': auto.info_adicionais or '',
            '{{outras_irregularidades}}': auto.outras_irregularidades or '',
            '{{dispositivos_legais}}': auto.dispositivos_legais or '',
            '{{responsavel_nome}}': auto.responsavel_nome or '',
            '{{fiscal_nome_1}}': auto.fiscal_nome_1 or '',
            '{{responsavel_cpf}}': auto.responsavel_cpf or '',
            '{{cb_origem_acao}}': mark_checkbox(origem == 'acao'),
            '{{cb_origem_denuncia}}': mark_checkbox(origem == 'denuncia'),
            '{{cb_origem_forca_tarefa}}': mark_checkbox(origem == 'forca_tarefa'),
            '{{cb_origem_outros}}': mark_checkbox(origem == 'outros'),
            'ESTADO: AM': f"ESTADO: {auto.estado or 'AM'}",
        }

        def apply_replacements(text_value):
            for key, value in replacements.items():
                if key in text_value:
                    text_value = text_value.replace(key, value)
            return text_value

        for paragraph in doc.paragraphs:
            updated = apply_replacements(paragraph.text)
            if updated != paragraph.text:
                paragraph.text = updated

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    updated = apply_replacements(cell.text)
                    if updated != cell.text:
                        cell.text = updated

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        safe_numero = (auto.numero or '').replace('/', '_')
        response['Content-Disposition'] = f'attachment; filename="auto_posto_{safe_numero}.docx"'

        return response
        
    except Exception as e:
        return Response({
            'error': str(e)
        }, status=500)


@api_view(['GET'])
def gerar_documento_supermercado(request, pk):
    """
    Gera documento Word para Auto de Supermercado.
    """
    try:
        auto = get_object_or_404(AutoSupermercado, pk=pk)

        template_path = os.path.join(
            settings.BASE_DIR,
            'fiscalizacao',
            'templates',
            'docs',
            'AutoConstatacaoSupermercado.docx',
        )
        if not os.path.exists(template_path):
            return Response({'error': 'Template do Auto de Supermercado nao encontrado.'}, status=500)

        doc = Document(template_path)

        def format_date(value):
            if not value:
                return '____/____/____'
            if isinstance(value, str):
                return value
            return value.strftime('%d/%m/%Y')

        def format_time(value):
            if not value:
                return '_____:_____'
            if isinstance(value, str):
                return value[:5]
            return value.strftime('%H:%M')

        def mark_checkbox(value):
            return '(X)' if value else '( )'

        def mark_yes_no(value):
            if value is True:
                return '(X)', '( )'
            if value is False:
                return '( )', '(X)'
            return '( )', '( )'

        def build_origem_text():
            origem = auto.origem or ''
            outros_texto = auto.origem_outros or ''
            return (
                'EM CUMPRIMENTO À:\n'
                f"{mark_checkbox(origem == 'acao')} AÇÃO FISCALIZATÓRIA  "
                f"{mark_checkbox(origem == 'denuncia')} DENÚNCIA  "
                f"{mark_checkbox(origem == 'forca_tarefa')} FORÇA TAREFA\n\n"
                f"OUTROS: {outros_texto}"
            )

        def build_cominacao_text():
            data_texto = format_date(auto.data_fiscalizacao)
            hora_texto = format_time(auto.hora_fiscalizacao)
            texto = (
                'COMINAÇÃO LEGAL:\n'
                f"Às {hora_texto} horas do dia {data_texto}, no exercício das competências dispostas no art. 55 e seguintes da "
                'Lei Federal nº 8.078/90, legalmente atribuídas ao Instituto de Defesa do Consumidor - PROCON AMAZONAS, '
                'neste ato fiscalizatório, constatamos que:\n'
                f"          {mark_checkbox(auto.nada_consta)} Não houve nenhuma irregularidade consumerista.\n"
                'O estabelecimento visitado praticou as seguintes irregularidades e/ou violou as seguintes disposições legais:\n'
                f"          {mark_checkbox(auto.comercializar_produtos_vencidos)} Comercializar produtos vencidos - art. 8º caput; art. 18 caput, §6º, incisos I e III; art. 39, inciso VIII, todos do CDC;\n"
                f"          {mark_checkbox(auto.comercializar_embalagem_violada)} Comercializar produtos com embalagem violada - art. 8º caput; art. 18 caput, §6º, incisos II e III; art. 39, inciso VIII, todos do CDC;\n"
                f"          {mark_checkbox(auto.comercializar_lata_amassada)} Comercializar produtos com lata amassada - art. 8º caput; art. 18 caput, §6º, incisos II e III; art. 39, inciso VIII, todos do CDC;\n"
                f"          {mark_checkbox(auto.comercializar_sem_validade)} Comercializar produtos sem validade ou com validade ilegível - art. 8º caput; art. 18 caput, §6º, incisos II e III; art. 31 caput e parágrafo único; art. 37 caput, §1º e §3º; art. 39, inciso VIII, todos do CDC;\n"
                f"          {mark_checkbox(auto.comercializar_mal_armazenados)} Comercializar produtos mal armazenados - art. 8º caput; art. 18 caput, §6º, inciso II; art. 39, inciso VIII do CDC;\n"
                f"          {mark_checkbox(auto.comercializar_descongelados)} Comercializar produtos parcialmente/totalmente descongelados - art. 8º caput; art. 18 caput, §6º, inciso II; art. 39, inciso VIII do CDC;\n"
                f"          {mark_checkbox(auto.publicidade_enganosa)} Publicidade enganosa - art. 37 caput, §1º e §3º; art. 38 do CDC;\n"
                f"          {mark_checkbox(auto.obstrucao_monitor)} Obstrução do monitor, impossibilitando o consumidor de visualizá-lo - art. 1º; art. 2º e art. 3º da Lei Estadual nº 4.777/2019;\n"
                f"          {mark_checkbox(auto.afixacao_precos_fora_padrao)} Afixação de preços fora do padrão estabelecido na Lei nº 10.962/2004 - art. 2º, incisos I e II da Lei nº 10.962/2004;\n"
                f"          {mark_checkbox(auto.ausencia_afixacao_precos)} Ausência de afixação de preços - art. 2º, incisos I e II da Lei nº 10.962/2004;\n"
                f"          {mark_checkbox(auto.afixacao_precos_fracionados_fora_padrao)} Afixação de preços na venda a varejo de produtos fracionados fora do padrão estabelecido na Lei nº 10.962/2004 - art. 2º-A da Lei nº 10.962/2010;\n"
                f"          {mark_checkbox(auto.ausencia_visibilidade_descontos)} Ausência de visibilidade de descontos oferecidos em função do prazo ou instrumento de pagamento utilizado - art. 5º-A da Lei nº 10.962/2010;\n"
                f"          {mark_checkbox(auto.ausencia_placas_promocao_vencimento)} Descumprimento do Art. 1º da Lei nº 7.355/2025 - ausência das placas ou cartazes informativos acerca da data de validade de produtos em promoção que estiverem a menos de dez dias do seu vencimento."
            )
            if auto.cominacao_legal:
                texto = f"{texto}\n{auto.cominacao_legal}"
            return texto

        def build_instrucoes_text():
            prazo = auto.prazo_cumprimento_dias if auto.prazo_cumprimento_dias is not None else '____'
            texto = (
                'De acordo com a Lei Federal Nº 8.078/90, Decreto Estadual Nº 18.606/98 e Legislação Complementar, '
                f"fica o autuado intimado ao cumprimento da obrigação abaixo descrita no prazo de {prazo} dia(s) a partir da lavratura do presente Auto."
            )
            if auto.instrucoes_fiscalizado:
                texto = f"{texto}\n\n{auto.instrucoes_fiscalizado}"
            return texto

        def build_outras_irregularidades_text():
            outras = auto.outras_irregularidades or ''
            return 'OUTRAS IRREGULARIDADES CONSTATADAS/OUTRAS COMINAÇÕES LEGAIS:\n\n' + outras

        def build_narrativa_text():
            narrativa = auto.narrativa_fatos or ''
            return 'NARRATIVA DOS FATOS\n\n' + narrativa

        def build_anexos_text():
            anexo_sim, anexo_nao = mark_yes_no(auto.possui_anexo)
            apreensao_sim, apreensao_nao = mark_yes_no(auto.auto_apreensao)
            pericia_sim, pericia_nao = mark_yes_no(auto.necessita_pericia)
            vicios = mark_checkbox(auto.vicios_aparentes)
            numero = auto.auto_apreensao_numero or '____________'
            return (
                f"POSSUI ANEXO: SIM {anexo_sim} NÃO {anexo_nao}\n"
                f"POSSUI AUTO DE APREENSÃO / INUTILIZAÇÃO  SIM {apreensao_sim} NÃO {apreensao_nao}  Nº{numero}\n"
                'O procedimento administrativo será regulado nos termos do Decreto Estadual 43.614/21 para fins de prazos e impugnações.\n'
                f"Os itens apreendidos e/ou descartados necessitam de perícia; {pericia_sim} SIM  {pericia_nao} NÃO, todos os vícios que os tornam impróprios estavam aparentes. {vicios}"
            )

        replacements = {
            'RAZÃO SOCIAL:': f"RAZÃO SOCIAL: {auto.razao_social or ''}",
            'NOME FANTASIA:': f"NOME FANTASIA: {auto.nome_fantasia or ''}",
            'PORTE:': f"PORTE: {auto.porte or ''}",
            'ATUAÇÃO:': f"ATUAÇÃO: {auto.atuacao or ''}",
            'ATIVIDADE:': f"ATIVIDADE: {auto.atividade or ''}",
            'ENDEREÇO:': f"ENDEREÇO: {auto.endereco or ''}",
            'CEP:': f"CEP: {auto.cep or ''}",
            'MUNICÍPIO:': f"MUNICÍPIO: {auto.municipio or ''}",
            'ESTADO: AMAZONAS': f"ESTADO: {auto.estado or 'AM'}",
            'CNPJ:': f"CNPJ: {auto.cnpj or ''}",
            'TELEFONE:': f"TELEFONE: {auto.telefone or ''}",
        }

        def apply_replacements(text_value):
            text_value = re.sub(
                r'AUTO DE CONSTATAÇÃO Nº\s*/\s*\.',
                f"AUTO DE CONSTATAÇÃO Nº {auto.numero or ''}",
                text_value,
            )
            for key, value in replacements.items():
                if key in text_value:
                    text_value = text_value.replace(key, value)
            return text_value

        for paragraph in doc.paragraphs:
            updated = apply_replacements(paragraph.text)
            if 'EM CUMPRIMENTO' in updated:
                paragraph.text = build_origem_text()
                continue
            if 'COMINAÇÃO LEGAL' in updated:
                paragraph.text = build_cominacao_text()
                continue
            if 'De acordo com a Lei Federal' in updated:
                paragraph.text = build_instrucoes_text()
                continue
            if 'OUTRAS IRREGULARIDADES CONSTATADAS/OUTRAS COMINAÇÕES LEGAIS' in updated:
                paragraph.text = build_outras_irregularidades_text()
                continue
            if 'NARRATIVA DOS FATOS' in updated:
                paragraph.text = build_narrativa_text()
                continue
            if 'POSSUI ANEXO' in updated:
                paragraph.text = build_anexos_text()
                continue
            if updated != paragraph.text:
                paragraph.text = updated

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    texto = cell.text
                    if 'EM CUMPRIMENTO' in texto:
                        cell.text = build_origem_text()
                        continue
                    if 'COMINAÇÃO LEGAL' in texto:
                        cell.text = build_cominacao_text()
                        continue
                    if 'De acordo com a Lei Federal' in texto:
                        cell.text = build_instrucoes_text()
                        continue
                    if 'OUTRAS IRREGULARIDADES CONSTATADAS/OUTRAS COMINAÇÕES LEGAIS' in texto:
                        cell.text = build_outras_irregularidades_text()
                        continue
                    if 'NARRATIVA DOS FATOS' in texto:
                        cell.text = build_narrativa_text()
                        continue
                    if 'POSSUI ANEXO' in texto:
                        cell.text = build_anexos_text()
                        continue

                    updated = apply_replacements(texto)
                    if updated != texto:
                        cell.text = updated

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        safe_numero = (auto.numero or '').replace('/', '_')
        response['Content-Disposition'] = f'attachment; filename="auto_supermercado_{safe_numero}.docx"'

        return response

    except Exception as e:
        return Response({
            'error': str(e)
        }, status=500)


@api_view(['GET'])
def gerar_documento_diversos(request, pk):
    """
    Gera documento Word para Auto Diversos.
    """
    try:
        auto = get_object_or_404(AutoDiversos, pk=pk)

        template_path = os.path.join(
            settings.BASE_DIR,
            'fiscalizacao',
            'templates',
            'docs',
            'Diversos.docx',
        )
        if not os.path.exists(template_path):
            return Response({'error': 'Template do Auto Diversos nao encontrado.'}, status=500)

        doc = Document(template_path)

        def format_date(value):
            if not value:
                return '____/____/____'
            if isinstance(value, str):
                return value
            return value.strftime('%d/%m/%Y')

        def format_time(value):
            if not value:
                return '_____:_____'
            if isinstance(value, str):
                return value[:5]
            return value.strftime('%H:%M')

        def mark_checkbox(value):
            return '(X)' if value else '( )'

        origem = auto.origem or ''
        replacements = {
            '{{numero}}': auto.numero or '',
            '{{razao_social}}': auto.razao_social or '',
            '{{nome_fantasia}}': auto.nome_fantasia or '',
            '{{porte}}': auto.porte or '',
            '{{atuacao}}': auto.atuacao or '',
            '{{atividade}}': auto.atividade or '',
            '{{endereco}}': auto.endereco or '',
            '{{cep}}': auto.cep or '',
            '{{municipio}}': auto.municipio or '',
            '{{cnpj}}': auto.cnpj or '',
            '{{telefone}}': auto.telefone or '',
            '{{hora_fiscalizacao}}': format_time(auto.hora_fiscalizacao),
            '{{data_fiscalizacao}}': format_date(auto.data_fiscalizacao),
            '{{outras_irregularidades}}': auto.outras_irregularidades or '',
            '{{narrativa_fatos}}': auto.narrativa_fatos or '',
            '{{origem_outros}}': auto.origem_outros or '',
            '{{responsavel_nome}}': auto.responsavel_nome or '',
            '{{fiscal_nome_1}}': auto.fiscal_nome_1 or '',
            '{{responsavel_cpf}}': auto.responsavel_cpf or '',
            '{{cb_publicidade_enganosa}}': mark_checkbox(auto.publicidade_enganosa),
            '{{cb_afixacao_fora_padrao}}': mark_checkbox(auto.afixacao_precos_fora_padrao),
            '{{cb_ausencia_afixacao}}': mark_checkbox(auto.ausencia_afixacao_precos),
            '{{cb_eletronico_fora_padrao}}': mark_checkbox(auto.afixacao_precos_eletronico_fora_padrao),
            '{{cb_ausencia_eletronico}}': mark_checkbox(auto.ausencia_afixacao_precos_eletronico),
            '{{cb_fracionados_fora_padrao}}': mark_checkbox(auto.afixacao_precos_fracionados_fora_padrao),
            '{{cb_ausencia_descontos}}': mark_checkbox(auto.ausencia_visibilidade_descontos),
            '{{cb_ausencia_cdc}}': mark_checkbox(auto.ausencia_exemplar_cdc),
            '{{cb_substituicao_troco}}': mark_checkbox(auto.substituicao_troco),
            '{{cb_advertencia}}': mark_checkbox(auto.advertencia),
            '{{cb_origem_acao}}': mark_checkbox(origem == 'acao'),
            '{{cb_origem_denuncia}}': mark_checkbox(origem == 'denuncia'),
            '{{cb_origem_forca_tarefa}}': mark_checkbox(origem == 'forca_tarefa'),
            '{{cb_origem_outros}}': mark_checkbox(origem == 'outros'),
            'ESTADO: AM': f"ESTADO: {auto.estado or 'AM'}",
        }

        def apply_replacements(text_value):
            for key, value in replacements.items():
                if key in text_value:
                    text_value = text_value.replace(key, value)
            return text_value

        for paragraph in doc.paragraphs:
            updated = apply_replacements(paragraph.text)
            if updated != paragraph.text:
                paragraph.text = updated

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    updated = apply_replacements(cell.text)
                    if updated != cell.text:
                        cell.text = updated

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        safe_numero = (auto.numero or '').replace('/', '_')
        response['Content-Disposition'] = f'attachment; filename=\"auto_diversos_{safe_numero}.docx\"'

        return response
        
    except Exception as e:
        return Response({
            'error': str(e)
        }, status=500)


# ========================================
# RELATÓRIOS
# ========================================

@api_view(['GET'])
def relatorio_consolidado(request):
    """
    Gera relatório consolidado de todas as atividades.
    """
    try:
        data_inicio = request.GET.get('data_inicio')
        data_fim = request.GET.get('data_fim')
        
        if not data_inicio or not data_fim:
            return Response({
                'error': 'data_inicio e data_fim são obrigatórios'
            }, status=400)
        
        # Filtros de data
        filtro_data = {
            'data_fiscalizacao__range': [data_inicio, data_fim]
        }
        
        # Coletar dados
        dados = {
            'periodo': {
                'inicio': data_inicio,
                'fim': data_fim
            },
            'autos': {
                'banco': {
                    'total': AutoBanco.objects.filter(**filtro_data).count(),
                    'com_irregularidades': AutoBanco.objects.filter(
                        **filtro_data
                    ).exclude(Q(nada_consta=True) | Q(sem_irregularidades=True)).count()
                },
                'posto': {
                    'total': AutoPosto.objects.filter(**filtro_data).count(),
                    'com_irregularidades': AutoPosto.objects.filter(
                        **filtro_data
                    ).exclude(Q(nada_consta=True) | Q(sem_irregularidades=True)).count()
                },
                'supermercado': {
                    'total': AutoSupermercado.objects.filter(**filtro_data).count(),
                    'com_irregularidades': AutoSupermercado.objects.filter(
                        **filtro_data
                    ).exclude(nada_consta=True).count()
                },
                'diversos': {
                    'total': AutoDiversos.objects.filter(**filtro_data).count()
                }
            },
            'infracoes': {
                'total': AutoInfracao.objects.filter(**filtro_data).count(),
                'por_gravidade': dict(
                    AutoInfracao.objects.filter(**filtro_data).values('gravidade').annotate(
                        count=Count('id')
                    ).values_list('gravidade', 'count')
                )
            }
        }
        
        return Response(dados)
        
    except Exception as e:
        return Response({
            'error': str(e)
        }, status=500)
