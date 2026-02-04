"""
Views para Processos Administrativos - System Procon

Este módulo contém views relacionadas aos Processos Administrativos,
incluindo gestão completa do ciclo de vida dos processos.
"""

from django.http import JsonResponse, HttpResponse
from django.shortcuts import get_object_or_404
from django.db.models import Count, Q, Sum, Avg
from django.utils import timezone
from django.core.cache import cache
from django.conf import settings
from django.core.files.base import ContentFile
from rest_framework import generics, status
from rest_framework.decorators import api_view
from rest_framework.response import Response
from rest_framework.pagination import PageNumberPagination
from datetime import datetime, timedelta
from decimal import Decimal, InvalidOperation
from io import BytesIO
from pathlib import Path
import shutil
import subprocess
from openpyxl import load_workbook
from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt
import tempfile

from protocolo_tramitacao.models import ProtocoloDocumento, TipoDocumento, Setor, TramitacaoDocumento
from protocolo_tramitacao.services import workflow_service
from caixa_entrada.models import CaixaEntrada
from caixa_entrada.services import SETOR_LABELS

from ..models import (
    Processo,
    HistoricoProcesso,
    DocumentoProcesso,
    ParecerProcesso,
    AutoInfracao,
)

from cobranca.models import GuiaRecolhimentoMulta, ConfiguracaoCobranca

from ..serializers import (
    ProcessoSimpleSerializer,
    ProcessoDetailSerializer,
    ProcessoCreateUpdateSerializer,
    ProcessoEstatisticasSerializer,
    ProcessoResumoMensalSerializer,
    ProcessoFiltroSerializer,
    ProcessoBuscaSerializer,
    DocumentoUploadSerializer,
    DocumentoProcessoSerializer,
    ParecerProcessoSerializer,
    AtualizarStatusProcessoSerializer,
    HistoricoProcessoSerializer,
)


# ========================================
# VIEWS PRINCIPAIS DE PROCESSO
# ========================================

class ProcessoListCreateAPIView(generics.ListCreateAPIView):
    queryset = Processo.objects.all()
    
    def get_serializer_class(self):
        if self.request.method == 'POST':
            return ProcessoCreateUpdateSerializer
        return ProcessoSimpleSerializer
    
    def get_queryset(self):
        queryset = super().get_queryset()
        
        # Filtros
        status_filter = self.request.query_params.get('status')
        prioridade = self.request.query_params.get('prioridade')
        prazo_vencendo = self.request.query_params.get('prazo_vencendo')
        data_inicio = self.request.query_params.get('data_inicio')
        data_fim = self.request.query_params.get('data_fim')
        
        if status_filter:
            queryset = queryset.filter(status=status_filter)
        if prioridade:
            queryset = queryset.filter(prioridade=prioridade)
        if data_inicio:
            queryset = queryset.filter(criado_em__gte=data_inicio)
        if data_fim:
            queryset = queryset.filter(criado_em__lte=data_fim)
        if prazo_vencendo == 'true':
            hoje = timezone.now().date()
            limite = hoje + timedelta(days=3)
            queryset = queryset.filter(
                Q(prazo_defesa__lte=limite) | Q(prazo_recurso__lte=limite)
            )

        queryset = _filtrar_processos_por_usuario(queryset, self.request)
        return queryset


class ProcessoDetailAPIView(generics.RetrieveUpdateDestroyAPIView):
    queryset = Processo.objects.all()
    
    def get_serializer_class(self):
        if self.request.method in ['PUT', 'PATCH']:
            return ProcessoCreateUpdateSerializer
        return ProcessoDetailSerializer

    def get_queryset(self):
        queryset = super().get_queryset()
        return _filtrar_processos_por_usuario(queryset, self.request)


# ========================================
# VIEWS FUNCIONAIS PARA PROCESSOS
# ========================================

@api_view(['POST'])
def atualizar_status_processo(request, pk):
    """
    Atualiza status do processo e registra no histórico.
    """
    try:
        processo = get_object_or_404(Processo, pk=pk)
        
        serializer = AtualizarStatusProcessoSerializer(data=request.data)
        if not serializer.is_valid():
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        
        novo_status = serializer.validated_data['status']
        observacao = serializer.validated_data.get('observacao', '')
        usuario = serializer.validated_data.get('usuario', 'Sistema')
        
        # Atualizar usando método do modelo
        processo.atualizar_status(novo_status, observacao)
        
        return Response({
            'message': 'Status atualizado com sucesso',
            'processo': ProcessoDetailSerializer(processo).data
        })
        
    except Exception as e:
        return Response({
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


def _parse_decimal(valor):
    if valor in (None, ''):
        return None
    try:
        texto = str(valor).strip().replace(',', '.')
        return Decimal(texto)
    except (InvalidOperation, ValueError, TypeError):
        return None


def _nome_usuario(request):
    usuario = getattr(request, 'user', None)
    if usuario and getattr(usuario, 'is_authenticated', False):
        return usuario.get_full_name() or usuario.username
    return 'sistema'


def _registrar_evento_processo(processo: Processo, observacao: str, usuario: str):
    try:
        HistoricoProcesso.objects.create(
            processo=processo,
            status_anterior=processo.status,
            status_novo=processo.status,
            observacao=observacao or '',
            usuario=usuario,
        )
    except Exception:
        pass


def _filtrar_processos_por_usuario(queryset, request):
    """Restringe processos ao que estiver na caixa do usuario (setor ou pessoal)."""
    usuario = getattr(request, 'user', None)
    if not usuario or not getattr(usuario, 'is_authenticated', False):
        return queryset.none()
    if getattr(usuario, 'is_superuser', False) or getattr(usuario, 'is_staff', False):
        return queryset

    try:
        from caixa_entrada.views import filtrar_documentos_por_usuario  # import local evita ciclos
    except Exception:
        return queryset

    documentos = filtrar_documentos_por_usuario(CaixaEntrada.objects.all(), request, apenas_pessoal=False)
    protocolo_ids = documentos.exclude(protocolo_id__isnull=True).values_list('protocolo_id', flat=True)
    if not protocolo_ids:
        return queryset.none()
    processo_ids = (
        ProtocoloDocumento.objects.filter(id__in=protocolo_ids)
        .exclude(processo_fiscalizacao_id__isnull=True)
        .values_list('processo_fiscalizacao_id', flat=True)
    )
    if not processo_ids:
        return queryset.none()
    return queryset.filter(id__in=processo_ids)


@api_view(['POST'])
def registrar_calculo_multa(request, pk):
    """Registra o calculo inicial/atualizado da multa no processo."""
    processo = get_object_or_404(Processo, pk=pk)

    valor_decimal = _parse_decimal(request.data.get('valor_multa'))
    if valor_decimal is None or valor_decimal <= 0:
        return Response(
            {'detail': 'Informe um valor de multa valido.'},
            status=status.HTTP_400_BAD_REQUEST
        )

    observacao = (request.data.get('observacao') or '').strip()
    origem = (request.data.get('origem') or 'calculo').strip().lower()

    processo.valor_multa = valor_decimal
    if processo.valor_final in (None, ''):
        processo.valor_final = valor_decimal
    processo.save()

    usuario = _nome_usuario(request)
    texto_evento = f"Calculo de multa registrado ({origem}). Valor: R$ {valor_decimal}."
    if observacao:
        texto_evento = f"{texto_evento} {observacao}"
    _registrar_evento_processo(processo, texto_evento, usuario)

    return Response(
        ProcessoDetailSerializer(processo).data,
        status=status.HTTP_200_OK
    )


@api_view(['POST'])
def registrar_impugnacao_valor(request, pk):
    """Registra impugnacao do valor da multa, sem alterar o status."""
    processo = get_object_or_404(Processo, pk=pk)

    observacao = (request.data.get('observacao') or '').strip()
    if not observacao:
        return Response(
            {'detail': 'Descreva a impugnacao do valor da multa.'},
            status=status.HTTP_400_BAD_REQUEST
        )

    processo.observacoes_internas = (processo.observacoes_internas or '').strip()
    if processo.observacoes_internas:
        processo.observacoes_internas = f"{processo.observacoes_internas}\n\nImpugnacao: {observacao}"
    else:
        processo.observacoes_internas = f"Impugnacao: {observacao}"
    processo.save()

    usuario = _nome_usuario(request)
    _registrar_evento_processo(
        processo,
        f"Impugnacao do valor da multa registrada. {observacao}",
        usuario,
    )

    return Response(
        ProcessoDetailSerializer(processo).data,
        status=status.HTTP_200_OK
    )


@api_view(['POST'])
def registrar_recalculo_multa(request, pk):
    """Registra o recalculo da multa apos impugnacao/analise."""
    processo = get_object_or_404(Processo, pk=pk)

    valor_decimal = _parse_decimal(request.data.get('valor_multa'))
    if valor_decimal is None or valor_decimal <= 0:
        return Response(
            {'detail': 'Informe um valor de multa valido para o recalculo.'},
            status=status.HTTP_400_BAD_REQUEST
        )

    observacao = (request.data.get('observacao') or '').strip()

    processo.valor_multa = valor_decimal
    processo.valor_final = valor_decimal
    processo.save()

    usuario = _nome_usuario(request)
    texto_evento = f"Recalculo de multa registrado. Novo valor: R$ {valor_decimal}."
    if observacao:
        texto_evento = f"{texto_evento} {observacao}"
    _registrar_evento_processo(processo, texto_evento, usuario)

    return Response(
        ProcessoDetailSerializer(processo).data,
        status=status.HTTP_200_OK
    )


def _buscar_setor_por_termos(siglas, nomes):
    qs = Setor.objects.all()
    for sigla in siglas:
        setor = qs.filter(sigla__icontains=sigla).first()
        if setor:
            return setor
    for nome in nomes:
        setor = qs.filter(nome__icontains=nome).first()
        if setor:
            return setor
    return None


def _obter_usuario_responsavel(request):
    usuario = getattr(request, 'user', None)
    if usuario and getattr(usuario, 'is_authenticated', False):
        return usuario
    # fallback seguro
    from django.contrib.auth.models import User

    return User.objects.filter(is_staff=True, is_active=True).order_by('id').first()


def _obter_setor_juridico2():
    setor = _buscar_setor_por_termos(
        ['JUR2', 'JUR', 'REC'],
        ['Juridico 2', 'Recursos', SETOR_LABELS.get('JURIDICO_2_RECURSOS', '')],
    )
    if setor:
        return setor
    return Setor.objects.create(
        nome=SETOR_LABELS.get('JURIDICO_2_RECURSOS', 'Juridico 2 - Recursos'),
        sigla='JUR2',
        pode_protocolar=True,
        pode_tramitar=True,
    )


def _obter_setor_fiscalizacao():
    setor = _buscar_setor_por_termos(['FISC'], ['Fiscalizacao'])
    if setor:
        return setor
    return Setor.objects.create(
        nome=SETOR_LABELS.get('FISCALIZACAO', 'Fiscalizacao'),
        sigla='FISC',
        pode_protocolar=True,
        pode_tramitar=True,
    )


def _obter_protocolo_processo(processo: Processo, usuario, setor_origem: Setor):
    protocolo = (
        ProtocoloDocumento.objects.filter(processo_fiscalizacao=processo)
        .order_by('-data_protocolo')
        .first()
    )
    if protocolo:
        return protocolo, False

    tipo_doc, _ = TipoDocumento.objects.get_or_create(
        nome='Processo Administrativo',
        defaults={
            'descricao': 'Tramitacao do processo administrativo',
            'prazo_resposta_dias': 15,
            'requer_assinatura': False,
        },
    )

    assunto = f"Processo {processo.numero_processo}"
    descricao = (
        "Tramitacao do processo administrativo no fluxo interno. "
        f"Processo: {processo.numero_processo}."
    )

    protocolo = ProtocoloDocumento.objects.create(
        tipo_documento=tipo_doc,
        origem='INTERNO',
        assunto=assunto,
        descricao=descricao,
        status='EM_TRAMITACAO',
        remetente_nome=processo.autuado,
        remetente_documento=processo.cnpj,
        remetente_email='',
        remetente_telefone='',
        processo_fiscalizacao=processo,
        auto_infracao=processo.auto_infracao,
        setor_atual=setor_origem,
        setor_origem=setor_origem,
        protocolado_por=usuario,
        responsavel_atual=getattr(setor_origem, 'responsavel', None),
        observacoes='',
        sigiloso=False,
    )

    return protocolo, True


def _criar_despacho_daf(processo: Processo, usuario, observacao: str, prazo_dias: int):
    setor_daf = _buscar_setor_por_termos(['DAF'], ['Finance', 'Financeiro', 'Administrativa'])
    if not setor_daf:
        setor_daf = Setor.objects.create(nome='Diretoria Administrativa Financeira', sigla='DAF')

    setor_origem = _buscar_setor_por_termos(['DJUR', 'JUR', 'FISC'], ['Juridico', 'Fiscalizacao'])
    if not setor_origem:
        setor_origem = setor_daf

    tipo_doc, _ = TipoDocumento.objects.get_or_create(
        nome='Despacho ao DAF',
        defaults={
            'descricao': 'Despacho para emissao de GRM/encaminhamento financeiro',
            'prazo_resposta_dias': prazo_dias,
            'requer_assinatura': False,
        },
    )

    assunto = f"Despacho ao DAF - Processo {processo.numero_processo}"
    descricao = (
        f"Encaminhamento ao DAF para emissao de GRM/acoes financeiras. "
        f"Processo: {processo.numero_processo}."
    )
    if observacao:
        descricao = f"{descricao}\n\n{observacao.strip()}"

    prazo_resposta = timezone.now() + timedelta(days=prazo_dias)

    protocolo_existente = ProtocoloDocumento.objects.filter(
        processo_fiscalizacao=processo,
        setor_atual=setor_daf,
        assunto__icontains='Despacho ao DAF',
    ).order_by('-data_protocolo').first()

    if protocolo_existente:
        return protocolo_existente, None

    protocolo = ProtocoloDocumento.objects.create(
        tipo_documento=tipo_doc,
        origem='INTERNO',
        assunto=assunto,
        descricao=descricao,
        status='EM_TRAMITACAO',
        remetente_nome=processo.autuado,
        remetente_documento=processo.cnpj,
        remetente_email='',
        remetente_telefone='',
        processo_fiscalizacao=processo,
        auto_infracao=processo.auto_infracao,
        setor_atual=setor_daf,
        setor_origem=setor_origem,
        prazo_resposta=prazo_resposta,
        protocolado_por=usuario,
        responsavel_atual=None,
        observacoes=observacao or '',
        sigiloso=False,
    )

    TramitacaoDocumento.objects.create(
        protocolo=protocolo,
        acao='ENCAMINHADO',
        setor_origem=setor_origem,
        setor_destino=setor_daf,
        motivo=observacao or 'Despacho ao DAF para emissao de GRM.',
        observacoes='Encaminhamento automatico pelo modulo de processos.',
        prazo_dias=prazo_dias,
        usuario=usuario,
    )

    setor_destino_nome = SETOR_LABELS.get('DAF', 'Diretoria Administrativa Financeira')
    caixa = CaixaEntrada.objects.create(
        tipo_documento='DOCUMENTO_INTERNO',
        assunto=assunto,
        descricao=descricao,
        prioridade='NORMAL',
        remetente_nome=getattr(usuario, 'get_full_name', lambda: '')() or getattr(usuario, 'username', 'sistema'),
        remetente_documento=processo.cnpj,
        remetente_email='',
        remetente_telefone='',
        empresa_nome=processo.autuado,
        empresa_cnpj=processo.cnpj,
        setor_destino=setor_destino_nome,
        responsavel_atual=None,
        protocolo=protocolo,
        setor_lotacao=setor_destino_nome,
        origem='PROCESSOS',
        prazo_resposta=prazo_resposta,
        content_type=None,
        object_id=None,
    )

    return protocolo, caixa


def _gerar_despacho_daf_docx(processo: Processo, observacao: str, prazo_dias: int, usuario_nome: str):
    """Gera um despacho simples ao DAF em DOCX para anexar ao processo."""
    agora = timezone.now()
    data_fmt = agora.strftime('%d/%m/%Y')
    hora_fmt = agora.strftime('%H:%M')
    carimbo = agora.strftime('%Y%m%d_%H%M%S')
    numero_limpo = (processo.numero_processo or '').replace('/', '_')
    nome_arquivo = f"Despacho_DAF_{numero_limpo}_{carimbo}.docx"

    doc = Document()

    # Título e identificação do despacho
    doc.add_heading('DESPACHO AO DAF', level=1)
    doc.add_paragraph(f"Processo: {processo.numero_processo}")
    doc.add_paragraph(f"Autuado: {processo.autuado}")
    doc.add_paragraph(f"Data/Hora: {data_fmt} {hora_fmt}")

    doc.add_paragraph(
        (
            "Encaminho o presente processo ao Departamento Administrativo Financeiro "
            "(DAF) para emissao da GRM e demais providencias financeiras cabiveis, "
            "conforme decisao registrada no processo."
        )
    )

    doc.add_paragraph(
        f"Prazo sugerido para retorno/manifestacao do DAF: {prazo_dias} dia(s)."
    )

    if observacao:
        doc.add_paragraph("Observacoes:")
        doc.add_paragraph(observacao.strip())

    doc.add_paragraph("")
    doc.add_paragraph(f"Responsavel pelo despacho: {usuario_nome or 'sistema'}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return ContentFile(buffer.read(), name=nome_arquivo)


def _registrar_documento_despacho_daf(
    processo: Processo,
    observacao: str,
    prazo_dias: int,
    usuario_nome: str,
    arquivo_upload=None,
):
    """Anexa o despacho ao DAF como DocumentoProcesso."""
    agora = timezone.now()
    carimbo = agora.strftime('%d/%m/%Y %H:%M')
    titulo = f"Despacho ao DAF - {processo.numero_processo} - {carimbo}"
    descricao = (
        "Despacho para emissao de GRM/encaminhamento financeiro ao DAF. "
        f"Prazo sugerido: {prazo_dias} dia(s)."
    )
    if observacao:
        descricao = f"{descricao}\n\n{observacao.strip()}"

    if arquivo_upload:
        arquivo = arquivo_upload
        nome_arquivo = getattr(arquivo_upload, 'name', '')
        if nome_arquivo:
            titulo = f"Despacho ao DAF - {processo.numero_processo} - {nome_arquivo}"
    else:
        arquivo = _gerar_despacho_daf_docx(processo, observacao, prazo_dias, usuario_nome)

    documento = DocumentoProcesso.objects.create(
        processo=processo,
        tipo='outros',
        titulo=titulo,
        descricao=descricao,
        arquivo=arquivo,
        usuario_upload=usuario_nome or 'sistema',
    )
    return documento


def _criar_grm_automatico(processo: Processo, usuario_nome: str, observacao: str = ''):
    """Cria a GRM automaticamente ao despachar para o DAF."""
    if not processo:
        return None

    grm_existente = GuiaRecolhimentoMulta.objects.filter(processo=processo).order_by('-criado_em').first()
    if grm_existente:
        return grm_existente

    config = ConfiguracaoCobranca.objects.filter(ativo=True).first()
    auto_infracao = processo.auto_infracao

    valor_integral = processo.valor_final or processo.valor_multa
    if valor_integral is None and auto_infracao:
        valor_integral = auto_infracao.valor_multa

    vencimento = None
    prazo_vencimento = 30
    if config and getattr(config, 'prazo_vencimento_boleto', None):
        prazo_vencimento = config.prazo_vencimento_boleto
    vencimento = timezone.now().date() + timedelta(days=prazo_vencimento)

    multa_obj = None
    if auto_infracao and hasattr(auto_infracao, 'multa'):
        try:
            multa_obj = auto_infracao.multa
        except Exception:
            multa_obj = None

    grm = GuiaRecolhimentoMulta.objects.create(
        processo=processo,
        auto_infracao=auto_infracao,
        multa=multa_obj,
        departamento_emissor='DAF',
        recebedor_nome=getattr(config, 'beneficiario_nome', '') or 'Instituto de Defesa do Consumidor - PROCON AM',
        recebedor_cnpj=getattr(config, 'beneficiario_cnpj', '') or '',
        banco_nome=getattr(config, 'banco_nome', '') or '',
        banco_agencia=getattr(config, 'agencia', '') or '',
        banco_conta=getattr(config, 'conta', '') or '',
        autuado_nome=processo.autuado,
        autuado_documento=processo.cnpj,
        numero_auto_infracao=getattr(auto_infracao, 'numero_auto', '') if auto_infracao else '',
        numero_processo=processo.numero_processo,
        valor_integral=valor_integral,
        valor_a_vista=valor_integral,
        vencimento=vencimento,
        observacao_texto=observacao or '',
        criado_por=usuario_nome or 'sistema',
    )

    try:
        grm.gerar_docx(salvar=True)
    except Exception:
        pass

    _registrar_evento_processo(
        processo,
        f"GRM gerada automaticamente (DAF). Numero: {grm.numero_guia}.",
        usuario_nome or 'sistema',
    )

    return grm


@api_view(['POST'])
def despachar_para_daf(request, pk):
    """Despacha manualmente o processo para o DAF (emissao de GRM)."""
    processo = get_object_or_404(Processo, pk=pk)

    if processo.status not in {'finalizado_procedente', 'aguardando_recurso', 'recurso_apresentado', 'julgamento'}:
        return Response(
            {'detail': 'O despacho ao DAF deve ocorrer apos decisao procedente.'},
            status=status.HTTP_400_BAD_REQUEST
        )

    observacao = (request.data.get('observacao') or '').strip()
    prazo_dias_raw = request.data.get('prazo_dias')
    try:
        prazo_dias = int(prazo_dias_raw) if prazo_dias_raw not in (None, '') else 15
        if prazo_dias <= 0:
            prazo_dias = 15
    except (TypeError, ValueError):
        prazo_dias = 15

    usuario = _obter_usuario_responsavel(request)
    if not usuario:
        return Response(
            {'detail': 'Nenhum usuario disponivel para protocolar o despacho.'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )

    arquivo_upload = request.FILES.get('arquivo') or request.FILES.get('despacho_arquivo')
    if arquivo_upload:
        nome_arquivo = (getattr(arquivo_upload, 'name', '') or '').lower()
        if not any(nome_arquivo.endswith(ext) for ext in ['.doc', '.docx', '.pdf']):
            return Response(
                {'detail': 'Formato invalido. Envie DOC, DOCX ou PDF.'},
                status=status.HTTP_400_BAD_REQUEST
            )

    protocolo, caixa = _criar_despacho_daf(processo, usuario, observacao, prazo_dias)
    usuario_nome = _nome_usuario(request)
    documento = _registrar_documento_despacho_daf(
        processo,
        observacao,
        prazo_dias,
        usuario_nome,
        arquivo_upload=arquivo_upload,
    )
    documento_data = DocumentoProcessoSerializer(documento, context={'request': request}).data

    _registrar_evento_processo(
        processo,
        (
            "Despacho ao DAF registrado. "
            f"Protocolo: {protocolo.numero_protocolo}. "
            f"Documento: {documento.titulo}."
        ),
        usuario_nome,
    )

    return Response(
        {
            'success': True,
            'processo': processo.numero_processo,
            'protocolo': protocolo.numero_protocolo,
            'setor_destino': getattr(protocolo.setor_atual, 'nome', ''),
            'caixa_id': str(getattr(caixa, 'id', '')) if caixa else None,
            'documento': documento_data,
            'grm_id': None,
            'grm_numero': None,
        },
        status=status.HTTP_200_OK
    )


@api_view(['POST'])
def tramitar_processo(request, pk):
    """Tramita o processo para outro setor (padrão: Jurídico 2 - Recursos)."""
    processo = get_object_or_404(Processo, pk=pk)

    usuario = _obter_usuario_responsavel(request)
    if not usuario:
        return Response(
            {'detail': 'Nenhum usuario disponivel para tramitar o processo.'},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR,
        )

    setor_destino = None
    setor_destino_id = request.data.get('setor_destino_id') or request.data.get('setor_id')
    if setor_destino_id:
        try:
            setor_destino = Setor.objects.get(pk=int(setor_destino_id))
        except (Setor.DoesNotExist, ValueError, TypeError):
            setor_destino = None

    codigo = (request.data.get('setor_destino') or '').strip().upper().replace(' ', '_')
    if not setor_destino:
        if not codigo:
            codigo = 'JURIDICO_2_RECURSOS'
        if codigo in {'JURIDICO_2_RECURSOS', 'JURIDICO_2', 'REC', 'RECURSOS'}:
            setor_destino = _obter_setor_juridico2()
        else:
            nome_base = SETOR_LABELS.get(codigo, codigo.replace('_', ' ').title())
            setor_destino = _buscar_setor_por_termos([codigo], [nome_base])
            if not setor_destino:
                setor_destino = Setor.objects.create(
                    nome=nome_base,
                    sigla=codigo[:10] if codigo else 'SETOR',
                    pode_protocolar=True,
                    pode_tramitar=True,
                )

    setor_origem = _obter_setor_fiscalizacao()
    protocolo, criado = _obter_protocolo_processo(processo, usuario, setor_origem)

    if protocolo.setor_atual_id == setor_destino.id:
        return Response(
            {
                'success': True,
                'message': 'Processo ja esta no setor destino.',
                'protocolo': protocolo.numero_protocolo,
                'setor_destino': setor_destino.nome,
            },
            status=status.HTTP_200_OK,
        )

    motivo = (request.data.get('motivo') or '').strip() or 'Encaminhamento do processo para analise.'
    observacoes = (request.data.get('observacoes') or '').strip()

    prazo_dias_raw = request.data.get('prazo_dias')
    prazo_dias = None
    if prazo_dias_raw not in (None, ''):
        try:
            prazo_dias = int(prazo_dias_raw)
        except (TypeError, ValueError):
            prazo_dias = None

    destinatario_id = (
        request.data.get('destinatario_id')
        or request.data.get('destinatario_direto')
        or request.data.get('responsavel_id')
    )
    destinatario = None
    if destinatario_id:
        try:
            from django.contrib.auth import get_user_model
            UserModel = get_user_model()
            destinatario = UserModel.objects.get(pk=int(destinatario_id))
        except Exception:
            destinatario = None

    protocolo_atualizado, tramitacao = workflow_service.tramitar_protocolo(
        protocolo=protocolo,
        setor_destino=setor_destino,
        motivo=motivo,
        usuario=usuario,
        observacoes=observacoes,
        prazo_dias=prazo_dias,
    )

    if destinatario:
        protocolo_atualizado.responsavel_atual = destinatario
        protocolo_atualizado.save(update_fields=['responsavel_atual', 'atualizado_em'])
        caixa_doc = (
            CaixaEntrada.objects.filter(protocolo=protocolo_atualizado)
            .order_by('-data_entrada')
            .first()
        )
        if caixa_doc:
            caixa_doc.destinatario_direto = destinatario
            caixa_doc.responsavel_atual = destinatario
            if not caixa_doc.setor_lotacao:
                caixa_doc.setor_lotacao = caixa_doc.setor_destino
            caixa_doc.save(
                update_fields=[
                    'destinatario_direto',
                    'responsavel_atual',
                    'setor_lotacao',
                    'data_atualizacao',
                ]
            )

    documento_data = None
    arquivo_upload = request.FILES.get('arquivo')
    if arquivo_upload:
        nome_arquivo = (getattr(arquivo_upload, 'name', '') or '').lower()
        if not any(nome_arquivo.endswith(ext) for ext in ['.doc', '.docx', '.pdf']):
            return Response(
                {'detail': 'Formato invalido. Envie DOC, DOCX ou PDF.'},
                status=status.HTTP_400_BAD_REQUEST
            )
        titulo = (request.data.get('titulo') or '').strip()
        if not titulo:
            titulo = f"Documento de tramitacao - {processo.numero_processo}"
        descricao_doc = (request.data.get('descricao') or '').strip()
        if observacoes and not descricao_doc:
            descricao_doc = observacoes
        documento = DocumentoProcesso.objects.create(
            processo=processo,
            tipo='outros',
            titulo=titulo,
            descricao=descricao_doc,
            arquivo=arquivo_upload,
            usuario_upload=_nome_usuario(request),
        )
        documento_data = DocumentoProcessoSerializer(documento, context={'request': request}).data

    usuario_nome = _nome_usuario(request)
    _registrar_evento_processo(
        processo,
        (
            f"Processo tramitado para {setor_destino.nome}."
            f"{' Destinatario: ' + (destinatario.get_full_name() or destinatario.username) + '.' if destinatario else ''} "
            f"Protocolo: {protocolo_atualizado.numero_protocolo}."
        ),
        usuario_nome,
    )

    return Response(
        {
            'success': True,
            'protocolo': protocolo_atualizado.numero_protocolo,
            'setor_destino': setor_destino.nome,
            'tramitacao_id': tramitacao.id,
            'protocolo_criado': criado,
            'documento': documento_data,
        },
        status=status.HTTP_200_OK,
    )


def _extrair_valor_total_multa_excel(conteudo_bytes: bytes):
    """
    Extrai o valor total da multa a partir do Excel de dosimetria.
    Estratégia:
    1) Procura o rótulo "VALOR TOTAL DA MULTA" e lê a célula numérica à direita.
    2) Fallback para células conhecidas (E69 / D63).
    """
    if not conteudo_bytes:
        return None, None, None

    wb = load_workbook(filename=BytesIO(conteudo_bytes), data_only=True)
    sheet_name = 'DOSIMETRIA' if 'DOSIMETRIA' in wb.sheetnames else wb.sheetnames[0]
    ws = wb[sheet_name]

    alvo = 'VALOR TOTAL DA MULTA'

    max_row = min(ws.max_row or 0, 400)
    max_col = min(ws.max_column or 0, 60)

    # 1) Busca pelo rótulo e pega o valor à direita
    for row in ws.iter_rows(min_row=1, max_row=max_row, min_col=1, max_col=max_col):
        for cell in row:
            v = cell.value
            if isinstance(v, str) and alvo in v.upper():
                for offset in range(1, 6):
                    candidato = ws.cell(row=cell.row, column=cell.column + offset)
                    if isinstance(candidato.value, (int, float)) and candidato.value > 0:
                        return Decimal(str(candidato.value)), sheet_name, candidato.coordinate

    # 2) Fallback para células já vistas nesses modelos
    for coord in ('E69', 'D63'):
        try:
            v = ws[coord].value
        except Exception:
            v = None
        if isinstance(v, (int, float)) and v > 0:
            return Decimal(str(v)), sheet_name, coord

    return None, sheet_name, None


@api_view(['POST'])
def registrar_dosimetria_excel(request, pk):
    """
    Importa o Excel de dosimetria (individual/coletiva), extrai o valor total
    e atualiza o valor da multa no processo.
    """
    processo = get_object_or_404(Processo, pk=pk)

    arquivo = request.FILES.get('arquivo') or request.FILES.get('dosimetria')
    if not arquivo:
        return Response(
            {'detail': 'Envie o arquivo Excel da dosimetria.'},
            status=status.HTTP_400_BAD_REQUEST,
        )

    nome_arquivo = getattr(arquivo, 'name', 'dosimetria.xlsx')
    if not nome_arquivo.lower().endswith('.xlsx'):
        return Response(
            {'detail': 'Use o arquivo em formato .xlsx (Excel).'},
            status=status.HTTP_400_BAD_REQUEST,
        )

    try:
        conteudo = arquivo.read()
    except Exception:
        return Response(
            {'detail': 'Nao foi possivel ler o arquivo enviado.'},
            status=status.HTTP_400_BAD_REQUEST,
        )

    valor_multa, sheet_name, coord = _extrair_valor_total_multa_excel(conteudo)
    if not valor_multa:
        return Response(
            {
                'detail': (
                    'Nao foi possivel identificar o valor total da multa no Excel. '
                    'Verifique se a planilha contem o campo "VALOR TOTAL DA MULTA" '
                    'e se o arquivo foi salvo com os calculos atualizados.'
                )
            },
            status=status.HTTP_400_BAD_REQUEST,
        )

    # Salva o Excel como documento do processo
    usuario_nome = _nome_usuario(request)
    titulo = f"Dosimetria (Excel) - {processo.numero_processo}"
    descricao = f"Importado da planilha {nome_arquivo} ({sheet_name}{'!' + coord if coord else ''})."

    documento = DocumentoProcesso.objects.create(
        processo=processo,
        tipo='outros',
        titulo=titulo,
        descricao=descricao,
        arquivo=ContentFile(conteudo, name=nome_arquivo),
        usuario_upload=usuario_nome,
    )

    # Atualiza valores do processo
    processo.valor_multa = valor_multa
    processo.valor_final = valor_multa
    processo.save()

    _registrar_evento_processo(
        processo,
        f"Dosimetria importada ({sheet_name}{'!' + coord if coord else ''}). Valor: R$ {valor_multa}.",
        usuario_nome,
    )

    return Response(
        {
            'success': True,
            'processo': processo.numero_processo,
            'valor_multa': str(valor_multa),
            'sheet': sheet_name,
            'celula': coord,
            'documento_id': documento.id,
        },
        status=status.HTTP_200_OK,
    )


@api_view(['GET'])
def historico_processo(request, pk):
    """
    Retorna histórico completo de mudanças do processo.
    """
    try:
        processo = get_object_or_404(Processo, pk=pk)
        historico = processo.historico.all().order_by('-data_mudanca')
        
        serializer = HistoricoProcessoSerializer(historico, many=True)
        return Response(serializer.data)
        
    except Exception as e:
        return Response({
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
def processos_dashboard(request):
    """
    Dashboard específico para processos.
    """
    try:
        hoje = timezone.now().date()
        
        dashboard_data = {
            'resumo_geral': {
                'total_processos': Processo.objects.count(),
                'processos_ativos': Processo.objects.exclude(
                    status__in=['finalizado_procedente', 'finalizado_improcedente', 'arquivado']
                ).count(),
                'processos_vencendo': Processo.objects.filter(
                    Q(prazo_defesa__lte=hoje + timedelta(days=3)) |
                    Q(prazo_recurso__lte=hoje + timedelta(days=3))
                ).count(),
                'processos_finalizados_mes': Processo.objects.filter(
                    data_finalizacao__month=hoje.month,
                    data_finalizacao__year=hoje.year
                ).count(),
            },
            'por_status': dict(
                Processo.objects.values('status').annotate(
                    count=Count('id')
                ).values_list('status', 'count')
            ),
            'por_prioridade': dict(
                Processo.objects.values('prioridade').annotate(
                    count=Count('id')
                ).values_list('prioridade', 'count')
            ),
            'valores_financeiros': {
                'total_multas': Processo.objects.aggregate(
                    total=Sum('valor_multa')
                )['total'] or 0,
                'valor_medio_multa': Processo.objects.aggregate(
                    media=Avg('valor_multa')
                )['media'] or 0,
                'total_valores_finais': Processo.objects.aggregate(
                    total=Sum('valor_final')
                )['total'] or 0,
            },
            'prazos_criticos': list(
                Processo.objects.filter(
                    Q(prazo_defesa__lte=hoje + timedelta(days=3)) |
                    Q(prazo_recurso__lte=hoje + timedelta(days=3))
                ).values(
                    'id', 'numero_processo', 'autuado', 'prazo_defesa', 'prazo_recurso', 'status'
                )
            ),
        }
        
        return Response(dashboard_data)
        
    except Exception as e:
        return Response({
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
def relatorio_mensal_processos(request):
    """
    Relatório mensal detalhado dos processos.
    """
    try:
        mes = int(request.GET.get('mes', timezone.now().month))
        ano = int(request.GET.get('ano', timezone.now().year))
        
        # Filtrar processos do mês
        processos_mes = Processo.objects.filter(
            criado_em__month=mes,
            criado_em__year=ano
        )
        
        finalizados_mes = Processo.objects.filter(
            data_finalizacao__month=mes,
            data_finalizacao__year=ano
        )
        
        relatorio = {
            'periodo': {
                'mes': mes,
                'ano': ano,
                'nome_mes': [
                    'Janeiro', 'Fevereiro', 'Março', 'Abril', 'Maio', 'Junho',
                    'Julho', 'Agosto', 'Setembro', 'Outubro', 'Novembro', 'Dezembro'
                ][mes - 1]
            },
            'criados_no_mes': {
                'total': processos_mes.count(),
                'por_prioridade': dict(
                    processos_mes.values('prioridade').annotate(
                        count=Count('id')
                    ).values_list('prioridade', 'count')
                ),
                'valor_total_multas': processos_mes.aggregate(
                    total=Sum('valor_multa')
                )['total'] or 0,
            },
            'finalizados_no_mes': {
                'total': finalizados_mes.count(),
                'por_status': dict(
                    finalizados_mes.values('status').annotate(
                        count=Count('id')
                    ).values_list('status', 'count')
                ),
                'tempo_medio_tramitacao': finalizados_mes.aggregate(
                    tempo_medio=Avg('tempo_tramitacao')
                )['tempo_medio'],
            },
            'detalhes': ProcessoResumoMensalSerializer(processos_mes, many=True).data
        }
        
        return Response(relatorio)
        
    except Exception as e:
        return Response({
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
def buscar_processos(request):
    """
    Busca avançada de processos.
    """
    try:
        serializer = ProcessoBuscaSerializer(data=request.GET)
        if not serializer.is_valid():
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        
        filtros = serializer.validated_data
        queryset = Processo.objects.all()
        
        # Aplicar filtros
        if filtros.get('termo'):
            termo = filtros['termo']
            queryset = queryset.filter(
                Q(numero_processo__icontains=termo) |
                Q(autuado__icontains=termo) |
                Q(cnpj__icontains=termo)
            )
        
        if filtros.get('status'):
            queryset = queryset.filter(status__in=filtros['status'])
        
        if filtros.get('prioridade'):
            queryset = queryset.filter(prioridade__in=filtros['prioridade'])
        
        if filtros.get('data_inicio'):
            queryset = queryset.filter(criado_em__gte=filtros['data_inicio'])
        
        if filtros.get('data_fim'):
            queryset = queryset.filter(criado_em__lte=filtros['data_fim'])
        
        if filtros.get('valor_min'):
            queryset = queryset.filter(valor_multa__gte=filtros['valor_min'])
        
        if filtros.get('valor_max'):
            queryset = queryset.filter(valor_multa__lte=filtros['valor_max'])
        
        # Ordenação
        ordem = filtros.get('ordem', '-criado_em')
        queryset = queryset.order_by(ordem)
        
        # Paginação
        paginator = PageNumberPagination()
        paginator.page_size = filtros.get('limite', 20)
        resultado_paginado = paginator.paginate_queryset(queryset, request)
        
        serializer_resultado = ProcessoSimpleSerializer(resultado_paginado, many=True)
        return paginator.get_paginated_response(serializer_resultado.data)
        
    except Exception as e:
        return Response({
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
def opcoes_filtros(request):
    """
    Retorna opções disponíveis para filtros de processos.
    """
    try:
        opcoes = {
            'status': [
                {'value': choice[0], 'label': choice[1]}
                for choice in Processo.STATUS_CHOICES
            ],
            'prioridade': [
                {'value': choice[0], 'label': choice[1]}
                for choice in Processo.PRIORIDADE_CHOICES
            ],
            'anos_disponiveis': list(
                Processo.objects.dates('criado_em', 'year').values_list(
                    'criado_em__year', flat=True
                )
            ),
        }
        
        return Response(opcoes)
        
    except Exception as e:
        return Response({
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


# ========================================
# VIEWS DE DOCUMENTOS DO PROCESSO
# ========================================

class DocumentoProcessoListCreateAPIView(generics.ListCreateAPIView):
    serializer_class = DocumentoProcessoSerializer
    
    def get_queryset(self):
        processo_id = self.kwargs['processo_id']
        return DocumentoProcesso.objects.filter(processo_id=processo_id)
    
    def perform_create(self, serializer):
        processo_id = self.kwargs['processo_id']
        processo = get_object_or_404(Processo, id=processo_id)
        serializer.save(processo=processo)


@api_view(['POST'])
def upload_documento_processo(request, processo_id):
    """
    Upload de documento para um processo específico.
    """
    try:
        processo = get_object_or_404(Processo, id=processo_id)
        
        serializer = DocumentoUploadSerializer(data=request.data, context={'request': request})
        if not serializer.is_valid():
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        
        documento = serializer.save(processo=processo)
        
        return Response(
            DocumentoProcessoSerializer(documento, context={'request': request}).data,
            status=status.HTTP_201_CREATED
        )
        
    except Exception as e:
        return Response({
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


def _formatar_data_extenso(data):
    meses = [
        'Janeiro', 'Fevereiro', 'Março', 'Abril', 'Maio', 'Junho',
        'Julho', 'Agosto', 'Setembro', 'Outubro', 'Novembro', 'Dezembro'
    ]
    return f"{data.day:02d} de {meses[data.month - 1]} de {data.year}"


def _normalizar_titulo(texto):
    return (
        texto.upper()
        .replace('–', '-')
        .replace('—', '-')
        .replace('−', '-')
        .strip()
    )


def _set_paragraph_text(paragraph, text, compact=False):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ''
    else:
        paragraph.add_run(text)

    if compact:
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph_format.line_spacing = 1.0


def _set_cell_text(cell, text):
    cell.text = ''
    if cell.paragraphs:
        _set_paragraph_text(cell.paragraphs[0], text)
    else:
        cell.add_paragraph(text)


def _converter_docx_para_pdf(docx_bytes, nome_base):
    if not docx_bytes:
        return None
    safe_name = (nome_base or 'Parecer').replace('/', '_')
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = Path(tmpdir) / f"{safe_name}.docx"
        pdf_path = Path(tmpdir) / f"{safe_name}.pdf"
        docx_path.write_bytes(docx_bytes)

        try:
            from docx2pdf import convert

            convert(str(docx_path), str(tmpdir))
            if pdf_path.exists():
                return pdf_path.read_bytes()
        except Exception:
            pass

        converter = shutil.which('soffice') or shutil.which('libreoffice')
        if converter:
            try:
                subprocess.run(
                    [converter, '--headless', '--convert-to', 'pdf', '--outdir', tmpdir, str(docx_path)],
                    check=True,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                )
                if pdf_path.exists():
                    return pdf_path.read_bytes()
            except Exception:
                return None
    return None


def _gerar_docx_parecer_processo(parecer):
    template_path = Path(settings.BASE_DIR) / 'fiscalizacao' / 'templates' / 'docs' / 'ParecerFiscalizacao.docx'
    doc = Document(template_path)

    # Atualiza tabela de cabeçalho (processo e interessada)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texto = cell.text.strip()
                texto_upper = texto.upper()
                if 'PROCESSO' in texto_upper:
                    _set_cell_text(cell, f"PROCESSO Nº {parecer.processo.numero_processo}")
                if 'INTERESSADA' in texto_upper:
                    _set_cell_text(cell, f"INTERESSADA: {parecer.processo.autuado}")

    idx_sintese = None
    idx_parecer = None
    idx_decisao = None
    idx_data = None

    for i, paragraph in enumerate(doc.paragraphs):
        texto = paragraph.text.strip()
        if not texto:
            continue
        titulo = _normalizar_titulo(texto)

        if titulo.startswith('PARECER'):
            _set_paragraph_text(paragraph, f"PARECER {parecer.numero_parecer} - FISCALIZAÇÃO")
        elif titulo.startswith('I - SÍNTESE FÁTICA') or titulo.startswith('I - SINTESE FATICA'):
            idx_sintese = i
        elif titulo.startswith('II - PARECER'):
            idx_parecer = i
        elif titulo.startswith('III - DECISÃO') or titulo.startswith('III - DECISAO'):
            idx_decisao = i
        elif texto.lower().startswith('manaus,'):
            idx_data = i

    def _preencher_secao(start_idx, end_idx, texto_secao):
        if start_idx is None:
            return
        content_idx = start_idx + 1
        if content_idx >= len(doc.paragraphs):
            return
        _set_paragraph_text(doc.paragraphs[content_idx], texto_secao or '', compact=True)
        for j in range(content_idx + 1, end_idx or len(doc.paragraphs)):
            _set_paragraph_text(doc.paragraphs[j], '', compact=True)

    _preencher_secao(idx_sintese, idx_parecer, parecer.sintese_fatica or '')
    _preencher_secao(idx_parecer, idx_decisao, parecer.parecer or '')
    _preencher_secao(idx_decisao, idx_data, parecer.decisao or '')

    if idx_data is not None:
        data_extenso = _formatar_data_extenso(parecer.data_emissao)
        _set_paragraph_text(doc.paragraphs[idx_data], f"Manaus, {data_extenso}.")

        nome = parecer.elaborado_por_nome
        if not nome and parecer.elaborado_por:
            nome = parecer.elaborado_por.get_full_name() or parecer.elaborado_por.username
        cargo = parecer.cargo_elaborador or ''

        idx_nome = None
        idx_cargo = None
        for i in range(idx_data + 1, len(doc.paragraphs)):
            if doc.paragraphs[i].text.strip():
                idx_nome = i
                break
        if idx_nome is not None:
            _set_paragraph_text(doc.paragraphs[idx_nome], nome or '')
            for i in range(idx_nome + 1, len(doc.paragraphs)):
                if doc.paragraphs[i].text.strip():
                    idx_cargo = i
                    break
            if idx_cargo is not None:
                _set_paragraph_text(doc.paragraphs[idx_cargo], cargo)
            else:
                doc.add_paragraph(cargo)
        else:
            doc.add_paragraph(nome or '')
            if cargo:
                doc.add_paragraph(cargo)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _salvar_documento_parecer(parecer, usuario=None):
    try:
        docx_bytes = _gerar_docx_parecer_processo(parecer)
    except Exception:
        logger.exception("Falha ao gerar DOCX do parecer %s", parecer.id)
        return None

    if not docx_bytes:
        return None

    nome_usuario = ''
    if usuario and getattr(usuario, 'is_authenticated', False):
        nome_usuario = usuario.get_full_name() or usuario.username

    nome_arquivo = f"Parecer_{parecer.numero_parecer.replace('/', '_')}.docx"
    arquivo = ContentFile(docx_bytes, name=nome_arquivo)

    documento, _ = DocumentoProcesso.objects.update_or_create(
        processo=parecer.processo,
        tipo='parecer',
        titulo=f"Parecer {parecer.numero_parecer}",
        defaults={
            'descricao': 'Parecer técnico do processo',
            'arquivo': arquivo,
            'usuario_upload': nome_usuario,
        },
    )
    return documento


class ParecerProcessoListCreateAPIView(generics.ListCreateAPIView):
    serializer_class = ParecerProcessoSerializer

    def get_queryset(self):
        processo_id = self.kwargs['processo_id']
        return ParecerProcesso.objects.filter(processo_id=processo_id).order_by('-criado_em')

    def perform_create(self, serializer):
        processo_id = self.kwargs['processo_id']
        processo = get_object_or_404(Processo, id=processo_id)
        elaborador = self.request.user if self.request.user.is_authenticated else None
        parecer = serializer.save(processo=processo, elaborado_por=elaborador)
        _salvar_documento_parecer(parecer, usuario=elaborador)


class ParecerProcessoDetailAPIView(generics.RetrieveUpdateDestroyAPIView):
    queryset = ParecerProcesso.objects.all()
    serializer_class = ParecerProcessoSerializer

    def perform_update(self, serializer):
        parecer = serializer.save()
        usuario = self.request.user if self.request.user.is_authenticated else None
        _salvar_documento_parecer(parecer, usuario=usuario)


@api_view(['GET'])
def gerar_parecer_processo_docx(request, pk):
    parecer = get_object_or_404(ParecerProcesso, pk=pk)
    try:
        docx_bytes = _gerar_docx_parecer_processo(parecer)
    except Exception as exc:
        return Response({'error': str(exc)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    response = HttpResponse(
        docx_bytes,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    filename = f"Parecer_{parecer.numero_parecer.replace('/', '_')}.docx"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


@api_view(['GET'])
def gerar_parecer_processo_pdf(request, pk):
    parecer = get_object_or_404(ParecerProcesso, pk=pk)
    try:
        docx_bytes = _gerar_docx_parecer_processo(parecer)
        pdf_bytes = _converter_docx_para_pdf(docx_bytes, f"Parecer_{parecer.numero_parecer}")
        if not pdf_bytes:
            return Response({'error': 'Conversão para PDF indisponível.'}, status=status.HTTP_400_BAD_REQUEST)
    except Exception as exc:
        return Response({'error': str(exc)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    response = HttpResponse(pdf_bytes, content_type='application/pdf')
    filename = f"Parecer_{parecer.numero_parecer.replace('/', '_')}.pdf"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


# ========================================
# VIEWS DE VALIDAÇÃO E UTILITÁRIOS
# ========================================

@api_view(['POST'])
def validar_numero_processo(request):
    """
    Valida se um número de processo já existe.
    """
    try:
        numero = request.data.get('numero_processo')
        
        if not numero:
            return Response({
                'valido': False,
                'erro': 'Número não fornecido'
            })
        
        existe = Processo.objects.filter(numero_processo=numero).exists()
        
        return Response({
            'valido': not existe,
            'numero': numero,
            'existe': existe
        })
        
    except Exception as e:
        return Response({
            'error': str(e)
        }, status=500)


@api_view(['GET'])
def sugerir_numero_processo(request):
    """
    Sugere próximo número de processo disponível.
    """
    try:
        from ..utils import obter_proximo_numero_processo_sei_preview
        numero_sugerido = obter_proximo_numero_processo_sei_preview()
        
        return Response({
            'numero_sugerido': numero_sugerido,
            'ano': timezone.now().year
        })
        
    except Exception as e:
        return Response({
            'error': str(e)
        }, status=500)


# ========================================
# VIEWS EXTRAS E RELATÓRIOS
# ========================================

@api_view(['GET'])
def exportar_processos(request):
    """
    Exporta lista de processos em formato CSV.
    """
    # Implementação de exportação seria aqui
    pass


@api_view(['GET'])
def estatisticas_avancadas(request):
    """
    Estatísticas avançadas dos processos.
    """
    # Implementação de estatísticas avançadas seria aqui
    pass


@api_view(['POST'])
def operacoes_lote(request):
    """
    Operações em lote para múltiplos processos.
    """
    # Implementação de operações em lote seria aqui
    pass


# ========================================
# VIEWS UNIFICADAS E PERFORMANCE
# ========================================

@api_view(['GET'])
def buscar_processo_unificado(request, processo_id):
    """
    Busca unificada que retorna processo com todos os dados relacionados.
    """
    try:
        processo = get_object_or_404(Processo, id=processo_id)
        
        dados = {
            'processo': ProcessoDetailSerializer(processo).data,
            'historico': HistoricoProcessoSerializer(
                processo.historico.all()[:10], many=True
            ).data,
            'documentos': DocumentoProcessoSerializer(
                processo.documentos.all(), many=True
            ).data,
            'auto_infracao': None
        }
        
        # Incluir dados da infração se existir
        if hasattr(processo, 'auto_infracao'):
            from ..serializers import AutoInfracaoSimpleSerializer
            dados['auto_infracao'] = AutoInfracaoSimpleSerializer(
                processo.auto_infracao
            ).data
        
        return Response(dados)
        
    except Exception as e:
        return Response({
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
def listar_todos_processos(request):
    """
    Lista otimizada de todos os processos com informações essenciais.
    """
    try:
        # Usar select_related para otimizar queries
        queryset = Processo.objects.select_related('auto_infracao').all()
        
        # Aplicar filtros básicos
        status_filter = request.GET.get('status')
        if status_filter:
            queryset = queryset.filter(status=status_filter)
        
        limite = int(request.GET.get('limite', 50))
        offset = int(request.GET.get('offset', 0))
        
        processos = queryset[offset:offset + limite]
        
        # Serializar com dados essenciais
        dados = ProcessoSimpleSerializer(processos, many=True).data
        
        return Response({
            'processos': dados,
            'total': queryset.count(),
            'offset': offset,
            'limite': limite
        })
        
    except Exception as e:
        return Response({
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
def teste_performance_busca(request, processo_id):
    """
    Endpoint para testar performance de busca.
    """
    # Implementação de teste de performance seria aqui
    pass


@api_view(['GET'])
def busca_avancada_processos(request):
    """
    Busca avançada de processos com múltiplos filtros.
    """
    try:
        queryset = Processo.objects.select_related('auto_infracao').all()
        
        # Filtros de busca
        q = request.GET.get('q', '')
        if q:
            queryset = queryset.filter(
                Q(numero_processo__icontains=q) |
                Q(autuado__icontains=q) |
                Q(cnpj__icontains=q) |
                Q(observacoes__icontains=q)
            )
        
        # Filtros adicionais
        status = request.GET.get('status')
        if status:
            queryset = queryset.filter(status=status)
        
        prioridade = request.GET.get('prioridade')
        if prioridade:
            queryset = queryset.filter(prioridade=prioridade)
        
        # Limite de resultados
        limit = int(request.GET.get('limit', 50))
        offset = int(request.GET.get('offset', 0))
        
        total = queryset.count()
        processos = queryset[offset:offset + limit]
        
        dados = ProcessoSimpleSerializer(processos, many=True).data
        
        return Response({
            'resultados': dados,
            'total_encontrados': total,
            'offset': offset,
            'limite': limit
        })
        
    except Exception as e:
        return Response({
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
def processos_dashboard_cached(request):
    """
    Dashboard de processos com cache para melhor performance.
    """
    try:
        # Cache key baseada nos parâmetros da requisição
        cache_key = f"dashboard_processos_{request.GET.get('periodo', 'mes')}"
        
        # Tentar buscar do cache primeiro
        cached_data = cache.get(cache_key)
        if cached_data:
            return Response(cached_data)
        
        # Se não estiver em cache, calcular dados
        hoje = timezone.now().date()
        inicio_mes = hoje.replace(day=1)
        
        # Estatísticas básicas
        total_processos = Processo.objects.count()
        processos_abertos = Processo.objects.filter(status='aberto').count()
        processos_vencidos = Processo.objects.filter(
            Q(prazo_defesa__lt=hoje) | Q(prazo_recurso__lt=hoje)
        ).count()
        
        # Processos próximos do vencimento (próximos 3 dias)
        limite_vencimento = hoje + timedelta(days=3)
        processos_proximos_vencimento = Processo.objects.filter(
            Q(prazo_defesa__lte=limite_vencimento) | Q(prazo_recurso__lte=limite_vencimento)
        ).count()
        
        # Valor total em tramitação
        valor_total = Processo.objects.aggregate(
            total=Sum('valor_multa')
        )['total'] or 0
        
        dados = {
            'resumo': {
                'total_processos': total_processos,
                'processos_abertos': processos_abertos,
                'processos_vencidos': processos_vencidos,
                'processos_proximos_vencimento': processos_proximos_vencimento,
                'valor_total_tramitacao': float(valor_total),
                'tempo_medio_tramitacao': 0  # Seria calculado se necessário
            },
            'periodo': request.GET.get('periodo', 'mes'),
            'atualizado_em': timezone.now().isoformat()
        }
        
        # Salvar no cache por 5 minutos
        cache.set(cache_key, dados, 300)
        
        return Response(dados)
        
    except Exception as e:
        return Response({
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
