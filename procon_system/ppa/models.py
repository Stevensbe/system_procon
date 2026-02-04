from django.db import models
from django.contrib.auth.models import User
from django.contrib.contenttypes.fields import GenericForeignKey
from django.contrib.contenttypes.models import ContentType
from django.utils import timezone
import uuid
import datetime


class ProcedimentoPreAdministrativo(models.Model):
    """
    PPA - Procedimento Preliminar Administrativo
    Capa do processo onde se anexa todos os documentos (AC, AI, NOT, Defesas, etc)
    """
    
    STATUS_CHOICES = [
        ('criado', 'Criado'),
        ('em_analise', 'Em Análise'),
        ('notificado', 'Notificado'),
        ('aguardando_resposta', 'Aguardando Resposta'),
        ('com_defesa', 'Com Defesa'),
        ('parecer_elaborado', 'Parecer Elaborado'),
        ('concluido', 'Concluído'),
        ('arquivado', 'Arquivado'),
    ]
    
    DECISAO_CHOICES = [
        ('pendente', 'Pendente'),
        ('arquivado', 'Arquivado - Sem Fundamento'),
        ('auto_criado', 'Auto de Infração Criado'),
        ('encaminhado', 'Encaminhado para Outro Órgão'),
    ]
    
    SIGLA_CHOICES = [
        ('BANCO', 'Banco'),
        ('POSTO', 'Posto de Combustível'),
        ('SUPERMERCADO', 'Supermercado'),
        ('DIVERSOS', 'Diversos'),
        ('TELECOMUNICACOES', 'Telecomunicações'),
        ('ENERGIA', 'Energia'),
        ('PLANO_SAUDE', 'Plano de Saúde'),
        ('OUTROS', 'Outros'),
    ]
    
    # === IDENTIFICAÇÃO ===
    uuid = models.UUIDField("UUID", default=uuid.uuid4, editable=False, unique=True)
    numero = models.CharField("Número do PPA", max_length=20, unique=True, blank=True)
    # Ex: 001/2026
    
    # === CLASSIFICAÇÃO ===
    sigla = models.CharField("Sigla/Tipo", max_length=50, choices=SIGLA_CHOICES)
    assunto = models.TextField("Assunto")
    interessado = models.CharField("Interessado (Empresa)", max_length=255)
    cnpj_interessado = models.CharField("CNPJ", max_length=18, blank=True)
    endereco_interessado = models.TextField("Endereço", blank=True)
    
    # === RESPONSÁVEIS ===
    analista_responsavel = models.ForeignKey(
        User,
        on_delete=models.PROTECT,
        related_name='ppas_responsavel',
        verbose_name="Analista Responsável"
    )
    supervisor = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='ppas_supervisionados',
        verbose_name="Supervisor"
    )
    
    # === STATUS E DECISÃO ===
    status = models.CharField("Status", max_length=30, choices=STATUS_CHOICES, default='criado')
    decisao_final = models.CharField(
        "Decisão Final",
        max_length=20,
        choices=DECISAO_CHOICES,
        default='pendente'
    )
    
    # === PRAZOS ===
    prazo_analise = models.DateField("Prazo para Análise", null=True, blank=True)
    prazo_resposta = models.DateField("Prazo para Resposta Empresa", null=True, blank=True)
    data_conclusao = models.DateTimeField("Data de Conclusão", null=True, blank=True)
    
    # === OBSERVAÇÕES ===
    observacoes = models.TextField("Observações Gerais", blank=True)
    observacoes_internas = models.TextField("Observações Internas", blank=True)
    fundamentacao_decisao = models.TextField("Fundamentação da Decisão", blank=True)
    
    # === CONTROLE ===
    criado_em = models.DateTimeField("Criado em", auto_now_add=True)
    atualizado_em = models.DateTimeField("Atualizado em", auto_now=True)
    criado_por = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        related_name='ppas_criados',
        verbose_name="Criado por"
    )
    
    class Meta:
        verbose_name = "PPA - Procedimento Pré-Administrativo"
        verbose_name_plural = "PPAs - Procedimentos Pré-Administrativos"
        ordering = ['-criado_em']
        indexes = [
            models.Index(fields=['numero']),
            models.Index(fields=['status']),
            models.Index(fields=['analista_responsavel']),
            models.Index(fields=['sigla']),
            models.Index(fields=['criado_em']),
        ]
    
    def __str__(self):
        return f"{self.numero} - {self.interessado}"
    
    def save(self, *args, **kwargs):
        """Gera número automático do PPA"""
        if not self.numero:
            self.numero = self._gerar_numero_ppa()
        super().save(*args, **kwargs)
    
    def _gerar_numero_ppa(self):
        """Gera número sequencial para o PPA (ex: 001/2026)"""
        ano = timezone.now().year
        ultimo = ProcedimentoPreAdministrativo.objects.filter(
            numero__endswith=f'/{ano}'
        ).order_by('-id').first()

        seq = 1
        if ultimo:
            try:
                bruto = str(ultimo.numero).replace('PPA-', '').replace('PPA', '').strip()
                seq = int(bruto.split('/')[0]) + 1
            except (ValueError, IndexError):
                seq = 1

        return f"{seq:03d}/{ano}"
    
    @property
    def total_anexos(self):
        """Retorna total de anexos"""
        return self.anexos.count()
    
    @property
    def total_movimentacoes(self):
        """Retorna total de movimentações"""
        return self.movimentacoes.count()
    
    @property
    def esta_no_prazo(self):
        """Verifica se está dentro do prazo"""
        if not self.prazo_analise:
            return True
        return timezone.now().date() <= self.prazo_analise
    
    @property
    def dias_ate_prazo(self):
        """Dias restantes até o prazo"""
        if not self.prazo_analise:
            return None
        delta = self.prazo_analise - timezone.now().date()
        return delta.days if delta.days >= 0 else 0


class MovimentacaoPPA(models.Model):
    """
    Tabela de Movimentação do PPA
    Registra TODOS os eventos que acontecem no processo
    """
    
    TIPO_MOVIMENTACAO_CHOICES = [
        ('criacao', 'Criação do PPA'),
        ('anexo_ac', 'Anexo de Auto de Constatação'),
        ('anexo_ai', 'Anexo de Auto de Infração'),
        ('anexo_notificacao', 'Anexo de Notificação'),
        ('anexo_defesa', 'Anexo de Defesa'),
        ('anexo_parecer', 'Anexo de Parecer'),
        ('anexo_documento', 'Anexo de Documento'),
        ('mudanca_status', 'Mudança de Status'),
        ('observacao', 'Observação'),
        ('analise', 'Análise'),
        ('decisao', 'Decisão'),
        ('outros', 'Outros'),
    ]
    
    ppa = models.ForeignKey(
        ProcedimentoPreAdministrativo,
        on_delete=models.CASCADE,
        related_name='movimentacoes',
        verbose_name="PPA"
    )
    
    # === DADOS DA MOVIMENTAÇÃO ===
    data = models.DateField("Data", default=timezone.localdate)
    hora = models.TimeField("Hora", null=True, blank=True)
    tipo_movimentacao = models.CharField(
        "Tipo de Movimentação",
        max_length=30,
        choices=TIPO_MOVIMENTACAO_CHOICES,
        default='outros'
    )
    atendimento = models.TextField("Atendimento/Descrição")
    # Ex: "NOT 047/2024", "AC 367/2024", "Img pl defesa", "Análise"
    
    # === USUÁRIO ===
    usuario = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        verbose_name="Usuário"
    )
    
    # === CONTROLE ===
    criado_em = models.DateTimeField("Criado em", auto_now_add=True)
    
    class Meta:
        verbose_name = "Movimentação de PPA"
        verbose_name_plural = "Movimentações de PPA"
        ordering = ['data', 'hora', '-criado_em']
        indexes = [
            models.Index(fields=['ppa', 'data']),
            models.Index(fields=['tipo_movimentacao']),
        ]
    
    def __str__(self):
        return f"{self.ppa.numero} - {self.data} - {self.atendimento[:50]}"
    
    def save(self, *args, **kwargs):
        """Normaliza data/hora para evitar inconsistências"""
        if isinstance(self.data, datetime.datetime):
            # Garante que o campo DateField receba apenas date
            self.data = timezone.localtime(self.data).date()
        elif self.data is None:
            self.data = timezone.localdate()

        if not self.hora:
            self.hora = timezone.localtime().time()

        super().save(*args, **kwargs)


class AnexoPPA(models.Model):
    """
    Documentos anexados ao PPA
    Pode vincular AC, AI, Notificações, Defesas, Pareceres ou arquivos
    """
    
    TIPO_DOCUMENTO_CHOICES = [
        ('AC', 'Auto de Constatação'),
        ('AI', 'Auto de Infração'),
        ('NOT', 'Notificação'),
        ('DEFESA', 'Defesa'),
        ('PARECER', 'Parecer'),
        ('RESPOSTA', 'Resposta da Empresa'),
        ('DOCUMENTO', 'Documento Complementar'),
        ('IMAGEM', 'Imagem/Foto'),
        ('COMPROVANTE', 'Comprovante'),
        ('OUTROS', 'Outros'),
    ]
    
    ppa = models.ForeignKey(
        ProcedimentoPreAdministrativo,
        on_delete=models.CASCADE,
        related_name='anexos',
        verbose_name="PPA"
    )
    
    # === TIPO ===
    tipo_documento = models.CharField(
        "Tipo de Documento",
        max_length=20,
        choices=TIPO_DOCUMENTO_CHOICES
    )
    
    # === REFERÊNCIA GENÉRICA (para vincular AC, AI, etc) ===
    content_type = models.ForeignKey(
        ContentType,
        on_delete=models.CASCADE,
        null=True,
        blank=True
    )
    object_id = models.PositiveIntegerField(null=True, blank=True)
    documento_relacionado = GenericForeignKey('content_type', 'object_id')
    
    # === ARQUIVO FÍSICO ===
    arquivo = models.FileField(
        "Arquivo",
        upload_to='ppa/anexos/%Y/%m/',
        null=True,
        blank=True
    )
    nome_arquivo_original = models.CharField("Nome Original", max_length=255, blank=True)
    
    # === DESCRIÇÃO ===
    descricao = models.TextField("Descrição", blank=True)
    numero_documento = models.CharField("Número do Documento", max_length=50, blank=True)
    # Ex: "NOT 047/2024", "AC 367/2024"
    
    # === CONTROLE ===
    data_anexacao = models.DateTimeField("Data de Anexação", auto_now_add=True)
    anexado_por = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        verbose_name="Anexado por"
    )
    
    class Meta:
        verbose_name = "Anexo de PPA"
        verbose_name_plural = "Anexos de PPA"
        ordering = ['-data_anexacao']
        indexes = [
            models.Index(fields=['ppa', 'tipo_documento']),
            models.Index(fields=['data_anexacao']),
        ]
    
    def __str__(self):
        return f"{self.ppa.numero} - {self.get_tipo_documento_display()} - {self.numero_documento}"
    
    @property
    def tamanho_arquivo(self):
        """Retorna tamanho do arquivo em formato legível"""
        if not self.arquivo:
            return None
        
        tamanho = self.arquivo.size
        if tamanho < 1024:
            return f"{tamanho} B"
        elif tamanho < 1024**2:
            return f"{tamanho/1024:.1f} KB"
        elif tamanho < 1024**3:
            return f"{tamanho/(1024**2):.1f} MB"
        else:
            return f"{tamanho/(1024**3):.1f} GB"


class ParecerPPA(models.Model):
    """
    Parecer técnico elaborado durante o PPA
    """
    
    CONCLUSAO_CHOICES = [
        ('procedente', 'Procedente - Criar AI'),
        ('improcedente', 'Improcedente - Arquivar'),
        ('mais_informacoes', 'Necessita Mais Informações'),
        ('encaminhar', 'Encaminhar para Outro Órgão'),
    ]
    
    ppa = models.ForeignKey(
        ProcedimentoPreAdministrativo,
        on_delete=models.CASCADE,
        related_name='pareceres',
        verbose_name="PPA"
    )
    
    # === PARECER ===
    numero_parecer = models.CharField("Número do Parecer", max_length=50, blank=True)
    titulo = models.CharField("Título", max_length=300)
    
    # Estrutura do parecer conforme modelo institucional
    sintese_fatica = models.TextField(
        "I - Síntese Fática",
        help_text="Relato cronológico completo dos fatos do caso",
        blank=True
    )
    relatorio = models.TextField(
        "II - Parecer",
        help_text="Análise jurídica e técnica do caso"
    )
    fundamentacao = models.TextField("Fundamentação Legal")
    conclusao = models.CharField(
        "III - Decisão",
        max_length=30,
        choices=CONCLUSAO_CHOICES
    )
    recomendacoes = models.TextField("Recomendações", blank=True)
    
    # === RESPONSÁVEL ===
    elaborado_por = models.ForeignKey(
        User,
        on_delete=models.PROTECT,
        related_name='pareceres_ppa_elaborados',
        verbose_name="Elaborado por"
    )
    cargo_elaborador = models.CharField("Cargo", max_length=100, blank=True)
    
    # === APROVAÇÃO ===
    aprovado_por = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='pareceres_ppa_aprovados',
        verbose_name="Aprovado por"
    )
    data_aprovacao = models.DateTimeField("Data de Aprovação", null=True, blank=True)
    
    # === CONTROLE ===
    criado_em = models.DateTimeField("Criado em", auto_now_add=True)
    atualizado_em = models.DateTimeField("Atualizado em", auto_now=True)
    
    class Meta:
        verbose_name = "Parecer de PPA"
        verbose_name_plural = "Pareceres de PPA"
        ordering = ['-criado_em']
    
    def __str__(self):
        return f"Parecer {self.numero_parecer} - {self.ppa.numero}"
    
    def save(self, *args, **kwargs):
        """Gera número automático do parecer"""
        if not self.numero_parecer:
            self.numero_parecer = self._gerar_numero_parecer()
        super().save(*args, **kwargs)
    
    def _gerar_numero_parecer(self):
        """Gera número sequencial para o parecer no formato: PARECER 001/2025"""
        ano = timezone.now().year
        ultimo = ParecerPPA.objects.filter(
            numero_parecer__endswith=f'/{ano}'
        ).order_by('-id').first()
        
        seq = 1
        if ultimo:
            try:
                # Suporta ambos os formatos: "PARECER 001/2025" e "PAR-00001/2025"
                numero_limpo = ultimo.numero_parecer.strip()
                if numero_limpo.startswith('PARECER'):
                    # Formato novo: "PARECER 001/2025"
                    seq = int(numero_limpo.split()[1].split('/')[0]) + 1
                elif numero_limpo.startswith('PAR-'):
                    # Formato antigo: "PAR-00001/2025"
                    seq = int(numero_limpo.split('/')[0].split('-')[1]) + 1
                else:
                    # Tentar extrair número de qualquer formato
                    partes = numero_limpo.split('/')
                    if len(partes) > 0:
                        seq = int(''.join(filter(str.isdigit, partes[0]))) + 1
            except (ValueError, IndexError):
                seq = 1
        
        return f"PARECER {seq:03d}/{ano}"
    
    def gerar_documento_word(self):
        """
        Gera documento Word no formato institucional do parecer
        
        Returns:
            Document: Objeto Document do python-docx
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from datetime import datetime
            import locale
            
            # Tentar configurar locale para português
            try:
                locale.setlocale(locale.LC_TIME, 'pt_BR.UTF-8')
            except:
                try:
                    locale.setlocale(locale.LC_TIME, 'Portuguese_Brazil.1252')
                except:
                    pass
            
            doc = Document()
            
            # Configurar margens
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(0.7)
                section.right_margin = Inches(0.7)
            
            # Cabeçalho
            heading = doc.add_heading(f'PARECER {self.numero_parecer} - FISCALIZAÇÃO', 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()  # Linha em branco
            
            # I - SÍNTESE FÁTICA
            if self.sintese_fatica:
                doc.add_heading('I - SÍNTESE FÁTICA', level=1)
                para = doc.add_paragraph(self.sintese_fatica)
                para_format = para.paragraph_format
                para_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para_format.first_line_indent = Inches(0.5)
                doc.add_paragraph()  # Linha em branco
            
            # II - PARECER
            doc.add_heading('II - PARECER', level=1)
            para = doc.add_paragraph(self.relatorio)
            para_format = para.paragraph_format
            para_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para_format.first_line_indent = Inches(0.5)
            
            # Adicionar fundamentação se houver
            if self.fundamentacao:
                doc.add_paragraph()  # Linha em branco
                para = doc.add_paragraph(self.fundamentacao)
                para_format = para.paragraph_format
                para_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para_format.first_line_indent = Inches(0.5)
            
            doc.add_paragraph()  # Linha em branco
            
            # III - DECISÃO
            doc.add_heading('III - DECISÃO', level=1)
            
            # Adicionar conclusão
            conclusao_texto = self.get_conclusao_display()
            para = doc.add_paragraph(conclusao_texto)
            para_format = para.paragraph_format
            para_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para_format.first_line_indent = Inches(0.5)
            
            # Adicionar recomendações se houver
            if self.recomendacoes:
                doc.add_paragraph()  # Linha em branco
                para = doc.add_paragraph(self.recomendacoes)
                para_format = para.paragraph_format
                para_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para_format.first_line_indent = Inches(0.5)
            
            doc.add_paragraph()  # Linha em branco
            doc.add_paragraph()  # Linha em branco
            
            # Data e assinatura
            data_atual = timezone.now()
            try:
                meses = {
                    1: 'Janeiro', 2: 'Fevereiro', 3: 'Março', 4: 'Abril',
                    5: 'Maio', 6: 'Junho', 7: 'Julho', 8: 'Agosto',
                    9: 'Setembro', 10: 'Outubro', 11: 'Novembro', 12: 'Dezembro'
                }
                data_formatada = f"Manaus, {data_atual.day:02d} de {meses[data_atual.month]} de {data_atual.year}."
            except:
                data_formatada = f"Manaus, {data_atual.strftime('%d de %B de %Y')}."
            
            doc.add_paragraph(data_formatada)
            doc.add_paragraph()  # Linha em branco
            doc.add_paragraph()  # Linha em branco
            
            # Assinatura
            if self.elaborado_por:
                nome_completo = self.elaborado_por.get_full_name() or self.elaborado_por.username
                doc.add_paragraph(nome_completo)
            
            if self.cargo_elaborador:
                doc.add_paragraph(self.cargo_elaborador)
            
            return doc
            
        except ImportError:
            raise ImportError(
                "python-docx não está instalado. "
                "Instale com: pip install python-docx"
            )
        except Exception as e:
            raise Exception(f"Erro ao gerar documento Word: {str(e)}")
    
    def salvar_documento_word(self, caminho_arquivo=None):
        """
        Salva o documento Word em um arquivo
        
        Args:
            caminho_arquivo: Caminho completo do arquivo. Se None, gera automaticamente.
        
        Returns:
            str: Caminho do arquivo salvo
        """
        import os
        from django.conf import settings
        
        doc = self.gerar_documento_word()
        
        if not caminho_arquivo:
            # Gerar caminho automático
            nome_arquivo = f"Parecer_{self.numero_parecer.replace('/', '_')}.docx"
            diretorio = os.path.join(settings.MEDIA_ROOT, 'ppa', 'pareceres', str(timezone.now().year))
            os.makedirs(diretorio, exist_ok=True)
            caminho_arquivo = os.path.join(diretorio, nome_arquivo)
        
        doc.save(caminho_arquivo)
        return caminho_arquivo

