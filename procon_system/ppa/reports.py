"""
Gerador de Relatórios em PDF para o módulo PPA
"""
from io import BytesIO
import logging
import shutil
import subprocess
import tempfile
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, Image, Frame, PageTemplate
)
from reportlab.lib.colors import HexColor
from django.utils import timezone
from django.conf import settings
from pathlib import Path
from docx import Document
from .models import ProcedimentoPreAdministrativo

logger = logging.getLogger(__name__)


class PPAReportGenerator:
    """Gerador de relatórios PDF para PPAs"""
    
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()
    
    def _setup_custom_styles(self):
        """Define estilos personalizados"""
        # Título principal
        self.styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=16,
            textColor=HexColor('#2c3e50'),
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        ))
        
        # Subtítulo
        self.styles.add(ParagraphStyle(
            name='CustomHeading',
            parent=self.styles['Heading2'],
            fontSize=14,
            textColor=HexColor('#34495e'),
            spaceAfter=12,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        ))
        
        # Texto normal
        self.styles.add(ParagraphStyle(
            name='CustomBody',
            parent=self.styles['BodyText'],
            fontSize=10,
            alignment=TA_JUSTIFY,
            spaceAfter=6
        ))
        
        # Label
        self.styles.add(ParagraphStyle(
            name='Label',
            parent=self.styles['BodyText'],
            fontSize=9,
            textColor=HexColor('#7f8c8d'),
            fontName='Helvetica-Bold'
        ))
    
    def gerar_pdf_completo(self, ppa_id):
        """
        Gera PDF completo do PPA com todas as informações
        
        Args:
            ppa_id: ID do PPA
            
        Returns:
            BytesIO com o PDF gerado
        """
        try:
            ppa = ProcedimentoPreAdministrativo.objects.select_related(
                'analista_responsavel',
                'supervisor',
            ).prefetch_related(
                'movimentacoes',
                'anexos',
                'pareceres'
            ).get(id=ppa_id)
            
            buffer = BytesIO()
            doc = SimpleDocTemplate(
                buffer,
                pagesize=A4,
                rightMargin=inch*0.7,
                leftMargin=inch*0.7,
                topMargin=inch,
                bottomMargin=inch*0.7
            )
            
            story = []
            
            # Cabeçalho
            story.extend(self._criar_cabecalho(ppa))
            story.append(Spacer(1, 0.3*inch))
            
            # Dados do PPA
            story.extend(self._criar_dados_ppa(ppa))
            story.append(Spacer(1, 0.2*inch))
            
            # Tabela de Movimentações
            story.extend(self._criar_tabela_movimentacoes(ppa))
            story.append(Spacer(1, 0.2*inch))
            
            # Anexos
            if ppa.anexos.exists():
                story.extend(self._criar_lista_anexos(ppa))
                story.append(Spacer(1, 0.2*inch))
            
            # Pareceres
            if ppa.pareceres.exists():
                story.extend(self._criar_pareceres(ppa))
                story.append(Spacer(1, 0.2*inch))
            
            # Decisão Final
            if ppa.decisao_final != 'pendente':
                story.extend(self._criar_decisao_final(ppa))
            
            # Rodapé
            story.extend(self._criar_rodape())
            
            # Gerar PDF
            doc.build(story)
            buffer.seek(0)
            return buffer
            
        except ProcedimentoPreAdministrativo.DoesNotExist:
            raise ValueError(f'PPA {ppa_id} não encontrado')

    def gerar_pdf_capa(self, ppa_id):
        """
        Gera PDF de capa do PPA no layout simplificado.

        Args:
            ppa_id: ID do PPA
        Returns:
            BytesIO com o PDF gerado
        """
        try:
            ppa = ProcedimentoPreAdministrativo.objects.select_related(
                'analista_responsavel',
                'supervisor',
            ).prefetch_related(
                'movimentacoes',
                'anexos',
                'pareceres'
            ).get(id=ppa_id)

            buffer = BytesIO()
            doc = SimpleDocTemplate(
                buffer,
                pagesize=A4,
                rightMargin=inch*0.7,
                leftMargin=inch*0.7,
                topMargin=inch,
                bottomMargin=inch*0.7
            )

            story = []
            story.extend(self._criar_cabecalho_capa(ppa))
            story.append(Spacer(1, 0.2*inch))
            story.extend(self._criar_campos_capa(ppa))
            story.append(Spacer(1, 0.2*inch))
            story.extend(self._criar_tabela_movimento_capa(ppa))

            doc.build(story)
            buffer.seek(0)
            return buffer
        except ProcedimentoPreAdministrativo.DoesNotExist:
            raise ValueError(f'PPA {ppa_id} não encontrado')
    
    def _criar_cabecalho(self, ppa):
        """Cria o cabeçalho do relatório"""
        elements = []
        
        # Logo/Título do órgão
        elements.append(Paragraph(
            "INSTITUTO DE DEFESA DO CONSUMIDOR - PROCON-AM",
            self.styles['CustomTitle']
        ))
        
        elements.append(Paragraph(
            "DEPARTAMENTO DE FISCALIZAÇÃO",
            self.styles['Normal']
        ))
        
        elements.append(Spacer(1, 0.2*inch))
        
        # Número do PPA
        elements.append(Paragraph(
            f"<b>PROCEDIMENTO PRELIMINAR ADMINISTRATIVO</b><br/>{ppa.numero}",
            self.styles['CustomTitle']
        ))
        
        return elements

    def _criar_cabecalho_capa(self, ppa):
        """Cria o cabeçalho simples da capa"""
        elements = []

        tabela = Table(
            [
                [f"PROCESSO Nº: {ppa.numero}", f"SIGLA: {ppa.get_sigla_display()}"]
            ],
            colWidths=[4.2*inch, 2.3*inch]
        )
        tabela.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ]))
        elements.append(tabela)
        return elements
    
    def _criar_dados_ppa(self, ppa):
        """Cria seção com dados do PPA"""
        elements = []
        
        elements.append(Paragraph("DADOS DO PROCESSO", self.styles['CustomHeading']))
        
        # Tabela com dados
        dados = [
            ["<b>Número:</b>", ppa.numero],
            ["<b>Sigla:</b>", ppa.get_sigla_display()],
            ["<b>Status:</b>", ppa.get_status_display()],
            ["<b>Interessado:</b>", ppa.interessado],
        ]
        
        if ppa.cnpj_interessado:
            dados.append(["<b>CNPJ:</b>", ppa.cnpj_interessado])
        
        if ppa.endereco_interessado:
            dados.append(["<b>Endereço:</b>", ppa.endereco_interessado])
        
        dados.extend([
            ["<b>Analista:</b>", ppa.analista_responsavel.get_full_name()],
            ["<b>Data de Criação:</b>", ppa.criado_em.strftime('%d/%m/%Y %H:%M')],
        ])
        
        if ppa.prazo_analise:
            dados.append(["<b>Prazo de Análise:</b>", ppa.prazo_analise.strftime('%d/%m/%Y')])
        
        if ppa.data_conclusao:
            dados.append(["<b>Data de Conclusão:</b>", ppa.data_conclusao.strftime('%d/%m/%Y %H:%M')])
        
        table = Table(dados, colWidths=[2*inch, 4.5*inch])
        table.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('ALIGN', (0, 0), (0, -1), 'LEFT'),
            ('ALIGN', (1, 0), (1, -1), 'LEFT'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        
        elements.append(table)
        elements.append(Spacer(1, 0.15*inch))
        
        # Assunto
        elements.append(Paragraph("<b>ASSUNTO:</b>", self.styles['Label']))
        elements.append(Paragraph(ppa.assunto, self.styles['CustomBody']))
        
        if ppa.observacoes:
            elements.append(Spacer(1, 0.1*inch))
            elements.append(Paragraph("<b>OBSERVAÇÕES:</b>", self.styles['Label']))
            elements.append(Paragraph(ppa.observacoes, self.styles['CustomBody']))
        
        return elements

    def _criar_campos_capa(self, ppa):
        """Cria os campos da capa no estilo do template PPA"""
        elements = []

        elements.append(Paragraph("<b>ASSUNTO:</b>", self.styles['Label']))
        elements.append(Paragraph(ppa.assunto or "-", self.styles['CustomBody']))
        elements.append(Spacer(1, 0.1*inch))

        elements.append(Paragraph("<b>INTERESSADO:</b>", self.styles['Label']))
        elements.append(Paragraph(ppa.interessado or "-", self.styles['CustomBody']))
        elements.append(Spacer(1, 0.1*inch))

        anexos = []
        for anexo in ppa.anexos.order_by('data_anexacao'):
            numero = anexo.numero_documento or anexo.get_tipo_documento_display()
            detalhe = f"{numero}".strip()
            if anexo.descricao:
                detalhe = f"{detalhe} - {anexo.descricao}"
            anexos.append(detalhe)
        anexos_texto = "; ".join(anexos) if anexos else "-"

        elements.append(Paragraph("<b>ANEXO:</b>", self.styles['Label']))
        elements.append(Paragraph(anexos_texto, self.styles['CustomBody']))

        return elements

    def _criar_tabela_movimento_capa(self, ppa):
        """Cria a tabela de movimento do processo para a capa"""
        elements = []

        elements.append(Paragraph("MOVIMENTO DO PROCESSO", self.styles['CustomHeading']))

        movimentacoes = ppa.movimentacoes.order_by('data', 'hora')

        if not movimentacoes.exists():
            elements.append(Paragraph("Nenhuma movimentação registrada.", self.styles['Normal']))
            return elements

        data = [
            [
                Paragraph("<b>DATA</b>", self.styles['Normal']),
                Paragraph("<b>HORA</b>", self.styles['Normal']),
                Paragraph("<b>ATENDIMENTO</b>", self.styles['Normal'])
            ]
        ]

        for mov in movimentacoes:
            hora = mov.hora.strftime('%H:%M') if mov.hora else '-'
            data.append([
                mov.data.strftime('%d/%m/%Y') if mov.data else '-',
                hora,
                Paragraph(mov.atendimento, self.styles['CustomBody'])
            ])

        table = Table(data, colWidths=[1.2*inch, 0.9*inch, 4.4*inch])
        table.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
        ]))

        elements.append(table)
        return elements
    
    def _criar_tabela_movimentacoes(self, ppa):
        """Cria tabela de movimentações"""
        elements = []
        
        elements.append(Paragraph("MOVIMENTAÇÃO DO PROCESSO", self.styles['CustomHeading']))
        
        movimentacoes = ppa.movimentacoes.order_by('data', 'hora')
        
        if not movimentacoes.exists():
            elements.append(Paragraph("Nenhuma movimentação registrada.", self.styles['Normal']))
            return elements
        
        # Cabeçalho da tabela
        data = [
            [
                Paragraph("<b>DATA</b>", self.styles['Normal']),
                Paragraph("<b>HORA</b>", self.styles['Normal']),
                Paragraph("<b>ATENDIMENTO</b>", self.styles['Normal'])
            ]
        ]
        
        # Linhas da tabela
        for mov in movimentacoes:
            data.append([
                mov.data.strftime('%d/%m/%Y') if mov.data else '-',
                mov.hora or '-',
                Paragraph(mov.atendimento, self.styles['CustomBody'])
            ])
        
        table = Table(data, colWidths=[1*inch, 0.8*inch, 4.7*inch])
        table.setStyle(TableStyle([
            # Cabeçalho
            ('BACKGROUND', (0, 0), (-1, 0), HexColor('#34495e')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
            
            # Corpo
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('ALIGN', (0, 1), (1, -1), 'CENTER'),
            ('ALIGN', (2, 1), (2, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('FONTSIZE', (0, 1), (-1, -1), 9),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, HexColor('#f8f9fa')]),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 1), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
        ]))
        
        elements.append(table)
        
        return elements
    
    def _criar_lista_anexos(self, ppa):
        """Cria lista de anexos"""
        elements = []
        
        elements.append(Paragraph("ANEXOS", self.styles['CustomHeading']))
        
        anexos = ppa.anexos.order_by('data_anexacao')
        
        for idx, anexo in enumerate(anexos, 1):
            texto = f"{idx}. <b>{anexo.get_tipo_documento_display()}</b>"
            if anexo.numero_documento:
                texto += f" - {anexo.numero_documento}"
            if anexo.descricao:
                texto += f"<br/><i>{anexo.descricao}</i>"
            texto += f"<br/><font size=8 color='#7f8c8d'>Anexado em: {anexo.data_anexacao.strftime('%d/%m/%Y %H:%M')} por {anexo.anexado_por.get_full_name()}</font>"
            
            elements.append(Paragraph(texto, self.styles['CustomBody']))
            elements.append(Spacer(1, 0.1*inch))
        
        return elements
    
    def _criar_pareceres(self, ppa):
        """Cria seção de pareceres"""
        elements = []
        
        elements.append(Paragraph("PARECERES", self.styles['CustomHeading']))
        
        pareceres = ppa.pareceres.order_by('criado_em')
        
        for parecer in pareceres:
            elements.append(Paragraph(
                f"<b>{parecer.numero_parecer}</b> - {parecer.titulo}",
                self.styles['CustomBody']
            ))
            
            elements.append(Paragraph(
                f"<i>Elaborado por: {parecer.elaborado_por.get_full_name()} em {parecer.criado_em.strftime('%d/%m/%Y')}</i>",
                self.styles['Label']
            ))
            
            elements.append(Spacer(1, 0.05*inch))
            
            elements.append(Paragraph("<b>Relatório:</b>", self.styles['Label']))
            elements.append(Paragraph(parecer.relatorio, self.styles['CustomBody']))
            
            if parecer.fundamentacao:
                elements.append(Spacer(1, 0.05*inch))
                elements.append(Paragraph("<b>Fundamentação:</b>", self.styles['Label']))
                elements.append(Paragraph(parecer.fundamentacao, self.styles['CustomBody']))
            
            elements.append(Spacer(1, 0.05*inch))
            elements.append(Paragraph(
                f"<b>Conclusão:</b> {parecer.get_conclusao_display()}",
                self.styles['Label']
            ))
            
            if parecer.aprovado_por:
                elements.append(Paragraph(
                    f"<i>✓ Aprovado por {parecer.aprovado_por.get_full_name()} em {parecer.data_aprovacao.strftime('%d/%m/%Y')}</i>",
                    self.styles['Label']
                ))
            
            elements.append(Spacer(1, 0.15*inch))
        
        return elements
    
    def _criar_decisao_final(self, ppa):
        """Cria seção de decisão final"""
        elements = []
        
        elements.append(Paragraph("DECISÃO FINAL", self.styles['CustomHeading']))
        
        elements.append(Paragraph(
            f"<b>Decisão:</b> {ppa.get_decisao_final_display()}",
            self.styles['CustomBody']
        ))
        
        if ppa.fundamentacao_decisao:
            elements.append(Spacer(1, 0.05*inch))
            elements.append(Paragraph("<b>Fundamentação:</b>", self.styles['Label']))
            elements.append(Paragraph(ppa.fundamentacao_decisao, self.styles['CustomBody']))
        
        return elements
    
    def _criar_rodape(self):
        """Cria rodapé do relatório"""
        elements = []
        
        elements.append(Spacer(1, 0.3*inch))
        elements.append(Paragraph(
            f"<i>Relatório gerado automaticamente em {timezone.now().strftime('%d/%m/%Y às %H:%M')}</i>",
            ParagraphStyle(
                'Footer',
                parent=self.styles['Normal'],
                fontSize=8,
                textColor=HexColor('#7f8c8d'),
                alignment=TA_CENTER
            )
        ))
        
        return elements


def gerar_pdf_ppa(ppa_id):
    """
    Função auxiliar para gerar PDF do PPA (capa simplificada)
    
    Args:
        ppa_id: ID do PPA
        
    Returns:
        BytesIO com o PDF gerado
    """
    docx_bytes = gerar_docx_ppa(ppa_id)
    pdf_bytes = _converter_docx_para_pdf(docx_bytes, f"PPA_{ppa_id}")
    if pdf_bytes:
        buffer = BytesIO(pdf_bytes)
        buffer.seek(0)
        return buffer

    logger.warning("Conversao DOCX->PDF indisponivel. Usando renderizacao simplificada.")
    generator = PPAReportGenerator()
    return generator.gerar_pdf_capa(ppa_id)


def _set_cell_text(cell, text):
    cell.text = text or ''


def _set_cell_label_value(cell, label, value):
    cell.text = label
    if value:
        cell.add_paragraph(str(value))


def _converter_docx_para_pdf(docx_bytes, nome_base):
    """Converte um DOCX em PDF e retorna os bytes, ou None se falhar."""
    if not docx_bytes:
        return None
    safe_name = (nome_base or 'PPA').replace('/', '_')
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = Path(tmpdir) / f"{safe_name}.docx"
        pdf_path = Path(tmpdir) / f"{safe_name}.pdf"
        docx_path.write_bytes(docx_bytes)

        try:
            from docx2pdf import convert

            convert(str(docx_path), str(tmpdir))
            if pdf_path.exists():
                return pdf_path.read_bytes()
        except Exception:
            pass

        converter = shutil.which('soffice') or shutil.which('libreoffice')
        if converter:
            try:
                subprocess.run(
                    [converter, '--headless', '--convert-to', 'pdf', '--outdir', tmpdir, str(docx_path)],
                    check=True,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                )
                if pdf_path.exists():
                    return pdf_path.read_bytes()
            except Exception:
                return None
    return None


def gerar_docx_ppa(ppa_id):
    """
    Gera DOCX de capa do PPA usando template.
    """
    try:
        ppa = ProcedimentoPreAdministrativo.objects.select_related(
            'analista_responsavel',
            'supervisor',
        ).prefetch_related(
            'movimentacoes',
            'anexos',
            'pareceres'
        ).get(id=ppa_id)
    except ProcedimentoPreAdministrativo.DoesNotExist:
        raise ValueError(f'PPA {ppa_id} não encontrado')

    template_path = Path(settings.BASE_DIR) / 'ppa' / 'templates' / 'docs' / 'PPA.docx'
    doc = Document(template_path)

    if doc.tables:
        # Tabela 0: Processo e Sigla
        table0 = doc.tables[0]
        if len(table0.rows) > 1 and len(table0.columns) > 1:
            _set_cell_text(table0.cell(1, 0), ppa.numero)
            _set_cell_text(table0.cell(1, 1), ppa.get_sigla_display())

        # Assunto, Interessado, Anexo
        if len(doc.tables) > 1:
            _set_cell_label_value(doc.tables[1].cell(0, 0), "ASSUNTO:", ppa.assunto or "-")
        if len(doc.tables) > 2:
            _set_cell_label_value(doc.tables[2].cell(0, 0), "INTERESSADO:", ppa.interessado or "-")
        if len(doc.tables) > 3:
            anexos = []
            for anexo in ppa.anexos.order_by('data_anexacao'):
                numero = anexo.numero_documento or anexo.get_tipo_documento_display()
                detalhe = f"{numero}".strip()
                if anexo.descricao:
                    detalhe = f"{detalhe} - {anexo.descricao}"
                anexos.append(detalhe)
            anexos_texto = "; ".join(anexos) if anexos else "-"
            _set_cell_label_value(doc.tables[3].cell(0, 0), "ANEXO:", anexos_texto)

        # Tabela de movimento
        if len(doc.tables) > 4:
            table_mov = doc.tables[4]
            linhas_disponiveis = max(len(table_mov.rows) - 2, 0)
            movimentacoes = list(ppa.movimentacoes.order_by('data', 'hora'))

            # Adiciona linhas extras se necessário
            while len(movimentacoes) > linhas_disponiveis:
                table_mov.add_row()
                linhas_disponiveis += 1

            for idx, mov in enumerate(movimentacoes):
                row = table_mov.rows[2 + idx]
                _set_cell_text(row.cells[0], mov.data.strftime('%d/%m/%Y') if mov.data else '-')
                hora = mov.hora.strftime('%H:%M') if mov.hora else '-'
                _set_cell_text(row.cells[3], hora)
                _set_cell_text(row.cells[5], mov.atendimento or '')

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

